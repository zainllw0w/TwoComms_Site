from collections import OrderedDict
from datetime import datetime, timedelta
import json
import logging
import re
import time
from pathlib import Path

from django.conf import settings
from django.shortcuts import render, redirect
from django.utils import timezone
from django.utils.html import escape
from django.contrib.auth.decorators import login_required
from django.contrib.auth import get_user_model
from django.db import transaction
from django.db.models import Count, Q
from django.http import JsonResponse, FileResponse, HttpResponse
from django.core.exceptions import ValidationError
from django.core.files.base import ContentFile
from django.core.mail import EmailMultiAlternatives
from django.core.validators import validate_email
from django.template.loader import render_to_string
from django.urls import reverse
from django.views.decorators.http import require_POST

import requests
from io import BytesIO
import os
import secrets

from docx import Document

from .forms import CommercialOfferEmailForm
from .models import (
    Client,
    ClientFollowUp,
    CommercialOfferEmailLog,
    CommercialOfferEmailSettings,
    ContractSequence,
    ContractRejectionReasonRequest,
    InvoiceRejectionReasonRequest,
    ManagementContract,
    ManagementLead,
    ReminderRead,
    ReminderSent,
    Report,
    Shop,
)
from accounts.models import UserProfile
from django.views.decorators.csrf import csrf_exempt

from .email_templates.twocomms_cp import (
    build_twocomms_cp_email,
    get_twocomms_cp_unit_defaults,
    OPT_TIER_LABELS,
    OPT_TIER_WHOLESALE_TEE,
    OPT_TIER_WHOLESALE_HOODIE,
    DROP_FIXED_TEE_PRICE,
    DROP_FIXED_HOODIE_PRICE,
)

from .constants import LEAD_BASE_PROCESSING_PENALTY, POINTS, TARGET_CLIENTS_DAY, TARGET_POINTS_DAY
from .invoice_service import (
    InvoicePayloadError,
    build_management_invoice_workbook,
    get_management_wholesale_price_context,
    normalize_management_invoice_payload,
    serialize_management_invoice_payload,
)
from .lead_services import calc_client_points, get_user_lead_bonus_points
from .services.dtf_bridge import build_dtf_bridge_payload
from .services.forecast import build_admin_economics_summary, build_salary_simulator
from .services.followups import build_reminder_digest

_BOT_USERNAME_CACHE = {"username": "", "ts": 0, "token": ""}
logger = logging.getLogger(__name__)


class ContractPayloadError(Exception):
    def __init__(self, code, fields=None):
        super().__init__(code)
        self.code = code
        self.fields = fields or []


def _load_env_tokens():
    """Best-effort load bot tokens from .env.production if process env is empty (shared hosting)."""
    if os.environ.get("MANAGER_TG_BOT_TOKEN") or os.environ.get("MANAGEMENT_TG_BOT_TOKEN"):
        return
    env_path = Path(__file__).resolve().parents[2] / ".env.production"
    if env_path.exists():
        try:
            for line in env_path.read_text().splitlines():
                if not line or line.strip().startswith("#"):
                    continue
                if "=" not in line:
                    continue
                k, v = line.split("=", 1)
                k = k.strip()
                v = v.strip().strip('"').strip("'")
                if k and v and k not in os.environ:
                    os.environ[k] = v
        except Exception:
            pass


_load_env_tokens()


def _split_chat_ids(raw_value):
    if not raw_value:
        return []
    parts = re.split(r"[;,\s]+", str(raw_value))
    return [part for part in (p.strip() for p in parts) if part]


def _get_management_admin_chat_ids():
    return _split_chat_ids(os.environ.get("MANAGEMENT_TG_ADMIN_CHAT_ID"))


def user_is_management(user):
    if not user.is_authenticated:
        return False
    try:
        prof = user.userprofile
        is_manager = getattr(prof, 'is_manager', False)
    except UserProfile.DoesNotExist:
        is_manager = False
    return user.is_staff or user.is_superuser or is_manager


def calc_points(qs):
    return calc_client_points(qs)


def get_manager_bot_username():
    """Try MANAGER_TG_BOT_USERNAME, else fetch via token. Fall back to report bot token if needed."""
    name = os.environ.get("MANAGER_TG_BOT_USERNAME", "")
    if name:
        return name
    token = os.environ.get("MANAGER_TG_BOT_TOKEN") or os.environ.get("MANAGEMENT_TG_BOT_TOKEN")
    if not token:
        return ""
    now = time.time()
    if _BOT_USERNAME_CACHE["username"] and _BOT_USERNAME_CACHE["token"] == token and now - _BOT_USERNAME_CACHE["ts"] < 600:
        return _BOT_USERNAME_CACHE["username"]
    try:
        resp = requests.get(f"https://api.telegram.org/bot{token}/getMe", timeout=4)
        data = resp.json()
        if data.get("ok") and data.get("result", {}).get("username"):
            username = data["result"]["username"]
            _BOT_USERNAME_CACHE.update({"username": username, "ts": now, "token": token})
            return username
    except Exception:
        pass
    return ""


def get_today_range():
    """Localized start/end of current day for consistent filters."""
    now = timezone.localtime(timezone.now())
    start = now.replace(hour=0, minute=0, second=0, microsecond=0)
    end = start + timedelta(days=1)
    return start, end


def has_report_today(user):
    start, end = get_today_range()
    return Report.objects.filter(owner=user, created_at__gte=start, created_at__lt=end).exists()


def _local_date_from_dt(dt_value):
    try:
        return timezone.localtime(dt_value).date()
    except Exception:
        try:
            return timezone.localdate(dt_value)
        except Exception:
            return timezone.localdate()


def _sync_client_followup(client: Client, prev_next_call_at, new_next_call_at, now_dt):
    """
    Keep ClientFollowUp in sync with Client.next_call_at.
    - Creates a follow-up when next_call_at is set.
    - Closes open follow-ups when next_call_at changes/clears.
    """
    owner = client.owner
    if not owner:
        return

    if prev_next_call_at:
        if not new_next_call_at:
            ClientFollowUp.objects.filter(client=client, owner=owner, status=ClientFollowUp.Status.OPEN).update(
                status=ClientFollowUp.Status.CANCELLED,
                closed_at=now_dt,
            )
        elif new_next_call_at != prev_next_call_at:
            status = ClientFollowUp.Status.RESCHEDULED if now_dt < prev_next_call_at else ClientFollowUp.Status.DONE
            ClientFollowUp.objects.filter(client=client, owner=owner, status=ClientFollowUp.Status.OPEN).update(
                status=status,
                closed_at=now_dt,
            )

    if new_next_call_at and (not prev_next_call_at or new_next_call_at != prev_next_call_at):
        ClientFollowUp.objects.create(
            client=client,
            owner=owner,
            due_at=new_next_call_at,
            due_date=_local_date_from_dt(new_next_call_at),
            status=ClientFollowUp.Status.OPEN,
            meta={"source": "client.next_call_at"},
        )


def _close_followups_for_report(report: Report):
    """Freeze today's follow-ups when daily report is sent."""
    report_day = _local_date_from_dt(report.created_at)
    ClientFollowUp.objects.filter(
        owner=report.owner,
        due_date=report_day,
        status=ClientFollowUp.Status.OPEN,
    ).update(
        status=ClientFollowUp.Status.MISSED,
        closed_at=report.created_at,
        closed_by_report=report,
    )


def _time_label(dt_local, now):
    delta = (now - dt_local).total_seconds()
    if delta >= 0:
        minutes = int(delta // 60)
        hours = int(delta // 3600)
        days = int(delta // 86400)
        if minutes < 1:
            return "щойно"
        if minutes < 60:
            return f"{minutes} хв тому"
        if hours < 24:
            return f"{hours} год тому"
        return f"{days} дн тому"
    else:
        delta_abs = abs(delta)
        minutes = int(delta_abs // 60)
        hours = int(delta_abs // 3600)
        days = int(delta_abs // 86400)
        if minutes < 1:
            return "за кілька секунд"
        if minutes < 60:
            return f"через {minutes} хв"
        if hours < 24:
            return f"через {hours} год"
        return f"через {days} дн"


def get_reminders(user, stats=None, report_sent=False):
    """Return upcoming/due follow-ups + report reminder."""
    digest = build_reminder_digest(user, stats=stats, report_sent=report_sent)
    reminders = digest["reminders"]
    now = timezone.localtime(timezone.now())
    reminders.sort(key=lambda x: x.get('dt', now), reverse=True)
    return reminders


def get_user_stats(user):
    base_qs = Client.objects.filter(owner=user)
    today_start, today_end = get_today_range()
    today_qs = base_qs.filter(created_at__gte=today_start, created_at__lt=today_end)
    lead_bonus_today, lead_bonus_total = get_user_lead_bonus_points(user, today_start, today_end)
    points_today = calc_points(today_qs) + lead_bonus_today
    points_total = calc_points(base_qs) + lead_bonus_total
    return {
        'points_today': points_today,
        'points_total': points_total,
        'lead_bonus_today': lead_bonus_today,
        'lead_bonus_total': lead_bonus_total,
        'processed_today': today_qs.count(),
        'processed_total': base_qs.count(),
    }


def build_report_excel(user, stats, clients):
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Звіт"

    title_font = Font(size=14, bold=True)
    header_font = Font(size=12, bold=True)

    ws["A1"] = "Звіт менеджера"
    ws["A1"].font = title_font
    ws.merge_cells("A1:D1")

    ws["A2"] = "Менеджер"
    ws["B2"] = user.get_full_name() or user.username
    ws["A3"] = "Дата"
    ws["B3"] = timezone.localtime().strftime("%d.%m.%Y %H:%M")
    ws["A4"] = "Бали за сьогодні"
    ws["B4"] = stats['points_today']
    ws["A5"] = "Оброблено за сьогодні"
    ws["B5"] = stats['processed_today']
    ws["A6"] = "Бали всього"
    ws["B6"] = stats['points_total']
    ws["A7"] = "Оброблено всього"
    ws["B7"] = stats['processed_total']

    ws.append([])
    ws.append([
        "Магазин / Insta", "Телефон", "ПІБ", "Статус",
        "Джерело", "Підсумок", "Деталі", "Наступний дзвінок", "Створено"
    ])
    header_row = ws.max_row
    for col in range(1, 10):
        ws.cell(row=header_row, column=col).font = header_font
        ws.cell(row=header_row, column=col).fill = PatternFill(start_color="1f2937", end_color="1f2937", fill_type="solid")
        ws.cell(row=header_row, column=col).alignment = Alignment(horizontal="center")

    for c in clients:
        ws.append([
            c.shop_name,
            c.phone,
            c.full_name,
            c.get_role_display(),
            c.source,
            c.get_call_result_display(),
            c.call_result_details,
            timezone.localtime(c.next_call_at).strftime("%d.%m.%Y %H:%M") if c.next_call_at else "–",
            timezone.localtime(c.created_at).strftime("%d.%m.%Y %H:%M"),
        ])

    widths = [22, 16, 22, 12, 16, 20, 24, 18, 18]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)
    return bio.getvalue()


def send_telegram_report(user, stats, clients, file_bytes, filename):
    token = os.environ.get("MANAGEMENT_TG_BOT_TOKEN")
    chat_ids = _get_management_admin_chat_ids()
    if not token or not chat_ids:
        return
    text = (
        f"Звіт менеджера\n"
        f"👤 {user.get_full_name() or user.username}\n"
        f"📅 {timezone.localtime().strftime('%d.%m.%Y %H:%M')}\n"
        f"Бали за сьогодні: {stats['points_today']}\n"
        f"Оброблено за сьогодні: {stats['processed_today']}\n"
        f"Клієнтів всього: {stats['processed_total']}"
    )
    url = f"https://api.telegram.org/bot{token}/sendDocument"
    files = {
        'document': (filename, file_bytes, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    }
    for chat_id in chat_ids:
        data = {
            'chat_id': chat_id,
            'caption': text,
            'parse_mode': 'HTML'
        }
        try:
            requests.post(url, data=data, files=files, timeout=10)
        except Exception:
            pass


@login_required(login_url='management_login')
def home(request):
    if not user_is_management(request.user):
        return redirect('management_login')
    if request.method == 'POST':
        data = request.POST
        client_id = data.get('client_id')
        shop_name = data.get('shop_name', '').strip()
        phone = data.get('phone', '').strip()
        website_url = data.get('website_url', '').strip()
        full_name = data.get('full_name', '').strip()
        role = data.get('role', Client.Role.MANAGER)
        role_custom = data.get('role_custom', '').strip()
        source = data.get('source', '').strip()
        source_link = data.get('source_link', '').strip()
        source_other = data.get('source_other', '').strip()
        call_result = data.get('call_result', Client.CallResult.THINKING)
        call_result_other = data.get('call_result_other', '').strip()
        next_call_type = data.get('next_call_type', 'scheduled')
        next_call_date = data.get('next_call_date', '').strip()
        next_call_time = data.get('next_call_time', '').strip()
        today = timezone.localdate()

        source_display = {
            'instagram': 'Instagram',
            'prom_ua': 'Prom.ua',
            'google_maps': 'Google Карти',
            'forums': f"Сайти/Форуми: {source_link}" if source_link else 'Сайти/Форуми',
            'other': source_other or 'Інше',
        }.get(source, source or '')

        call_result_details_parts = []
        if role == Client.Role.OTHER and role_custom:
            call_result_details_parts.append(f"Роль: {role_custom}")
        if call_result == Client.CallResult.OTHER and call_result_other:
            call_result_details_parts.append(f"Інше: {call_result_other}")

        next_call_at = None
        if next_call_type == 'scheduled' and next_call_date and next_call_time:
            try:
                naive = datetime.strptime(f"{next_call_date} {next_call_time}", "%Y-%m-%d %H:%M")
                next_call_at = timezone.make_aware(naive, timezone.get_current_timezone())
            except ValueError:
                next_call_at = None
        elif next_call_type == 'no_follow':
            next_call_at = None

        if not full_name:
            full_name = 'ПІБ не вказано'

        if shop_name and phone:
            details = "\n".join(call_result_details_parts)
            if client_id:
                try:
                    client = Client.objects.get(id=client_id)
                    if request.user.is_staff or client.owner == request.user:
                        prev_next_call_at = client.next_call_at
                        client.shop_name = shop_name
                        client.phone = phone
                        client.website_url = website_url
                        client.full_name = full_name
                        client.role = role
                        client.source = source_display
                        client.call_result = call_result
                        client.call_result_details = details
                        client.next_call_at = next_call_at
                        client.owner = client.owner or request.user
                        if client.points_override is not None:
                            client.points_override = max(0, int(POINTS.get(call_result, 0)) - LEAD_BASE_PROCESSING_PENALTY)
                        client.save()
                        _sync_client_followup(client, prev_next_call_at, client.next_call_at, timezone.now())
                except Client.DoesNotExist:
                    pass
            else:
                client = Client.objects.create(
                    shop_name=shop_name,
                    phone=phone,
                    website_url=website_url,
                    full_name=full_name,
                    role=role,
                    source=source_display,
                    call_result=call_result,
                    call_result_details=details,
                    next_call_at=next_call_at,
                    owner=request.user,
                )
                _sync_client_followup(client, None, client.next_call_at, timezone.now())
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            # Сформируем актуальные данные после операции
            stats = get_user_stats(request.user)
            user_points_today = stats['points_today']
            user_points_total = stats['points_total']
            processed_today = stats['processed_today']
            progress_clients_pct = min(100, int(processed_today / TARGET_CLIENTS_DAY * 100)) if TARGET_CLIENTS_DAY else 0
            progress_points_pct = min(100, int(user_points_today / TARGET_POINTS_DAY * 100)) if TARGET_POINTS_DAY else 0

            latest = Client.objects.filter(owner=request.user).order_by('-created_at').first()
            latest_data = None
            if latest:
                latest_data = {
                    'id': latest.id,
                    'shop': latest.shop_name,
                    'phone': latest.phone,
                    'website_url': latest.website_url,
                    'full_name': latest.full_name,
                    'role': latest.role,
                    'role_display': latest.get_role_display(),
                    'source': latest.source,
                    'call_result': latest.call_result,
                    'call_result_display': latest.get_call_result_display(),
                    'call_result_details': latest.call_result_details,
                    'next_call': timezone.localtime(latest.next_call_at).strftime('%d.%m.%Y %H:%M') if latest.next_call_at else '–',
                    'created_date_label': 'Сьогодні' if timezone.localtime(latest.created_at).date() == today else timezone.localtime(latest.created_at).strftime('%d.%m.%Y'),
                }

            return JsonResponse({
                'success': True,
                'user_points_today': user_points_today,
                'user_points_total': user_points_total,
                'processed_today': processed_today,
                'target_clients': TARGET_CLIENTS_DAY,
                'target_points': TARGET_POINTS_DAY,
                'progress_clients_pct': progress_clients_pct,
                'progress_points_pct': progress_points_pct,
                'latest': latest_data,
            })
        return redirect('management_home')

    base_qs = Client.objects.select_related('owner').order_by('-created_at')
    # У основній панелі всі бачать тільки свої записи
    clients_qs = base_qs.filter(owner=request.user)
    clients = clients_qs

    today = timezone.localdate()
    today_start, today_end = get_today_range()
    yesterday = today - timedelta(days=1)
    grouped = OrderedDict()

    for client in clients:
        local_date = timezone.localtime(client.created_at).date()
        if local_date == today:
            label = 'Сьогодні'
        elif local_date == yesterday:
            label = 'Вчора'
        else:
            label = local_date.strftime('%d.%m.%Y')

        grouped.setdefault(label, []).append(client)

    grouped_clients = list(grouped.items())
    base_leads = ManagementLead.objects.filter(status=ManagementLead.Status.BASE).select_related('added_by').order_by('-created_at')[:400]

    user_stats = get_user_stats(request.user)
    user_points_today = user_stats['points_today']
    user_points_total = user_stats['points_total']
    processed_today = user_stats['processed_today']
    report_sent_today = has_report_today(request.user)
    reminders = get_reminders(request.user, stats=user_stats, report_sent=report_sent_today)

    progress_clients_pct = min(100, int(processed_today / TARGET_CLIENTS_DAY * 100)) if TARGET_CLIENTS_DAY else 0
    progress_points_pct = min(100, int(user_points_today / TARGET_POINTS_DAY * 100)) if TARGET_POINTS_DAY else 0

    bot_username = get_manager_bot_username()
    return render(request, 'management/home.html', {
        'grouped_clients': grouped_clients,
        'base_leads': base_leads,
        'user_points_today': user_points_today,
        'user_points_total': user_points_total,
        'processed_today': processed_today,
        'target_clients': TARGET_CLIENTS_DAY,
        'target_points': TARGET_POINTS_DAY,
        'progress_clients_pct': progress_clients_pct,
        'progress_points_pct': progress_points_pct,
        'has_report_today': report_sent_today,
        'reminders': reminders,
        'manager_bot_username': bot_username,
    })


@login_required(login_url='management_login')
def delete_client(request, client_id):
    if not user_is_management(request.user):
        return redirect('management_login')
    if request.method != 'POST':
        return redirect('management_home')
    try:
        client = Client.objects.get(id=client_id)
        if request.user.is_staff or client.owner == request.user:
            client.delete()
    except Client.DoesNotExist:
        pass
    return redirect('management_home')


@login_required(login_url='management_login')
def admin_overview(request):
    if not request.user.is_staff:
        return redirect('management_home')

    tab = (request.GET.get('tab') or 'managers').strip().lower()
    if tab not in ('managers', 'invoices', 'shops', 'payouts', 'dtf'):
        tab = 'managers'

    User = get_user_model()
    today_start, today_end = get_today_range()
    admin_user_data = []
    users = User.objects.filter(is_active=True).filter(
        Q(is_staff=True) | Q(userprofile__is_manager=True)
    ).annotate(
        total_clients=Count('management_clients', distinct=True),
        today_clients=Count('management_clients', filter=Q(management_clients__created_at__gte=today_start), distinct=True),
    ).distinct()

    stats = get_user_stats(request.user)
    user_points_today = stats['points_today']
    user_points_total = stats['points_total']
    processed_today = stats['processed_today']
    report_sent_today = has_report_today(request.user)
    progress_clients_pct = min(100, int(processed_today / TARGET_CLIENTS_DAY * 100)) if TARGET_CLIENTS_DAY else 0
    progress_points_pct = min(100, int(user_points_today / TARGET_POINTS_DAY * 100)) if TARGET_POINTS_DAY else 0

    if tab == 'managers':
        for u in users:
            last_login = u.last_login
            online = False
            last_login_local = None
            if last_login:
                last_login_local = timezone.localtime(last_login)
                online = (timezone.now() - last_login) <= timedelta(minutes=5)
            user_clients = Client.objects.filter(owner=u).order_by('-created_at')
            points_stats = get_user_stats(u)
            admin_user_data.append({
                'id': u.id,
                'name': u.get_full_name() or u.username,
                'role': 'Адмін' if u.is_staff else 'Менеджер',
                'today': u.today_clients,
                'total': u.total_clients,
                'points_today': points_stats['points_today'],
                'points_total': points_stats['points_total'],
                'online': online,
                'last_login': last_login_local.strftime('%d.%m.%Y %H:%M') if last_login_local else 'Немає даних',
            })

    bot_username = get_manager_bot_username()
    reminders = get_reminders(request.user, stats={
        'points_today': user_points_today,
        'processed_today': processed_today,
    }, report_sent=report_sent_today)

    ctx = {
        'admin_tab': tab,
        'admin_user_data': admin_user_data,
        'target_clients': TARGET_CLIENTS_DAY,
        'target_points': TARGET_POINTS_DAY,
        'user_points_today': user_points_today,
        'user_points_total': user_points_total,
        'processed_today': processed_today,
        'progress_clients_pct': progress_clients_pct,
        'progress_points_pct': progress_points_pct,
        'reminders': reminders,
        'manager_bot_username': bot_username,
    }

    if tab == 'managers':
        ctx['admin_analytics'] = build_admin_economics_summary()

    if tab == 'dtf':
        ctx['dtf_bridge'] = build_dtf_bridge_payload()

    if tab == 'invoices':
        from orders.models import WholesaleInvoice
        invoices_for_review = WholesaleInvoice.objects.filter(
            created_by__isnull=False,
            review_status='pending',
        ).select_related('created_by', 'created_by__userprofile').order_by('-created_at')[:200]
        ctx['invoices_for_review'] = invoices_for_review

    if tab == 'payouts':
        import re
        from decimal import Decimal

        from django.db import models
        from django.db.models import Max, Min, Sum
        from django.db.models.functions import Coalesce

        from management.models import ManagerCommissionAccrual, ManagerPayoutRequest

        money_field = models.DecimalField(max_digits=12, decimal_places=2)
        zero = Decimal('0')
        now = timezone.now()

        payout_users = User.objects.filter(
            is_active=True,
            userprofile__is_manager=True,
        ).select_related('userprofile').order_by('id')

        accr_rows = ManagerCommissionAccrual.objects.filter(owner__in=payout_users).values('owner_id').annotate(
            total=Coalesce(Sum('amount'), zero, output_field=money_field),
            frozen=Coalesce(Sum('amount', filter=Q(frozen_until__gt=now)), zero, output_field=money_field),
            total_count=Count('id'),
            frozen_count=Count('id', filter=Q(frozen_until__gt=now)),
            last_accrual_at=Max('created_at'),
            next_unfreeze_at=Min('frozen_until', filter=Q(frozen_until__gt=now)),
        )
        accr_map = {row['owner_id']: row for row in accr_rows}

        pay_rows = ManagerPayoutRequest.objects.filter(owner__in=payout_users).values('owner_id').annotate(
            paid=Coalesce(Sum('amount', filter=Q(status=ManagerPayoutRequest.Status.PAID)), zero, output_field=money_field),
            reserved=Coalesce(Sum('amount', filter=Q(status__in=[ManagerPayoutRequest.Status.PROCESSING, ManagerPayoutRequest.Status.APPROVED])), zero, output_field=money_field),
            total_requests=Count('id'),
            last_request_at=Max('created_at'),
            last_paid_at=Max('paid_at', filter=Q(status=ManagerPayoutRequest.Status.PAID)),
        )
        pay_map = {row['owner_id']: row for row in pay_rows}

        active_reqs = ManagerPayoutRequest.objects.filter(
            owner__in=payout_users,
            status__in=[ManagerPayoutRequest.Status.PROCESSING, ManagerPayoutRequest.Status.APPROVED],
        ).select_related('owner', 'owner__userprofile').order_by('owner_id', '-created_at')
        active_map = {}
        for r in active_reqs:
            if r.owner_id not in active_map:
                active_map[r.owner_id] = r

        frozen_details_map = {}
        frozen_rows = (
            ManagerCommissionAccrual.objects.filter(owner__in=payout_users, frozen_until__gt=now)
            .select_related('invoice')
            .order_by('owner_id', 'frozen_until')
        )
        for row in frozen_rows:
            items = frozen_details_map.setdefault(row.owner_id, [])
            invoice_number = ''
            if row.invoice_id:
                invoice_number = getattr(row.invoice, 'invoice_number', '') or ''
            reason = (row.note or '').strip()
            if not reason:
                if invoice_number:
                    reason = f"Накладна #{invoice_number}"
                else:
                    reason = "Ручне коригування"
            frozen_until_local = timezone.localtime(row.frozen_until) if row.frozen_until else None
            days_left = None
            if frozen_until_local:
                try:
                    days_left = max(0, (frozen_until_local.date() - timezone.localdate()).days)
                except Exception:
                    days_left = None
            items.append({
                'amount': row.amount,
                'frozen_until': frozen_until_local.strftime('%d.%m.%Y') if frozen_until_local else '—',
                'days_left': days_left,
                'reason': reason,
            })

        def mask_card(raw):
            digits = re.sub(r'\D+', '', (raw or '').strip())
            if len(digits) < 4:
                return '—'
            return '**** ' + digits[-4:]

        payouts_payload = []
        summary = {
            'managers': 0,
            'balance': zero,
            'available': zero,
            'frozen': zero,
            'reserved': zero,
            'paid_total': zero,
            'active_requests': 0,
        }
        for u in payout_users:
            prof = getattr(u, 'userprofile', None)
            a = accr_map.get(u.id) or {}
            p = pay_map.get(u.id) or {}
            total_accrued = a.get('total') or zero
            frozen_amount = a.get('frozen') or zero
            paid_total = p.get('paid') or zero
            reserved_amount = p.get('reserved') or zero

            balance = total_accrued - paid_total
            available = balance - frozen_amount - reserved_amount
            if available < 0:
                available = zero

            next_unfreeze_at = a.get('next_unfreeze_at')
            next_unfreeze_date = None
            next_unfreeze_label = None
            if next_unfreeze_at:
                next_unfreeze_local = timezone.localtime(next_unfreeze_at)
                next_unfreeze_date = next_unfreeze_local.strftime('%d.%m.%Y')
                try:
                    days_left = max(0, (next_unfreeze_local.date() - timezone.localdate()).days)
                except Exception:
                    days_left = None
                if days_left is None:
                    next_unfreeze_label = '—'
                elif days_left == 0:
                    next_unfreeze_label = 'Сьогодні'
                elif days_left == 1:
                    next_unfreeze_label = 'Через 1 день'
                else:
                    next_unfreeze_label = f"Через {days_left} дн"

            started_at = getattr(prof, 'manager_started_at', None) if prof else None
            if not started_at:
                try:
                    started_at = timezone.localtime(u.date_joined).date() if u.date_joined else None
                except Exception:
                    started_at = None

            days_worked = None
            try:
                if started_at:
                    days_worked = max(0, (timezone.localdate() - started_at).days)
            except Exception:
                days_worked = None

            summary['managers'] += 1
            summary['balance'] += balance
            summary['available'] += available
            summary['frozen'] += frozen_amount
            summary['reserved'] += reserved_amount
            summary['paid_total'] += paid_total
            if active_map.get(u.id):
                summary['active_requests'] += 1

            payouts_payload.append({
                'id': u.id,
                'name': u.get_full_name() or u.username,
                'position': (getattr(prof, 'manager_position', '') or '').strip() or '—',
                'base_salary': getattr(prof, 'manager_base_salary_uah', 0) if prof else 0,
                'percent': getattr(prof, 'manager_commission_percent', 0) if prof else 0,
                'started_at': started_at,
                'days_worked': days_worked,
                'card_mask': mask_card(getattr(prof, 'payment_details', '') if prof else ''),
                'balance': balance,
                'available': available,
                'frozen': frozen_amount,
                'reserved': reserved_amount,
                'active_request': active_map.get(u.id),
                'total_accrued': total_accrued,
                'paid_total': paid_total,
                'accruals_count': a.get('total_count') or 0,
                'frozen_count': a.get('frozen_count') or 0,
                'last_request_at': p.get('last_request_at'),
                'last_paid_at': p.get('last_paid_at'),
                'total_requests': p.get('total_requests') or 0,
                'next_unfreeze_date': next_unfreeze_date,
                'next_unfreeze_label': next_unfreeze_label,
                'frozen_items': frozen_details_map.get(u.id, []),
            })

        payout_requests = ManagerPayoutRequest.objects.select_related('owner', 'owner__userprofile').order_by('-created_at')[:200]

        ctx['payouts_payload'] = payouts_payload
        ctx['payout_requests'] = payout_requests
        ctx['payouts_summary'] = summary

    if tab == 'shops':
        from decimal import Decimal
        from datetime import time as dt_time, timedelta as dt_timedelta

        from django.core.paginator import Paginator
        from django.db import models
        from django.db.models import Sum
        from django.db.models.functions import Coalesce

        qs = Shop.objects.all().annotate(
            total_amount=Coalesce(
                Sum("shipments__invoice_total_amount"),
                Decimal("0"),
                output_field=models.DecimalField(max_digits=12, decimal_places=2),
            )
        )
        qs = qs.prefetch_related("phones", "shipments__wholesale_invoice").order_by("-created_at")
        paginator = Paginator(qs, 12)
        page_obj = paginator.get_page(request.GET.get("page") or 1)

        now_dt = timezone.localtime(timezone.now())
        shops_payload = []
        for shop in page_obj.object_list:
            primary_phone = None
            phones = []
            for p in shop.phones.all():
                phones.append(
                    {
                        "id": p.id,
                        "role": p.role,
                        "role_other": p.role_other,
                        "phone": p.phone,
                        "is_primary": bool(p.is_primary),
                        "sort_order": p.sort_order,
                    }
                )
                if p.is_primary and not primary_phone:
                    primary_phone = p.phone
            if not primary_phone and phones:
                primary_phone = phones[0]["phone"]

            shipments = []
            for s in shop.shipments.all():
                invoice_kind = "none"
                invoice_title = ""
                invoice_download_url = ""
                invoice_total = s.invoice_total_amount
                if s.wholesale_invoice_id:
                    invoice_kind = "system"
                    inv_num = ""
                    try:
                        inv_num = s.wholesale_invoice.invoice_number if s.wholesale_invoice else ""
                    except Exception:
                        inv_num = ""
                    invoice_title = f"#{inv_num}" if inv_num else "Накладна"
                    invoice_download_url = reverse("management_invoices_download", args=[s.wholesale_invoice_id])
                elif s.uploaded_invoice_file:
                    invoice_kind = "upload"
                    invoice_title = s.uploaded_invoice_file.name.split("/")[-1]
                    invoice_download_url = reverse("management_shop_shipment_invoice_download", args=[s.id])

                shipments.append(
                    {
                        "id": s.id,
                        "ttn_number": s.ttn_number,
                        "shipped_at": s.shipped_at.isoformat() if s.shipped_at else "",
                        "invoice_kind": invoice_kind,
                        "wholesale_invoice_id": s.wholesale_invoice_id,
                        "invoice_title": invoice_title,
                        "invoice_total_amount": str(invoice_total) if invoice_total is not None else "",
                        "invoice_download_url": invoice_download_url,
                    }
                )

            timer = None
            if shop.shop_type == Shop.ShopType.TEST and shop.test_connected_at:
                end_date = shop.test_connected_at + dt_timedelta(days=int(shop.test_period_days or 14))
                end_dt = timezone.make_aware(datetime.combine(end_date, dt_time.min), timezone.get_current_timezone())
                remaining = end_dt - now_dt
                remaining_seconds = int(remaining.total_seconds())
                if remaining_seconds <= 0:
                    timer = {"status": "expired", "label": "Тест завершено", "seconds": remaining_seconds}
                elif remaining_seconds <= 86400:
                    hours = max(1, remaining_seconds // 3600)
                    timer = {"status": "urgent", "label": f"Залишилось {hours} год", "seconds": remaining_seconds}
                else:
                    days = remaining_seconds // 86400
                    timer = {"status": "active", "label": f"Залишилось {days} дн", "seconds": remaining_seconds}

            shops_payload.append(
                {
                    "id": shop.id,
                    "name": shop.name,
                    "photo_url": shop.photo.url if shop.photo else "",
                    "owner_full_name": shop.owner_full_name,
                    "shop_type": shop.shop_type,
                    "registration_place": shop.registration_place,
                    "is_physical": bool(shop.is_physical),
                    "city": shop.city,
                    "address": shop.address,
                    "website_url": shop.website_url,
                    "instagram_url": shop.instagram_url,
                    "prom_url": shop.prom_url,
                    "other_sales_channel": shop.other_sales_channel,
                    "test_product_id": shop.test_product_id,
                    "test_package": shop.test_package or {},
                    "test_contract_name": shop.test_contract_file.name.split("/")[-1] if shop.test_contract_file else "",
                    "test_contract_download_url": reverse("management_shop_contract_download", args=[shop.id]) if shop.test_contract_file else "",
                    "test_connected_at": shop.test_connected_at.isoformat() if shop.test_connected_at else "",
                    "test_period_days": int(shop.test_period_days or 14),
                    "next_contact_at": timezone.localtime(shop.next_contact_at).isoformat() if shop.next_contact_at else "",
                    "notes": shop.notes,
                    "primary_phone": primary_phone or "",
                    "phones": phones,
                    "shipments": shipments,
                    "total_amount": str(getattr(shop, "total_amount", "") or ""),
                    "timer": timer,
                    "created_by": (shop.created_by.get_full_name() or shop.created_by.username) if shop.created_by else "",
                    "managed_by": (shop.managed_by.get_full_name() or shop.managed_by.username) if shop.managed_by else "",
                    "can_delete": True,
                }
            )

        try:
            from storefront.models import Product

            test_products = list(
                Product.objects.all()
                .select_related("category")
                .order_by("category__name", "title")
                .values("id", "title", "category__name")[:600]
            )
        except Exception:
            test_products = []

        ctx["page_obj"] = page_obj
        ctx["shops_payload"] = shops_payload
        ctx["test_products"] = test_products

    return render(request, 'management/admin.html', ctx)


@login_required(login_url='management_login')
@require_POST
def admin_invoice_approve_api(request, invoice_id):
    if not request.user.is_staff:
        return JsonResponse({'ok': False}, status=403)

    from orders.models import WholesaleInvoice

    invoice = WholesaleInvoice.objects.filter(id=invoice_id).first()
    if not invoice:
        return JsonResponse({'ok': False, 'error': 'Накладна не знайдена'}, status=404)

    if invoice.review_status != 'pending':
        return JsonResponse({'ok': False, 'error': 'Накладна вже оброблена'}, status=400)

    invoice.review_status = 'approved'
    invoice.review_reject_reason = ''
    invoice.reviewed_at = timezone.now()
    invoice.reviewed_by = request.user
    invoice.is_approved = True
    invoice.save(update_fields=['review_status', 'review_reject_reason', 'reviewed_at', 'reviewed_by', 'is_approved'])

    InvoiceRejectionReasonRequest.objects.filter(invoice=invoice, is_active=True).update(is_active=False)
    _try_update_admin_invoice_message(invoice, final=True)
    _notify_manager_invoice(
        invoice,
        title="✅ <b>Накладна підтверджена</b>",
        body_lines=[
            f"№: <code>{escape(invoice.invoice_number)}</code>",
            f"Компанія: {escape(invoice.company_name)}",
            "Статус: ✅ Підтверджено",
            "Далі: сформуйте посилання на оплату у вкладці накладних.",
        ],
    )

    return JsonResponse({'ok': True})


@login_required(login_url='management_login')
@require_POST
def admin_invoice_reject_api(request, invoice_id):
    if not request.user.is_staff:
        return JsonResponse({'ok': False}, status=403)

    import json
    from orders.models import WholesaleInvoice

    invoice = WholesaleInvoice.objects.filter(id=invoice_id).first()
    if not invoice:
        return JsonResponse({'ok': False, 'error': 'Накладна не знайдена'}, status=404)

    if invoice.review_status != 'pending':
        return JsonResponse({'ok': False, 'error': 'Накладна вже оброблена'}, status=400)

    try:
        payload = json.loads(request.body.decode('utf-8') or '{}')
    except Exception:
        payload = {}
    reason = (payload.get('reason') or '').strip()
    if not reason:
        return JsonResponse({'ok': False, 'error': 'Вкажіть причину відхилення'}, status=400)

    invoice.review_status = 'rejected'
    invoice.review_reject_reason = reason
    invoice.reviewed_at = timezone.now()
    invoice.reviewed_by = request.user
    invoice.is_approved = False
    invoice.payment_url = None
    invoice.monobank_invoice_id = None
    invoice.save(update_fields=[
        'review_status',
        'review_reject_reason',
        'reviewed_at',
        'reviewed_by',
        'is_approved',
        'payment_url',
        'monobank_invoice_id',
    ])

    InvoiceRejectionReasonRequest.objects.filter(invoice=invoice, is_active=True).update(is_active=False)
    _try_update_admin_invoice_message(invoice, final=True)
    _notify_manager_invoice(
        invoice,
        title="❌ <b>Накладна відхилена</b>",
        body_lines=[
            f"№: <code>{escape(invoice.invoice_number)}</code>",
            f"Компанія: {escape(invoice.company_name)}",
            f"Причина: {escape(reason)}",
        ],
    )

    return JsonResponse({'ok': True})


@login_required(login_url='management_login')
def admin_user_clients(request, user_id):
    """AJAX: список клієнтів обраного менеджера для модалки 'Огляд'."""
    if not request.user.is_staff:
        return JsonResponse({'ok': False}, status=403)

    User = get_user_model()
    try:
        target = User.objects.get(id=user_id)
    except User.DoesNotExist:
        return JsonResponse({'ok': False}, status=404)

    try:
        offset = int(request.GET.get('offset', '0'))
    except ValueError:
        offset = 0
    try:
        limit = int(request.GET.get('limit', '50'))
    except ValueError:
        limit = 50

    offset = max(0, offset)
    limit = max(1, min(200, limit))

    query = (request.GET.get('q') or '').strip()

    qs = Client.objects.filter(owner=target).order_by('-created_at')
    if query:
        qs = qs.filter(
            Q(shop_name__icontains=query) |
            Q(phone__icontains=query) |
            Q(full_name__icontains=query)
        )

    total = qs.count()
    page = list(qs[offset:offset + limit])

    now = timezone.localtime(timezone.now())
    today = now.date()
    yesterday = today - timedelta(days=1)

    clients = []
    for c in page:
        created_local = timezone.localtime(c.created_at) if c.created_at else now
        if created_local.date() == today:
            created_label = 'Сьогодні'
        elif created_local.date() == yesterday:
            created_label = 'Вчора'
        else:
            created_label = created_local.strftime('%d.%m.%Y')

        next_call_label = '–'
        next_call_status = 'none'
        if c.next_call_at:
            next_local = timezone.localtime(c.next_call_at)
            next_call_label = next_local.strftime('%d.%m.%Y %H:%M')
            delta = (next_local - now).total_seconds()
            if delta <= 0:
                next_call_status = 'due'
            elif delta <= 300:
                next_call_status = 'soon'
            else:
                next_call_status = 'scheduled'

        clients.append({
            'id': c.id,
            'shop': c.shop_name,
            'phone': c.phone,
            'full_name': c.full_name,
            'role_code': c.role,
            'role': c.get_role_display(),
            'source': c.source or '',
            'result_code': c.call_result,
            'result': c.get_call_result_display(),
            'details': c.call_result_details or '',
            'points': POINTS.get(c.call_result, 0),
            'next_call': next_call_label,
            'next_call_status': next_call_status,
            'created': created_local.strftime('%d.%m.%Y %H:%M'),
            'created_date_label': created_label,
        })

    has_more = offset + limit < total
    return JsonResponse({
        'ok': True,
        'total': total,
        'offset': offset,
        'limit': limit,
        'has_more': has_more,
        'clients': clients,
    })


@login_required(login_url='management_login')
def reports(request):
    if not user_is_management(request.user):
        return redirect('management_home')
    if not (request.user.is_staff or Client.objects.filter(owner=request.user).exists()):
        return render(request, 'management/reports.html', {'denied': True})

    today_start, today_end = get_today_range()
    report_sent_today = has_report_today(request.user)

    qs = Report.objects.select_related('owner').order_by('-created_at')
    if not request.user.is_staff:
        qs = qs.filter(owner=request.user)

    reports_list = []
    for r in qs:
        report_day_start = timezone.localtime(r.created_at).replace(hour=0, minute=0, second=0, microsecond=0)
        report_day_end = report_day_start + timedelta(days=1)
        clients_qs = Client.objects.filter(
            owner=r.owner,
            created_at__gte=report_day_start,
            created_at__lt=report_day_end
        ).order_by('-created_at')
        clients = list(clients_qs[:100])
        # fallback: якщо записів за дату немає, але processed > 0, показуємо останніх клієнтів
        if not clients and r.processed > 0:
            fallback_limit = max(1, min(100, r.processed))
            clients = list(Client.objects.filter(owner=r.owner).order_by('-created_at')[:fallback_limit])
        reports_list.append({
            'id': r.id,
            'created_at': r.created_at,
            'owner': r.owner,
            'points': r.points,
            'processed': r.processed,
            'file': r.file,
            'clients': [
                {
                    'shop': c.shop_name,
                    'full_name': c.full_name,
                    'phone': c.phone,
                    'role': c.get_role_display(),
                    'role_code': c.role,
                    'source': c.source,
                    'result': c.get_call_result_display(),
                    'result_code': c.call_result,
                    'details': c.call_result_details,
                    'next_call': timezone.localtime(c.next_call_at).strftime('%d.%m.%Y %H:%M') if c.next_call_at else '–',
                    'created': timezone.localtime(c.created_at).strftime('%d.%m.%Y %H:%M'),
                } for c in clients
            ]
        })

    stats = get_user_stats(request.user)
    progress_clients_pct = min(100, int(stats['processed_today'] / TARGET_CLIENTS_DAY * 100)) if TARGET_CLIENTS_DAY else 0
    progress_points_pct = min(100, int(stats['points_today'] / TARGET_POINTS_DAY * 100)) if TARGET_POINTS_DAY else 0
    bot_username = get_manager_bot_username()
    reminders = get_reminders(request.user, stats=stats, report_sent=report_sent_today)

    return render(request, 'management/reports.html', {
        'reports': reports_list,
        'user_points_today': stats['points_today'],
        'user_points_total': stats['points_total'],
        'processed_today': stats['processed_today'],
        'target_clients': TARGET_CLIENTS_DAY,
        'target_points': TARGET_POINTS_DAY,
        'progress_clients_pct': progress_clients_pct,
        'progress_points_pct': progress_points_pct,
        'has_report_today': report_sent_today,
        'reminders': reminders,
        'manager_bot_username': bot_username,
    })


@login_required(login_url='management_login')
def send_report(request):
    if request.method != 'POST':
        return redirect('management_reports')
    if not user_is_management(request.user):
        return redirect('management_home')
    if not (request.user.is_staff or Client.objects.filter(owner=request.user).exists()):
        return redirect('management_reports')

    today_start, today_end = get_today_range()
    clients_today = Client.objects.filter(
        owner=request.user,
        created_at__gte=today_start,
        created_at__lt=today_end
    ).order_by('-created_at')
    stats = get_user_stats(request.user)

    file_bytes = build_report_excel(request.user, stats, clients_today)
    filename = f"report_{request.user.username}_{timezone.localtime().strftime('%Y%m%d_%H%M')}.xlsx"

    report = Report.objects.create(
        owner=request.user,
        points=stats['points_today'],
        processed=stats['processed_today'],
    )
    report.file.save(filename, ContentFile(file_bytes), save=True)

    _close_followups_for_report(report)

    send_telegram_report(request.user, stats, clients_today, file_bytes, filename)

    return redirect('management_reports')


@login_required(login_url='management_login')
def reminder_read(request):
    if not user_is_management(request.user):
        return JsonResponse({'ok': False}, status=403)
    if request.method != 'POST':
        return JsonResponse({'ok': False}, status=400)
    try:
        body = json.loads(request.body.decode('utf-8'))
    except Exception:
        return JsonResponse({'ok': False}, status=400)
    key = body.get('key')
    if not key:
        return JsonResponse({'ok': False}, status=400)
    ReminderRead.objects.get_or_create(user=request.user, key=key)
    return JsonResponse({'ok': True})


@login_required(login_url='management_login')
def reminder_feed(request):
    if not user_is_management(request.user):
        return JsonResponse({'reminders': []})
    stats = get_user_stats(request.user)
    report_sent = has_report_today(request.user)
    reminders = get_reminders(request.user, stats=stats, report_sent=report_sent)
    _send_manager_bot_notifications(request.user, reminders)
    serialized = []
    for r in reminders:
        eta = int(r.get('eta_seconds', 0) or 0)
        if eta > 300:
            eta = 0
        serialized.append({
            'shop': r.get('shop', ''),
            'title': r.get('title', ''),
            'name': r.get('name', ''),
            'phone': r.get('phone', ''),
            'time_label': r.get('time_label', ''),
            'status': r.get('status', ''),
            'eta_seconds': eta,
            'key': r.get('key', ''),
            'read': r.get('read', False),
            'dt_iso': r.get('dt_iso', ''),
        })
    return JsonResponse({'reminders': serialized})


@login_required(login_url='management_login')
@require_POST
def profile_update(request):
    if not user_is_management(request.user):
        return JsonResponse({'ok': False}, status=403)

    import re
    from datetime import date
    from urllib.parse import urlparse

    profile, _ = UserProfile.objects.get_or_create(user=request.user)

    errors = {}

    def set_trimmed(attr: str, raw, *, max_len: int | None = None):
        if raw is None:
            return
        val = (raw or '').strip()
        if max_len is not None and len(val) > max_len:
            errors[attr] = 'Занадто довге значення.'
            return
        setattr(profile, attr, val)

    full_name = request.POST.get('full_name')
    if full_name is not None:
        set_trimmed('full_name', full_name, max_len=200)

    email = request.POST.get('email')
    if email is not None:
        email = (email or '').strip()
        if email:
            try:
                validate_email(email)
            except ValidationError:
                errors['email'] = 'Некоректний email.'
            else:
                profile.email = email
        else:
            profile.email = ''

    phone = request.POST.get('phone')
    if phone is not None:
        phone = (phone or '').strip()
        if len(phone) > 32:
            errors['phone'] = 'Занадто довгий телефон.'
        else:
            profile.phone = phone

    city = request.POST.get('city')
    if city is not None:
        set_trimmed('city', city, max_len=100)

    def normalize_instagram(value: str) -> str:
        raw = (value or '').strip()
        if not raw:
            return ''
        if raw.startswith('@'):
            raw = raw[1:].strip()
        if 'instagram.com' in raw:
            try:
                parsed = urlparse(raw if raw.startswith('http') else f'https://{raw.lstrip("/")}')
                path = (parsed.path or '').strip('/')
                if path:
                    raw = path.split('/', 1)[0]
            except Exception:
                pass
        raw = raw.strip().lstrip('@')
        if not raw:
            return ''
        if len(raw) > 30:
            raise ValueError('Instagram: занадто довгий логін.')
        if not re.fullmatch(r'[A-Za-z0-9._]+', raw):
            raise ValueError('Instagram: дозволені лише літери, цифри, крапка та підкреслення.')
        return raw

    instagram = request.POST.get('instagram')
    if instagram is not None:
        instagram = (instagram or '').strip()
        if instagram:
            try:
                profile.instagram = normalize_instagram(instagram)
            except ValueError as exc:
                errors['instagram'] = str(exc)
        else:
            profile.instagram = ''

    def normalize_messenger(value: str) -> str:
        v = (value or '').strip()
        if not v:
            return ''
        if len(v) > 100:
            raise ValueError('Занадто довге значення.')
        if not re.fullmatch(r'[0-9+()\-\s@._]+', v):
            raise ValueError('Недопустимі символи.')
        return v

    whatsapp_has = bool(request.POST.get('whatsapp_has'))
    whatsapp = request.POST.get('whatsapp', '')
    if not whatsapp_has:
        profile.whatsapp = ''
    else:
        try:
            profile.whatsapp = normalize_messenger(whatsapp)
        except ValueError as exc:
            errors['whatsapp'] = str(exc)

    viber_has = bool(request.POST.get('viber_has'))
    viber = request.POST.get('viber', '')
    if not viber_has:
        profile.viber = ''
    else:
        try:
            profile.viber = normalize_messenger(viber)
        except ValueError as exc:
            errors['viber'] = str(exc)

    day_raw = (request.POST.get('birth_day') or '').strip()
    month_raw = (request.POST.get('birth_month') or '').strip()
    year_raw = (request.POST.get('birth_year') or '').strip()

    if day_raw or month_raw or year_raw:
        if not (day_raw and month_raw and year_raw):
            errors['birth_day'] = 'Вкажіть дату повністю.'
            errors['birth_month'] = 'Вкажіть дату повністю.'
            errors['birth_year'] = 'Вкажіть дату повністю.'
        else:
            try:
                profile.birth_date = date(int(year_raw), int(month_raw), int(day_raw))
            except Exception:
                errors['birth_day'] = 'Некоректна дата.'
                errors['birth_month'] = 'Некоректна дата.'
                errors['birth_year'] = 'Некоректна дата.'
    else:
        profile.birth_date = None

    payment_card = request.POST.get('payment_card')
    if payment_card is not None:
        card_raw = (payment_card or '').strip()
        if card_raw:
            digits = re.sub(r'\D+', '', card_raw)

            def luhn_ok(num: str) -> bool:
                if not num or not num.isdigit():
                    return False
                total = 0
                rev = num[::-1]
                for i, ch in enumerate(rev):
                    d = int(ch)
                    if i % 2 == 1:
                        d *= 2
                        if d > 9:
                            d -= 9
                    total += d
                return total % 10 == 0

            if len(digits) < 12 or len(digits) > 19:
                errors['payment_card'] = 'Номер картки виглядає некоректно.'
            elif not luhn_ok(digits):
                errors['payment_card'] = 'Номер картки виглядає некоректно.'
            else:
                profile.payment_method = 'card'
                profile.payment_details = ' '.join(digits[i:i+4] for i in range(0, len(digits), 4))

    if errors:
        return JsonResponse({'ok': False, 'error': 'Перевірте поля форми.', 'errors': errors}, status=400)

    profile.save()
    return JsonResponse({'ok': True})


@login_required(login_url='management_login')
@require_POST
def profile_bind_code(request):
    if not user_is_management(request.user):
        return JsonResponse({'ok': False}, status=403)
    profile, _ = UserProfile.objects.get_or_create(user=request.user)
    code = secrets.token_hex(3)
    expires_at = timezone.now() + timedelta(minutes=10)
    profile.tg_manager_bind_code = code
    profile.tg_manager_bind_expires_at = expires_at
    profile.save()
    bot_username = get_manager_bot_username()
    return JsonResponse({
        'ok': True,
        'code': code,
        'expires': expires_at.strftime('%d.%m.%Y %H:%M'),
        'bot_username': bot_username,
    })


def _send_telegram_message(bot_token, chat_id, text):
    if not bot_token or not chat_id:
        return
    url = f"https://api.telegram.org/bot{bot_token}/sendMessage"
    try:
        requests.post(url, data={'chat_id': chat_id, 'text': text}, timeout=5)
    except Exception:
        pass


def _tg_api_post(bot_token, method, payload, *, as_json=True, timeout=10):
    if not bot_token:
        return None
    url = f"https://api.telegram.org/bot{bot_token}/{method}"
    try:
        if as_json:
            resp = requests.post(url, json=payload, timeout=timeout)
        else:
            resp = requests.post(url, data=payload, timeout=timeout)
        return resp.json()
    except Exception:
        return None


def _tg_send_message(bot_token, chat_id, text, *, reply_markup=None, parse_mode='HTML'):
    if not bot_token or not chat_id:
        return None
    payload = {
        'chat_id': chat_id,
        'text': text,
        'parse_mode': parse_mode,
        'disable_web_page_preview': True,
    }
    if reply_markup is not None:
        payload['reply_markup'] = reply_markup
    data = _tg_api_post(bot_token, 'sendMessage', payload, as_json=True)
    if data and data.get('ok'):
        return data.get('result')
    return None


def _tg_send_document(
    bot_token,
    chat_id,
    *,
    file_path=None,
    file_bytes=None,
    filename=None,
    mime_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    caption='',
    reply_markup=None,
    parse_mode='HTML',
    timeout=20,
):
    if not bot_token or not chat_id:
        return None

    try:
        chat_id_int = int(chat_id)
    except Exception:
        chat_id_int = chat_id

    url = f"https://api.telegram.org/bot{bot_token}/sendDocument"
    data = {
        'chat_id': chat_id_int,
        'caption': caption or '',
        'parse_mode': parse_mode,
        'disable_web_page_preview': True,
    }
    if reply_markup is not None:
        try:
            data['reply_markup'] = json.dumps(reply_markup, ensure_ascii=False)
        except Exception:
            data['reply_markup'] = json.dumps(reply_markup)

    if file_path:
        try:
            with open(file_path, 'rb') as f:
                send_name = filename or os.path.basename(file_path) or 'invoice.xlsx'
                files = {'document': (send_name, f, mime_type)}
                resp = requests.post(url, data=data, files=files, timeout=timeout)
        except Exception:
            return None
    elif file_bytes is not None:
        send_name = filename or 'invoice.xlsx'
        files = {'document': (send_name, file_bytes, mime_type)}
        try:
            resp = requests.post(url, data=data, files=files, timeout=timeout)
        except Exception:
            return None
    else:
        return None

    try:
        payload = resp.json()
    except Exception:
        return None
    if payload and payload.get('ok'):
        return payload.get('result')
    return None


def _tg_edit_message(bot_token, chat_id, message_id, text, *, reply_markup=None, parse_mode='HTML'):
    if not bot_token or not chat_id or not message_id:
        return None
    payload = {
        'chat_id': chat_id,
        'message_id': message_id,
        'text': text,
        'parse_mode': parse_mode,
        'disable_web_page_preview': True,
    }
    if reply_markup is not None:
        payload['reply_markup'] = reply_markup
    data = _tg_api_post(bot_token, 'editMessageText', payload, as_json=True)
    if data and data.get('ok'):
        return data.get('result')
    return None


def _tg_edit_caption(bot_token, chat_id, message_id, caption, *, reply_markup=None, parse_mode='HTML'):
    if not bot_token or not chat_id or not message_id:
        return None
    payload = {
        'chat_id': chat_id,
        'message_id': message_id,
        'caption': caption,
        'parse_mode': parse_mode,
    }
    if reply_markup is not None:
        payload['reply_markup'] = reply_markup
    data = _tg_api_post(bot_token, 'editMessageCaption', payload, as_json=True)
    if data and data.get('ok'):
        return data.get('result')
    return None


def _tg_answer_callback(bot_token, callback_query_id, text=''):
    if not bot_token or not callback_query_id:
        return None
    payload = {'callback_query_id': callback_query_id}
    if text:
        payload['text'] = text
        payload['show_alert'] = False
    return _tg_api_post(bot_token, 'answerCallbackQuery', payload, as_json=True)


def _get_management_admin_bot_config():
    token = os.environ.get("MANAGEMENT_TG_BOT_TOKEN") or os.environ.get("MANAGER_TG_BOT_TOKEN")
    chat_ids = _get_management_admin_chat_ids()
    return token, (chat_ids[0] if chat_ids else None)


def _get_manager_bot_token():
    return os.environ.get("MANAGER_TG_BOT_TOKEN") or os.environ.get("MANAGEMENT_TG_BOT_TOKEN")


def _format_admin_invoice_message(invoice, *, status_line=None, include_links=True, include_excel_link=True):
    manager_name = ""
    if getattr(invoice, 'created_by', None):
        manager_name = invoice.created_by.get_full_name() or invoice.created_by.username
    company_name = escape(invoice.company_name or '')
    company_number = escape(invoice.company_number or '—')
    phone = escape(invoice.contact_phone or '—')
    address = escape(invoice.delivery_address or '—')
    amount = escape(str(invoice.total_amount))
    tshirts = invoice.total_tshirts or 0
    hoodies = invoice.total_hoodies or 0

    lines = [
        "🧾 <b>Накладна на перевірку</b>",
        "",
        f"<b>№</b>: <code>{escape(invoice.invoice_number)}</code>",
        f"<b>Менеджер</b>: {escape(manager_name) if manager_name else '—'}",
        f"<b>Компанія</b>: {company_name}",
        f"<b>ЄДРПОУ/ІПН</b>: {company_number}",
        f"<b>Телефон</b>: {phone}",
        f"<b>Адреса</b>: {address}",
        f"<b>Сума</b>: {amount} грн",
        f"<b>Футболки</b>: {tshirts} • <b>Худі</b>: {hoodies}",
    ]
    if status_line:
        lines += ["", status_line]

    if include_links:
        lines += [
            "",
            f"🌐 Адмін-панель: <a href=\"https://management.twocomms.shop/admin-panel/?tab=invoices\">відкрити</a>",
        ]
        if include_excel_link:
            lines += [
                f"📥 Excel: <a href=\"https://management.twocomms.shop/invoices/{invoice.id}/download/\">скачати</a>",
            ]

    return "\n".join(lines)


def _notify_manager_invoice(invoice, *, title, body_lines):
    try:
        manager = invoice.created_by
        profile = manager.userprofile
    except Exception:
        return
    chat_id = getattr(profile, 'tg_manager_chat_id', None)
    bot_token = _get_manager_bot_token()
    if not bot_token or not chat_id:
        return
    text = "\n".join([title, *body_lines])
    _tg_send_message(bot_token, chat_id, text, parse_mode='HTML')


def _normalize_invoice_admin_message_ref(chat_id, message_id):
    try:
        normalized_chat_id = int(chat_id)
        normalized_message_id = int(message_id)
    except (TypeError, ValueError):
        return None
    if not normalized_chat_id or not normalized_message_id:
        return None
    return {
        "chat_id": normalized_chat_id,
        "message_id": normalized_message_id,
    }


def _get_invoice_admin_message_refs(invoice):
    refs = []
    seen = set()

    primary_ref = _normalize_invoice_admin_message_ref(
        getattr(invoice, "admin_tg_chat_id", None),
        getattr(invoice, "admin_tg_message_id", None),
    )
    if primary_ref:
        refs.append(primary_ref)
        seen.add((primary_ref["chat_id"], primary_ref["message_id"]))

    order_details = invoice.order_details if isinstance(invoice.order_details, dict) else {}
    raw_refs = order_details.get("admin_tg_messages") or []
    if isinstance(raw_refs, list):
        for raw_ref in raw_refs:
            if not isinstance(raw_ref, dict):
                continue
            normalized_ref = _normalize_invoice_admin_message_ref(
                raw_ref.get("chat_id"),
                raw_ref.get("message_id"),
            )
            if not normalized_ref:
                continue
            ref_key = (normalized_ref["chat_id"], normalized_ref["message_id"])
            if ref_key in seen:
                continue
            refs.append(normalized_ref)
            seen.add(ref_key)

    return refs


def _save_invoice_admin_message_refs(invoice, refs):
    normalized_refs = []
    seen = set()
    for ref in refs:
        normalized_ref = None
        if isinstance(ref, dict):
            normalized_ref = _normalize_invoice_admin_message_ref(
                ref.get("chat_id"),
                ref.get("message_id"),
            )
        elif isinstance(ref, (list, tuple)) and len(ref) >= 2:
            normalized_ref = _normalize_invoice_admin_message_ref(ref[0], ref[1])
        if not normalized_ref:
            continue
        ref_key = (normalized_ref["chat_id"], normalized_ref["message_id"])
        if ref_key in seen:
            continue
        normalized_refs.append(normalized_ref)
        seen.add(ref_key)

    order_details = dict(invoice.order_details) if isinstance(invoice.order_details, dict) else {}
    if normalized_refs:
        order_details["admin_tg_messages"] = normalized_refs
        invoice.admin_tg_chat_id = normalized_refs[0]["chat_id"]
        invoice.admin_tg_message_id = normalized_refs[0]["message_id"]
    else:
        order_details.pop("admin_tg_messages", None)
        invoice.admin_tg_chat_id = None
        invoice.admin_tg_message_id = None

    invoice.order_details = order_details
    invoice.save(update_fields=["order_details", "admin_tg_chat_id", "admin_tg_message_id"])
    return normalized_refs


def _remember_invoice_admin_message_ref(invoice, chat_id, message_id):
    normalized_ref = _normalize_invoice_admin_message_ref(chat_id, message_id)
    if not normalized_ref:
        return None

    existing_refs = _get_invoice_admin_message_refs(invoice)
    ref_key = (normalized_ref["chat_id"], normalized_ref["message_id"])
    existing_keys = {(ref["chat_id"], ref["message_id"]) for ref in existing_refs}
    if ref_key not in existing_keys:
        existing_refs.append(normalized_ref)

    return _save_invoice_admin_message_refs(invoice, existing_refs)


def _try_update_admin_invoice_message(invoice, *, bot_token=None, final=False):
    refs = _get_invoice_admin_message_refs(invoice)
    if not refs:
        return
    token = bot_token or (os.environ.get("MANAGEMENT_TG_BOT_TOKEN") or os.environ.get("MANAGER_TG_BOT_TOKEN"))
    if not token:
        return
    status_line = None
    if invoice.review_status == 'approved':
        status_line = "✅ <b>ПІДТВЕРДЖЕНО</b>"
    elif invoice.review_status == 'rejected':
        reason = escape((invoice.review_reject_reason or '').strip() or '—')
        status_line = f"❌ <b>ВІДХИЛЕНО</b>\n<b>Причина</b>: {reason}"
    elif invoice.review_status == 'pending':
        status_line = "⏳ <b>НА ПЕРЕВІРЦІ</b>"
    else:
        status_line = None

    text = _format_admin_invoice_message(invoice, status_line=status_line, include_links=True, include_excel_link=False)
    reply_markup = {'inline_keyboard': []} if final else None
    fallback_text = _format_admin_invoice_message(invoice, status_line=status_line, include_links=True, include_excel_link=True)
    for ref in refs:
        updated = _tg_edit_caption(
            token,
            ref["chat_id"],
            ref["message_id"],
            text,
            reply_markup=reply_markup,
            parse_mode='HTML',
        )
        if not updated:
            _tg_edit_message(
                token,
                ref["chat_id"],
                ref["message_id"],
                fallback_text,
                reply_markup=reply_markup,
                parse_mode='HTML',
            )


def _send_invoice_review_request_to_admin(invoice, *, request=None):
    token = os.environ.get("MANAGEMENT_TG_BOT_TOKEN") or os.environ.get("MANAGER_TG_BOT_TOKEN")
    raw_chat_ids = _get_management_admin_chat_ids()
    if not token or not raw_chat_ids:
        return

    chat_ids = []
    seen_chat_ids = set()
    for raw_chat_id in raw_chat_ids:
        try:
            chat_id_int = int(raw_chat_id)
        except (TypeError, ValueError):
            continue
        if chat_id_int in seen_chat_ids:
            continue
        chat_ids.append(chat_id_int)
        seen_chat_ids.add(chat_id_int)

    if not chat_ids:
        return

    caption = _format_admin_invoice_message(invoice, status_line="Оберіть дію нижче ⬇️", include_links=True, include_excel_link=False)
    keyboard = {
        'inline_keyboard': [[
            {'text': '✅ Підтвердити', 'callback_data': f'inv:approve:{invoice.id}'},
            {'text': '❌ Відхилити', 'callback_data': f'inv:reject:{invoice.id}'},
        ]]
    }
    fallback_text = _format_admin_invoice_message(invoice, status_line="Оберіть дію нижче ⬇️", include_links=True, include_excel_link=True)
    sent_refs = []

    for chat_id_int in chat_ids:
        sent = None
        try:
            if invoice.file_path and os.path.exists(invoice.file_path):
                sent = _tg_send_document(
                    token,
                    chat_id_int,
                    file_path=invoice.file_path,
                    filename=os.path.basename(invoice.file_path),
                    caption=caption,
                    reply_markup=keyboard,
                    parse_mode='HTML',
                )
        except Exception:
            sent = None
        if not sent:
            sent = _tg_send_message(token, chat_id_int, fallback_text, reply_markup=keyboard, parse_mode='HTML')
        if sent and sent.get('message_id'):
            sent_refs.append((chat_id_int, sent.get('message_id')))

    if sent_refs:
        _save_invoice_admin_message_refs(invoice, sent_refs)


def _format_admin_contract_message(contract, *, status_line=None, include_links=True):
    manager_name = ""
    if getattr(contract, 'created_by', None):
        manager_name = contract.created_by.get_full_name() or contract.created_by.username
    payload = contract.payload or {}
    realizer = payload.get("realizer_name") or contract.realizer_name or ""
    recipient = payload.get("recipient_name") or realizer or ""
    recipient_phone = payload.get("recipient_phone") or payload.get("realizer_phone") or ""
    delivery_address = payload.get("delivery_address") or ""
    product_title = payload.get("product_title") or payload.get("product_print") or ""
    total_sum = payload.get("total_sum")
    total_sum_text = str(total_sum) if total_sum not in (None, "") else "—"
    contract_date = contract.contract_date.strftime('%d.%m.%Y') if contract.contract_date else ""

    lines = [
        "📄 <b>Договір на перевірку</b>",
        "",
        f"<b>№</b>: <code>{escape(contract.contract_number)}</code>",
        f"<b>Дата</b>: {escape(contract_date) if contract_date else '—'}",
        f"<b>Менеджер</b>: {escape(manager_name) if manager_name else '—'}",
        f"<b>Реалізатор</b>: {escape(realizer) if realizer else '—'}",
        f"<b>Одержувач</b>: {escape(recipient) if recipient else '—'}",
        f"<b>Телефон</b>: {escape(recipient_phone) if recipient_phone else '—'}",
        f"<b>Адреса доставки</b>: {escape(delivery_address) if delivery_address else '—'}",
        f"<b>Товар</b>: {escape(product_title) if product_title else '—'}",
        f"<b>Сума партії</b>: {escape(total_sum_text)} грн",
    ]
    if status_line:
        lines += ["", status_line]

    if include_links:
        lines += [
            "",
            "🌐 Адмін-панель: <a href=\"https://management.twocomms.shop/admin-panel/?tab=contracts\">відкрити</a>",
            f"📥 DOCX: <a href=\"https://management.twocomms.shop/contracts/{contract.id}/download/\">скачати</a>",
        ]

    return "\n".join(lines)


def _notify_manager_contract(contract, *, title, body_lines):
    try:
        manager = contract.created_by
        profile = manager.userprofile
    except Exception:
        return
    chat_id = getattr(profile, 'tg_manager_chat_id', None)
    bot_token = _get_manager_bot_token()
    if not bot_token or not chat_id:
        return
    text = "\n".join([title, *body_lines])
    _tg_send_message(bot_token, chat_id, text, parse_mode='HTML')


def _try_update_admin_contract_message(contract, *, bot_token=None, final=False):
    if not contract.admin_tg_chat_id or not contract.admin_tg_message_id:
        return
    token = bot_token or (os.environ.get("MANAGEMENT_TG_BOT_TOKEN") or os.environ.get("MANAGER_TG_BOT_TOKEN"))
    if not token:
        return
    status_line = None
    if contract.review_status == 'approved':
        status_line = "✅ <b>ПІДТВЕРДЖЕНО</b>"
    elif contract.review_status == 'rejected':
        reason = escape((contract.review_reject_reason or '').strip() or '—')
        status_line = f"❌ <b>ВІДХИЛЕНО</b>\n<b>Причина</b>: {reason}"
    elif contract.review_status == 'pending':
        status_line = "⏳ <b>НА ПЕРЕВІРЦІ</b>"

    text = _format_admin_contract_message(contract, status_line=status_line, include_links=True)
    reply_markup = {'inline_keyboard': []} if final else None
    updated = _tg_edit_caption(token, contract.admin_tg_chat_id, contract.admin_tg_message_id, text, reply_markup=reply_markup, parse_mode='HTML')
    if not updated:
        _tg_edit_message(token, contract.admin_tg_chat_id, contract.admin_tg_message_id, text, reply_markup=reply_markup, parse_mode='HTML')


def _send_contract_review_request_to_admin(contract, *, request=None):
    token, chat_id = _get_management_admin_bot_config()
    if not token or not chat_id:
        return
    try:
        chat_id_int = int(chat_id)
    except Exception:
        return

    caption = _format_admin_contract_message(contract, status_line="Оберіть дію нижче ⬇️", include_links=True)
    keyboard = {
        'inline_keyboard': [[
            {'text': '✅ Підтвердити', 'callback_data': f'ctr:approve:{contract.id}'},
            {'text': '❌ Відхилити', 'callback_data': f'ctr:reject:{contract.id}'},
        ]]
    }
    sent = None
    try:
        if contract.file_path and os.path.exists(contract.file_path):
            sent = _tg_send_document(
                token,
                chat_id_int,
                file_path=contract.file_path,
                filename=os.path.basename(contract.file_path),
                caption=caption,
                reply_markup=keyboard,
                parse_mode='HTML',
            )
    except Exception:
        sent = None
    if not sent:
        sent = _tg_send_message(token, chat_id_int, caption, reply_markup=keyboard, parse_mode='HTML')
    if sent and sent.get('message_id'):
        contract.admin_tg_chat_id = chat_id_int
        contract.admin_tg_message_id = sent.get('message_id')
        contract.save(update_fields=['admin_tg_chat_id', 'admin_tg_message_id'])


def _format_admin_payout_message(payout_req, *, status_line=None, include_links=True):
    import re

    manager_name = ''
    try:
        if getattr(payout_req, 'owner', None):
            manager_name = payout_req.owner.get_full_name() or payout_req.owner.username
    except Exception:
        manager_name = ''

    def mask_card(raw):
        s = (raw or '').strip()
        digits = re.sub(r'\D+', '', s)
        if len(digits) < 4:
            return '—'
        return '**** ' + digits[-4:]

    card_raw = ''
    try:
        prof = payout_req.owner.userprofile
        card_raw = getattr(prof, 'payment_details', '')
    except Exception:
        card_raw = ''

    created_label = ''
    try:
        created_label = timezone.localtime(payout_req.created_at).strftime('%d.%m.%Y %H:%M')
    except Exception:
        created_label = ''

    lines = [
        '💸 <b>Запит на виплату</b>',
        '',
        f"<b>ID</b>: <code>{escape(str(getattr(payout_req, 'id', '')))}</code>",
        f"<b>Менеджер</b>: {escape(manager_name) if manager_name else '—'}",
        f"<b>Сума</b>: {escape(str(getattr(payout_req, 'amount', '0')))} грн",
        f"<b>Картка</b>: {escape(mask_card(card_raw))}",
    ]
    if created_label:
        lines.append(f"<b>Створено</b>: {escape(created_label)}")

    if status_line:
        lines += ['', status_line]

    if include_links:
        lines += [
            '',
            '🌐 Адмін-панель: <a href="https://management.twocomms.shop/admin-panel/?tab=payouts">відкрити</a>',
        ]

    return "\n".join(lines)


def _admin_payout_keyboard(payout_req):
    try:
        from management.models import ManagerPayoutRequest
    except Exception:
        return {'inline_keyboard': []}

    status = getattr(payout_req, 'status', '')
    if status == ManagerPayoutRequest.Status.PROCESSING:
        return {
            'inline_keyboard': [[
                {'text': '✅ Схвалити', 'callback_data': f'pay:approve:{payout_req.id}'},
                {'text': '❌ Відхилити', 'callback_data': f'pay:reject:{payout_req.id}'},
            ]]
        }

    if status == ManagerPayoutRequest.Status.APPROVED:
        return {
            'inline_keyboard': [[
                {'text': '💳 Виплачено', 'callback_data': f'pay:paid:{payout_req.id}'},
                {'text': '❌ Відхилити', 'callback_data': f'pay:reject:{payout_req.id}'},
            ]]
        }

    return {'inline_keyboard': []}


def _notify_manager_payout(payout_req, *, title, body_lines):
    try:
        manager = payout_req.owner
        profile = manager.userprofile
    except Exception:
        return

    chat_id = getattr(profile, 'tg_manager_chat_id', None)
    bot_token = _get_manager_bot_token()
    if not bot_token or not chat_id:
        return

    text = "\n".join([title, *body_lines])
    _tg_send_message(bot_token, chat_id, text, parse_mode='HTML')


def _try_update_admin_payout_message(payout_req, *, bot_token=None, final=False):
    if not getattr(payout_req, 'admin_tg_chat_id', None) or not getattr(payout_req, 'admin_tg_message_id', None):
        return

    token = bot_token or (os.environ.get('MANAGEMENT_TG_BOT_TOKEN') or os.environ.get('MANAGER_TG_BOT_TOKEN'))
    if not token:
        return

    status = getattr(payout_req, 'status', '')
    status_line = None
    try:
        from management.models import ManagerPayoutRequest
        if status == ManagerPayoutRequest.Status.PROCESSING:
            status_line = '⏳ <b>В ОБРОБЦІ</b>'
        elif status == ManagerPayoutRequest.Status.APPROVED:
            status_line = '✅ <b>СХВАЛЕНО</b>'
        elif status == ManagerPayoutRequest.Status.REJECTED:
            reason = escape((getattr(payout_req, 'rejection_reason', '') or '').strip() or '—')
            status_line = f"❌ <b>ВІДХИЛЕНО</b>\n<b>Причина</b>: {reason}"
        elif status == ManagerPayoutRequest.Status.PAID:
            status_line = '💳 <b>ВИПЛАЧЕНО</b>'
    except Exception:
        status_line = None

    text = _format_admin_payout_message(payout_req, status_line=status_line, include_links=True)
    reply_markup = {'inline_keyboard': []} if final else _admin_payout_keyboard(payout_req)
    _tg_edit_message(token, payout_req.admin_tg_chat_id, payout_req.admin_tg_message_id, text, reply_markup=reply_markup, parse_mode='HTML')


def _send_payout_request_to_admin(payout_req):
    token, chat_id = _get_management_admin_bot_config()
    if not token or not chat_id:
        return

    try:
        chat_id_int = int(chat_id)
    except Exception:
        return

    text = _format_admin_payout_message(payout_req, status_line='Оберіть дію нижче ⬇️', include_links=True)
    keyboard = _admin_payout_keyboard(payout_req)
    sent = _tg_send_message(token, chat_id_int, text, reply_markup=keyboard, parse_mode='HTML')
    if sent and sent.get('message_id'):
        payout_req.admin_tg_chat_id = chat_id_int
        payout_req.admin_tg_message_id = sent.get('message_id')
        payout_req.save(update_fields=['admin_tg_chat_id', 'admin_tg_message_id'])


def _send_manager_bot_notifications(user, reminders):
    """Send new reminders (<=5 мин) to manager bot."""
    if not reminders:
        return
    try:
        profile = user.userprofile
    except UserProfile.DoesNotExist:
        return
    chat_id = profile.tg_manager_chat_id
    bot_token = os.environ.get("MANAGER_TG_BOT_TOKEN") or os.environ.get("MANAGEMENT_TG_BOT_TOKEN")
    if not chat_id or not bot_token:
        return
    for r in reminders:
        eta = int(r.get('eta_seconds') or 0)
        status = r.get('status')
        if eta == 0 and status not in ('due', 'report'):
            continue
        if eta > 300:
            continue
        key = r.get('key')
        if not key:
            continue
        if ReminderSent.objects.filter(key=key, chat_id=chat_id).exists():
            continue
        when = 'зараз' if eta == 0 else (f"через {eta//60} хв {eta % 60:02d} с" if eta >= 60 else f"через {eta} с")
        kind = r.get('kind')
        contact_label = "Контакт" if kind == 'shop' else "Клієнт"
        text = (
            "Нагадування\n"
            f"Магазин: {r.get('shop', '')}\n"
            f"{contact_label}: {r.get('name', '')}\n"
            f"Телефон: {r.get('phone', '')}\n"
            f"Коли: {when}\n"
        )
        _send_telegram_message(bot_token, chat_id, text)
        ReminderSent.objects.create(key=key, chat_id=chat_id)


@csrf_exempt
def management_bot_webhook(request, token):
    manager_token = os.environ.get("MANAGER_TG_BOT_TOKEN")
    admin_token = os.environ.get("MANAGEMENT_TG_BOT_TOKEN")
    valid_tokens = {t for t in (manager_token, admin_token) if t}
    if not valid_tokens or token not in valid_tokens:
        return JsonResponse({'ok': False}, status=403)
    bot_token = token
    try:
        payload = json.loads(request.body.decode('utf-8'))
    except Exception:
        return JsonResponse({'ok': False}, status=400)

    # ==================== CALLBACK QUERIES (ADMIN INVOICE REVIEW) ====================
    callback = payload.get('callback_query')
    if callback:
        cb_id = callback.get('id')
        data = (callback.get('data') or '').strip()
        msg = callback.get('message') or {}
        msg_chat = msg.get('chat') or {}
        chat_id = msg_chat.get('id')
        message_id = msg.get('message_id')

        admin_chat_cfg = _get_management_admin_chat_ids()
        if admin_chat_cfg and str(chat_id) not in {str(v) for v in admin_chat_cfg}:
            _tg_answer_callback(bot_token, cb_id, "Недостатньо прав")
            return JsonResponse({'ok': True})

        # ==================== CALLBACK QUERIES (ADMIN PAYOUTS) ====================
        if data.startswith('pay:'):
            parts = data.split(':', 2)
            if len(parts) != 3:
                _tg_answer_callback(bot_token, cb_id, 'Невідома дія')
                return JsonResponse({'ok': True})

            action = parts[1]
            try:
                payout_id = int(parts[2])
            except Exception:
                _tg_answer_callback(bot_token, cb_id, 'Невірний ID')
                return JsonResponse({'ok': True})

            from django.db import transaction
            from management.models import ManagerPayoutRequest, PayoutRejectionReasonRequest

            req = None
            prompt_req = None

            with transaction.atomic():
                req = ManagerPayoutRequest.objects.select_for_update().select_related('owner', 'owner__userprofile').filter(id=payout_id).first()
                if not req:
                    _tg_answer_callback(bot_token, cb_id, 'Запит не знайдено')
                    return JsonResponse({'ok': True})

                # Persist message reference for later sync (site actions)
                if chat_id and message_id and (not req.admin_tg_chat_id or not req.admin_tg_message_id):
                    req.admin_tg_chat_id = chat_id
                    req.admin_tg_message_id = message_id
                    req.save(update_fields=['admin_tg_chat_id', 'admin_tg_message_id'])

                if action == 'approve':
                    if req.status == ManagerPayoutRequest.Status.PAID:
                        _tg_answer_callback(bot_token, cb_id, 'Вже виплачено')
                        return JsonResponse({'ok': True})
                    if req.status == ManagerPayoutRequest.Status.REJECTED:
                        _tg_answer_callback(bot_token, cb_id, 'Вже відхилено')
                        return JsonResponse({'ok': True})
                    if req.status == ManagerPayoutRequest.Status.APPROVED:
                        _tg_answer_callback(bot_token, cb_id, 'Вже схвалено')
                        return JsonResponse({'ok': True})

                    req.status = ManagerPayoutRequest.Status.APPROVED
                    req.approved_at = timezone.now()
                    req.rejected_at = None
                    req.rejection_reason = ''
                    req.processed_by = None
                    req.save(update_fields=['status', 'approved_at', 'rejected_at', 'rejection_reason', 'processed_by'])
                    PayoutRejectionReasonRequest.objects.filter(payout_request=req, is_active=True).update(is_active=False)

                elif action == 'reject':
                    if req.status in (ManagerPayoutRequest.Status.PAID, ManagerPayoutRequest.Status.REJECTED):
                        _tg_answer_callback(bot_token, cb_id, 'Вже оброблено')
                        return JsonResponse({'ok': True})

                    PayoutRejectionReasonRequest.objects.filter(admin_chat_id=chat_id, is_active=True).update(is_active=False)
                    prompt_req = PayoutRejectionReasonRequest.objects.create(payout_request=req, admin_chat_id=chat_id, is_active=True)

                elif action == 'paid':
                    if req.status == ManagerPayoutRequest.Status.PAID:
                        _tg_answer_callback(bot_token, cb_id, 'Вже виплачено')
                        return JsonResponse({'ok': True})
                    if req.status == ManagerPayoutRequest.Status.REJECTED:
                        _tg_answer_callback(bot_token, cb_id, 'Запит відхилено')
                        return JsonResponse({'ok': True})
                    if req.status != ManagerPayoutRequest.Status.APPROVED:
                        _tg_answer_callback(bot_token, cb_id, 'Спочатку схваліть')
                        return JsonResponse({'ok': True})

                    req.status = ManagerPayoutRequest.Status.PAID
                    req.paid_at = timezone.now()
                    req.processed_by = None
                    req.save(update_fields=['status', 'paid_at', 'processed_by'])
                    PayoutRejectionReasonRequest.objects.filter(payout_request=req, is_active=True).update(is_active=False)

                else:
                    _tg_answer_callback(bot_token, cb_id, 'Невідома дія')
                    return JsonResponse({'ok': True})

            # Outside transaction: Telegram edits / notifications
            if action == 'approve' and req:
                _try_update_admin_payout_message(req, bot_token=bot_token, final=False)

                import re
                card_raw = ''
                try:
                    card_raw = getattr(req.owner.userprofile, 'payment_details', '')
                except Exception:
                    card_raw = ''
                digits = re.sub(r'\D+', '', (card_raw or ''))
                card_mask = ('**** ' + digits[-4:]) if len(digits) >= 4 else '—'

                _notify_manager_payout(
                    req,
                    title='✅ <b>Виплату схвалено</b>',
                    body_lines=[
                        f"Сума: <b>{escape(str(req.amount))} грн</b>.",
                        f"Протягом 3 годин сума буде зарахована на картку <code>{escape(card_mask)}</code>.",
                    ],
                )

                _tg_answer_callback(bot_token, cb_id, 'Схвалено')
                return JsonResponse({'ok': True})

            if action == 'reject' and req and prompt_req:
                waiting_text = _format_admin_payout_message(req, status_line='✍️ <b>Ожидаю причину отклонения</b>', include_links=True)
                _tg_edit_message(bot_token, chat_id, message_id, waiting_text, reply_markup={'inline_keyboard': []}, parse_mode='HTML')

                prompt = _tg_send_message(
                    bot_token,
                    chat_id,
                    f"❌ <b>Відхилення виплати</b> <code>{escape(str(req.id))}</code>\n\nНапишіть причину одним повідомленням.",
                    reply_markup={'force_reply': True, 'input_field_placeholder': 'Причина…'},
                    parse_mode='HTML',
                )
                if prompt and prompt.get('message_id'):
                    prompt_req.prompt_message_id = prompt.get('message_id')
                    prompt_req.save(update_fields=['prompt_message_id'])

                _tg_answer_callback(bot_token, cb_id, 'Вкажіть причину')
                return JsonResponse({'ok': True})

            if action == 'paid' and req:
                _try_update_admin_payout_message(req, bot_token=bot_token, final=True)

                import re
                card_raw = ''
                try:
                    card_raw = getattr(req.owner.userprofile, 'payment_details', '')
                except Exception:
                    card_raw = ''
                digits = re.sub(r'\D+', '', (card_raw or ''))
                card_mask = ('**** ' + digits[-4:]) if len(digits) >= 4 else '—'

                _notify_manager_payout(
                    req,
                    title='💳 <b>Виплату здійснено</b>',
                    body_lines=[
                        f"Сума: <b>{escape(str(req.amount))} грн</b>.",
                        f"Зачислено на карту <code>{escape(card_mask)}</code>.",
                    ],
                )

                _tg_answer_callback(bot_token, cb_id, 'Готово')
                return JsonResponse({'ok': True})

            _tg_answer_callback(bot_token, cb_id, 'Готово')
            return JsonResponse({'ok': True})

        if data.startswith('ctr:'):
            parts = data.split(':', 2)
            if len(parts) != 3:
                _tg_answer_callback(bot_token, cb_id, "Невідома дія")
                return JsonResponse({'ok': True})

            action = parts[1]
            try:
                contract_id = int(parts[2])
            except Exception:
                _tg_answer_callback(bot_token, cb_id, "Невірний ID")
                return JsonResponse({'ok': True})

            contract = ManagementContract.objects.filter(id=contract_id).select_related('created_by', 'created_by__userprofile').first()
            if not contract:
                _tg_answer_callback(bot_token, cb_id, "Договір не знайдено")
                return JsonResponse({'ok': True})

            if chat_id and message_id and (not contract.admin_tg_chat_id or not contract.admin_tg_message_id):
                contract.admin_tg_chat_id = chat_id
                contract.admin_tg_message_id = message_id
                contract.save(update_fields=['admin_tg_chat_id', 'admin_tg_message_id'])

            if contract.review_status != 'pending':
                _tg_answer_callback(bot_token, cb_id, "Вже оброблено")
                _try_update_admin_contract_message(contract, bot_token=bot_token, final=True)
                return JsonResponse({'ok': True})

            if action == 'approve':
                contract.review_status = 'approved'
                contract.review_reject_reason = ''
                contract.reviewed_at = timezone.now()
                contract.reviewed_by = None
                contract.is_approved = True
                contract.save(update_fields=['review_status', 'review_reject_reason', 'reviewed_at', 'reviewed_by', 'is_approved'])
                ContractRejectionReasonRequest.objects.filter(contract=contract, is_active=True).update(is_active=False)
                _try_update_admin_contract_message(contract, bot_token=bot_token, final=True)
                _notify_manager_contract(
                    contract,
                    title="✅ <b>Договір підтверджено</b>",
                    body_lines=[
                        f"№: <code>{escape(contract.contract_number)}</code>",
                        f"Реалізатор: {escape(contract.realizer_name or '—')}",
                        "Статус: ✅ Підтверджено",
                    ],
                )
                _tg_answer_callback(bot_token, cb_id, "Підтверджено")
                return JsonResponse({'ok': True})

            if action == 'reject':
                ContractRejectionReasonRequest.objects.filter(admin_chat_id=chat_id, is_active=True).update(is_active=False)
                req = ContractRejectionReasonRequest.objects.create(contract=contract, admin_chat_id=chat_id, is_active=True)

                waiting_text = _format_admin_contract_message(contract, status_line="✍️ <b>Очікую причину відхилення</b>", include_links=True)
                updated = _tg_edit_caption(bot_token, chat_id, message_id, waiting_text, reply_markup={'inline_keyboard': []}, parse_mode='HTML')
                if not updated:
                    _tg_edit_message(bot_token, chat_id, message_id, waiting_text, reply_markup={'inline_keyboard': []}, parse_mode='HTML')

                prompt = _tg_send_message(
                    bot_token,
                    chat_id,
                    f"❌ <b>Відхилення договору</b> <code>{escape(contract.contract_number)}</code>\n\nНапишіть причину одним повідомленням.",
                    reply_markup={'force_reply': True, 'input_field_placeholder': 'Причина відхилення…'},
                    parse_mode='HTML',
                )
                if prompt and prompt.get('message_id'):
                    req.prompt_message_id = prompt.get('message_id')
                    req.save(update_fields=['prompt_message_id'])

                _tg_answer_callback(bot_token, cb_id, "Вкажіть причину")
                return JsonResponse({'ok': True})

            _tg_answer_callback(bot_token, cb_id, "Невідома дія")
            return JsonResponse({'ok': True})

        if not data.startswith('inv:') or ':' not in data:
            _tg_answer_callback(bot_token, cb_id, "Невідома дія")
            return JsonResponse({'ok': True})

        parts = data.split(':', 2)
        if len(parts) != 3:
            _tg_answer_callback(bot_token, cb_id, "Невідома дія")
            return JsonResponse({'ok': True})

        action = parts[1]
        try:
            invoice_id = int(parts[2])
        except Exception:
            _tg_answer_callback(bot_token, cb_id, "Невірний ID")
            return JsonResponse({'ok': True})

        from orders.models import WholesaleInvoice
        invoice = WholesaleInvoice.objects.filter(id=invoice_id).select_related('created_by', 'created_by__userprofile').first()
        if not invoice:
            _tg_answer_callback(bot_token, cb_id, "Накладна не знайдена")
            return JsonResponse({'ok': True})

        # Persist every admin message reference so status sync updates all copies.
        if chat_id and message_id:
            _remember_invoice_admin_message_ref(invoice, chat_id, message_id)

        if invoice.review_status != 'pending':
            _tg_answer_callback(bot_token, cb_id, "Вже оброблено")
            _try_update_admin_invoice_message(invoice, bot_token=bot_token, final=True)
            return JsonResponse({'ok': True})

        if action == 'approve':
            invoice.review_status = 'approved'
            invoice.review_reject_reason = ''
            invoice.reviewed_at = timezone.now()
            invoice.reviewed_by = None
            invoice.is_approved = True
            invoice.save(update_fields=['review_status', 'review_reject_reason', 'reviewed_at', 'reviewed_by', 'is_approved'])
            InvoiceRejectionReasonRequest.objects.filter(invoice=invoice, is_active=True).update(is_active=False)
            _try_update_admin_invoice_message(invoice, bot_token=bot_token, final=True)
            _notify_manager_invoice(
                invoice,
                title="✅ <b>Накладна підтверджена</b>",
                body_lines=[
                    f"№: <code>{escape(invoice.invoice_number)}</code>",
                    f"Компанія: {escape(invoice.company_name)}",
                    "Статус: ✅ Підтверджено",
                    "Далі: сформуйте посилання на оплату у вкладці накладних.",
                ],
            )
            _tg_answer_callback(bot_token, cb_id, "Підтверджено")
            return JsonResponse({'ok': True})

        if action == 'reject':
            InvoiceRejectionReasonRequest.objects.filter(admin_chat_id=chat_id, is_active=True).update(is_active=False)
            req = InvoiceRejectionReasonRequest.objects.create(invoice=invoice, admin_chat_id=chat_id, is_active=True)

            # Update original message: remove buttons and mark waiting reason
            waiting_text = _format_admin_invoice_message(invoice, status_line="✍️ <b>Очікую причину відхилення</b>", include_links=True, include_excel_link=False)
            updated = _tg_edit_caption(bot_token, chat_id, message_id, waiting_text, reply_markup={'inline_keyboard': []}, parse_mode='HTML')
            if not updated:
                _tg_edit_message(bot_token, chat_id, message_id, waiting_text, reply_markup={'inline_keyboard': []}, parse_mode='HTML')

            prompt = _tg_send_message(
                bot_token,
                chat_id,
                f"❌ <b>Відхилення накладної</b> <code>{escape(invoice.invoice_number)}</code>\n\nНапишіть причину одним повідомленням.",
                reply_markup={'force_reply': True, 'input_field_placeholder': 'Причина відхилення…'},
                parse_mode='HTML',
            )
            if prompt and prompt.get('message_id'):
                req.prompt_message_id = prompt.get('message_id')
                req.save(update_fields=['prompt_message_id'])

            _tg_answer_callback(bot_token, cb_id, "Вкажіть причину")
            return JsonResponse({'ok': True})

        _tg_answer_callback(bot_token, cb_id, "Невідома дія")
        return JsonResponse({'ok': True})

    # ==================== MESSAGES ====================
    message = payload.get('message') or {}
    text = (message.get('text') or '').strip()
    chat = message.get('chat') or {}
    from_user = message.get('from') or {}
    chat_id = chat.get('id')
    username = from_user.get('username') or ''

    # Admin: rejection reason flow
    if chat_id and text and not text.startswith('/'):
        active_req = InvoiceRejectionReasonRequest.objects.select_related('invoice').filter(
            admin_chat_id=chat_id,
            is_active=True
        ).first()
        if active_req:
            invoice = active_req.invoice
            active_req.is_active = False
            active_req.save(update_fields=['is_active'])

            if invoice and invoice.review_status == 'pending':
                reason = text.strip()
                invoice.review_status = 'rejected'
                invoice.review_reject_reason = reason
                invoice.reviewed_at = timezone.now()
                invoice.reviewed_by = None
                invoice.is_approved = False
                invoice.payment_url = None
                invoice.monobank_invoice_id = None
                invoice.save(update_fields=[
                    'review_status',
                    'review_reject_reason',
                    'reviewed_at',
                    'reviewed_by',
                    'is_approved',
                    'payment_url',
                    'monobank_invoice_id',
                ])
                _try_update_admin_invoice_message(invoice, bot_token=bot_token, final=True)
                _notify_manager_invoice(
                    invoice,
                    title="❌ <b>Накладна відхилена</b>",
                    body_lines=[
                        f"№: <code>{escape(invoice.invoice_number)}</code>",
                        f"Компанія: {escape(invoice.company_name)}",
                        f"Причина: {escape(reason)}",
                    ],
                )
                _tg_send_message(bot_token, chat_id, "Готово ✅ Накладну відхилено.", parse_mode='HTML')
            else:
                _tg_send_message(bot_token, chat_id, "Накладна вже оброблена або недоступна.", parse_mode='HTML')
            return JsonResponse({'ok': True})

        contract_req = ContractRejectionReasonRequest.objects.select_related('contract').filter(
            admin_chat_id=chat_id,
            is_active=True,
        ).first()
        if contract_req:
            contract = contract_req.contract
            contract_req.is_active = False
            contract_req.save(update_fields=['is_active'])

            if contract and contract.review_status == 'pending':
                reason = text.strip()
                contract.review_status = 'rejected'
                contract.review_reject_reason = reason
                contract.reviewed_at = timezone.now()
                contract.reviewed_by = None
                contract.is_approved = False
                contract.save(update_fields=[
                    'review_status',
                    'review_reject_reason',
                    'reviewed_at',
                    'reviewed_by',
                    'is_approved',
                ])
                _try_update_admin_contract_message(contract, bot_token=bot_token, final=True)
                _notify_manager_contract(
                    contract,
                    title="❌ <b>Договір відхилено</b>",
                    body_lines=[
                        f"№: <code>{escape(contract.contract_number)}</code>",
                        f"Реалізатор: {escape(contract.realizer_name or '—')}",
                        f"Причина: {escape(reason)}",
                    ],
                )
                _tg_send_message(bot_token, chat_id, "Готово ✅ Договір відхилено.", parse_mode='HTML')
            else:
                _tg_send_message(bot_token, chat_id, "Договір вже оброблено або недоступний.", parse_mode='HTML')
            return JsonResponse({'ok': True})

        # Payouts: rejection reason flow
        try:
            from management.models import ManagerPayoutRequest, PayoutRejectionReasonRequest
        except Exception:
            ManagerPayoutRequest = None
            PayoutRejectionReasonRequest = None

        if ManagerPayoutRequest and PayoutRejectionReasonRequest:
            pay_req = PayoutRejectionReasonRequest.objects.select_related(
                'payout_request',
                'payout_request__owner',
                'payout_request__owner__userprofile',
            ).filter(
                admin_chat_id=chat_id,
                is_active=True,
            ).first()

            if pay_req:
                req = pay_req.payout_request
                pay_req.is_active = False
                pay_req.save(update_fields=['is_active'])

                if req and req.status not in (ManagerPayoutRequest.Status.REJECTED, ManagerPayoutRequest.Status.PAID):
                    reason = text.strip()
                    req.status = ManagerPayoutRequest.Status.REJECTED
                    req.rejection_reason = reason
                    req.rejected_at = timezone.now()
                    req.approved_at = None
                    req.paid_at = None
                    req.processed_by = None
                    req.save(update_fields=['status', 'rejection_reason', 'rejected_at', 'approved_at', 'paid_at', 'processed_by'])

                    _try_update_admin_payout_message(req, bot_token=bot_token, final=True)
                    _notify_manager_payout(
                        req,
                        title='❌ <b>Виплату відхилено</b>',
                        body_lines=[
                            f"Сума: <b>{escape(str(req.amount))} грн</b>.",
                            f"Причина: {escape(reason)}",
                        ],
                    )
                    _tg_send_message(bot_token, chat_id, 'Готово ✅ Запит відхилено.', parse_mode='HTML')
                else:
                    _tg_send_message(bot_token, chat_id, 'Запит вже оброблено або недоступний.', parse_mode='HTML')

                return JsonResponse({'ok': True})
    # Manager bot binding
    if text.startswith('/start'):
        code = ''
        if ' ' in text:
            code = text.split(' ', 1)[1].strip()
        now = timezone.now()
        if code:
            profile = UserProfile.objects.filter(
                tg_manager_bind_code=code,
                tg_manager_bind_expires_at__gte=now
            ).first()
            if profile:
                profile.tg_manager_chat_id = chat_id
                profile.tg_manager_username = username
                profile.tg_manager_bind_code = ''
                profile.tg_manager_bind_expires_at = None
                profile.save()
                _send_telegram_message(bot_token, chat_id, "Готово! Бот привʼязаний. Нагадування будуть тут.")
            else:
                _send_telegram_message(bot_token, chat_id, "Код не дійсний або прострочений. Згенеруйте новий у кабінеті і натисніть 'Привʼязати бота'.")
        else:
            _send_telegram_message(bot_token, chat_id, "Не отримав код. Згенеруйте його у кабінеті й натисніть 'Привʼязати бота' або надішліть /start <код>.")
    return JsonResponse({'ok': True})


def _get_wholesale_products_for_invoice():
    from django.db.models import Q
    from storefront.models import Product, Category

    tshirt_categories = Category.objects.filter(
        Q(name__icontains='футболка') |
        Q(name__icontains='tshirt') |
        Q(name__icontains='t-shirt') |
        Q(slug__icontains='футболка') |
        Q(slug__icontains='tshirt') |
        Q(slug__icontains='t-shirt')
    )
    hoodie_categories = Category.objects.filter(
        Q(name__icontains='худи') |
        Q(name__icontains='hoodie') |
        Q(name__icontains='hooded') |
        Q(slug__icontains='худи') |
        Q(slug__icontains='hoodie') |
        Q(slug__icontains='hooded')
    )

    tshirt_products = Product.objects.filter(category__in=tshirt_categories).select_related('category').prefetch_related(
        'color_variants__color',
        'color_variants__images',
    )
    hoodie_products = Product.objects.filter(category__in=hoodie_categories).select_related('category').prefetch_related(
        'color_variants__color',
        'color_variants__images',
    )
    return tshirt_products, hoodie_products


def _get_wholesale_price_context():
    price_context = get_management_wholesale_price_context()
    tshirt_prices = price_context["tshirt"]
    hoodie_prices = price_context["hoodie"]
    return tshirt_prices, hoodie_prices


_UA_MONTHS = {
    1: "січня",
    2: "лютого",
    3: "березня",
    4: "квітня",
    5: "травня",
    6: "червня",
    7: "липня",
    8: "серпня",
    9: "вересня",
    10: "жовтня",
    11: "листопада",
    12: "грудня",
}


def _contract_template_path():
    base_path = Path(settings.BASE_DIR) / "static"
    primary = base_path / "Договір Рудь.docx"
    fallback = base_path / "Договір.docx"
    return primary if primary.exists() else fallback


def _ua_month_name(month_num):
    return _UA_MONTHS.get(month_num, "")


def _format_ua_date(date_obj):
    return f"«{date_obj.day}» {_ua_month_name(date_obj.month)} {date_obj.year}р."


def _format_contract_number(seq_num, date_obj):
    return f"{seq_num:02d}\\{date_obj.day:02d}{date_obj.month:02d}\\{str(date_obj.year)[-2:]}"


def _format_price_uah(value):
    return f"{int(value)},00"


def _ua_number_to_words(value):
    value = int(value)
    if value == 0:
        return "нуль"
    units_m = ["", "один", "два", "три", "чотири", "п’ять", "шість", "сім", "вісім", "дев’ять"]
    units_f = ["", "одна", "дві", "три", "чотири", "п’ять", "шість", "сім", "вісім", "дев’ять"]
    teens = [
        "десять",
        "одинадцять",
        "дванадцять",
        "тринадцять",
        "чотирнадцять",
        "п’ятнадцять",
        "шістнадцять",
        "сімнадцять",
        "вісімнадцять",
        "дев’ятнадцять",
    ]
    tens = ["", "", "двадцять", "тридцять", "сорок", "п’ятдесят", "шістдесят", "сімдесят", "вісімдесят", "дев’яносто"]
    hundreds = [
        "",
        "сто",
        "двісті",
        "триста",
        "чотириста",
        "п’ятсот",
        "шістсот",
        "сімсот",
        "вісімсот",
        "дев’ятсот",
    ]

    parts = []
    thousands = value // 1000
    if thousands:
        parts.append(units_f[thousands])
        if thousands % 10 == 1 and thousands % 100 != 11:
            parts.append("тисяча")
        elif 2 <= thousands % 10 <= 4 and not (12 <= thousands % 100 <= 14):
            parts.append("тисячі")
        else:
            parts.append("тисяч")
    value = value % 1000
    if value // 100:
        parts.append(hundreds[value // 100])
    value = value % 100
    if 10 <= value <= 19:
        parts.append(teens[value - 10])
    else:
        if value // 10:
            parts.append(tens[value // 10])
        if value % 10:
            parts.append(units_m[value % 10])
    return " ".join([p for p in parts if p])


def _replace_placeholders(text, replacements):
    for key, value in replacements.items():
        text = text.replace(key, value)
    return text


def _replace_text_in_paragraph(paragraph, search_text, replace_text):
    if search_text is None or search_text == "":
        return False
    replace_text = "" if replace_text is None else str(replace_text)
    if search_text == replace_text:
        return True
    if search_text not in paragraph.text:
        return False
    replaced = False
    max_iter = paragraph.text.count(search_text)
    while search_text in paragraph.text and max_iter > 0:
        full_text = "".join(run.text for run in paragraph.runs)
        start_idx = full_text.find(search_text)
        if start_idx < 0:
            break
        end_idx = start_idx + len(search_text)
        offset = 0
        for run in paragraph.runs:
            run_text = run.text
            run_len = len(run_text)
            run_start = offset
            run_end = offset + run_len
            if run_end <= start_idx or run_start >= end_idx:
                offset = run_end
                continue
            local_start = max(start_idx, run_start) - run_start
            local_end = min(end_idx, run_end) - run_start
            if run_start <= start_idx < run_end:
                run.text = run_text[:local_start] + replace_text + run_text[local_end:]
            else:
                run.text = run_text[:local_start] + run_text[local_end:]
            offset = run_end
        replaced = True
        max_iter -= 1
    return replaced


def _replace_text_in_cell(cell, search_text, replace_text):
    replaced = False
    for paragraph in cell.paragraphs:
        if _replace_text_in_paragraph(paragraph, search_text, replace_text):
            replaced = True
    return replaced


def _replace_text_in_doc(doc, search_text, replace_text):
    replaced = False
    for paragraph in doc.paragraphs:
        if _replace_text_in_paragraph(paragraph, search_text, replace_text):
            replaced = True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if _replace_text_in_cell(cell, search_text, replace_text):
                    replaced = True
    return replaced


def _safe_contract_filename(realizer_name, contract_number):
    safe_name = re.sub(r"[\\\\/:*?\"<>|]+", "", realizer_name or "").strip()
    safe_name = re.sub(r"\s+", "_", safe_name) or "Реалізатор"
    safe_number = re.sub(r"[\\\\/:*?\"<>|]+", "-", contract_number)
    return f"Договір_{safe_name}_{safe_number}.docx"


def _guess_product_type(product):
    slug = ""
    title = ""
    try:
        slug = (product.category.slug or "").lower() if getattr(product, "category", None) else ""
    except Exception:
        slug = ""
    try:
        title = (product.title or "").lower()
    except Exception:
        title = ""
    hoodie_slugs = {"hoodie", "hoodies"}
    tshirt_slugs = {"tshirts", "tshirt", "t-shirt", "tee"}
    if slug in hoodie_slugs or "худи" in title or "hoodie" in title or "флис" in title:
        return "hoodie"
    if slug in tshirt_slugs or "футболка" in title or "t-shirt" in title or "tshirt" in title:
        return "tshirt"
    return ""


def _update_contract_paragraphs(doc, data):
    type_forms = data["type_forms"]
    in_realizer_block = False
    in_delivery_block = False
    placeholder_map = {
        "{{contract_number}}": data["contract_number"],
        "{{contract_date}}": data["contract_date_text"],
        "{{realizer_name}}": data["realizer_name"],
        "{{realizer_code}}": data["realizer_code"],
        "{{realizer_address}}": data["realizer_address"],
        "{{realizer_iban}}": data["realizer_iban"],
        "{{realizer_phone}}": data["realizer_phone"],
        "{{realizer_email}}": data["realizer_email"],
        "{{delivery_address}}": data["delivery_address"],
        "{{recipient_name}}": data["recipient_name"],
        "{{recipient_phone}}": data["recipient_phone"],
        "{{product_type_single}}": type_forms["single"],
        "{{product_type_plural}}": type_forms["plural"],
        "{{product_type_gen}}": type_forms["gen"],
        "{{product_print}}": data["product_print"],
        "{{price_str}}": data["price_str"],
        "{{price_words}}": data["price_words"],
        "{{total_sum}}": str(data["total_sum"]),
        "{{product_table_name}}": data["product_table_name"],
        "{{delivery_role}}": "отримувача" if data["recipient_diff"] else "Реалізатора",
    }
    for paragraph in doc.paragraphs:
        raw_text = paragraph.text or ""
        text = raw_text.strip()
        if not text:
            continue
        if "{{" in raw_text and "}}" in raw_text:
            for key, value in placeholder_map.items():
                if key in paragraph.text:
                    _replace_text_in_paragraph(paragraph, key, value)
            continue
        if text.startswith("№ "):
            match = re.search(r"№\s*(.+)", text)
            if match:
                _replace_text_in_paragraph(paragraph, match.group(1), data["contract_number"])
            continue
        if text.startswith("м.") and "«" in text and "р." in text:
            match = re.search(r"«\d{1,2}»\s+[^\s]+\s+\d{4}р\.", text)
            if match:
                _replace_text_in_paragraph(paragraph, match.group(0), data["contract_date_text"])
            continue
        if text.startswith("Фізична особа – підприємець") and "Реалізатор" in text:
            name_match = re.search(r"підприємець\s+(.+?),\s+код\s+([0-9A-Za-z]+)", text)
            if name_match:
                _replace_text_in_paragraph(paragraph, name_match.group(1), data["realizer_name"])
                _replace_text_in_paragraph(paragraph, name_match.group(2), data["realizer_code"])
            continue
        if text.startswith("1.1. Постачальник передає"):
            match = re.search(r"товару –\s+(.+?)\s+\(далі", text)
            if match:
                _replace_text_in_paragraph(paragraph, match.group(1), type_forms["gen"])
            continue
        if text.startswith("1.2.1. Орієнтовно:"):
            match = re.search(r"Орієнтовно:\s+(.+?)\s+з", text)
            if match:
                _replace_text_in_paragraph(paragraph, match.group(1), type_forms["plural"])
            print_match = re.search(r'\["(.+?)"\]', text)
            if print_match:
                _replace_text_in_paragraph(paragraph, print_match.group(1), data["product_print"])
            continue
        if text.startswith("1.2.2. Базова ціна"):
            price_match = re.search(r"(\d+[\d\s]*,\d{2})", text)
            if price_match:
                _replace_text_in_paragraph(paragraph, price_match.group(1), data["price_str"])
            words_match = re.search(r"\(([^)]+)\)", text)
            if words_match:
                _replace_text_in_paragraph(paragraph, words_match.group(1), data["price_words"])
            continue
        if text.startswith("3.1. Товар відправляється") and data["recipient_diff"]:
            _replace_text_in_paragraph(paragraph, "ПІБ Реалізатора", "ПІБ отримувача")
            continue
        if text.startswith("РЕАЛІЗАТОР:"):
            in_realizer_block = True
            continue
        if text.startswith("Додаток №1"):
            in_realizer_block = False
        if in_realizer_block:
            if text.startswith("ФОП:"):
                value = text.split("ФОП:", 1)[1].strip()
                if value:
                    _replace_text_in_paragraph(paragraph, value, data["realizer_name"])
                continue
            if text.startswith("РНОКПП/ЄДРПОУ:"):
                value = text.split("РНОКПП/ЄДРПОУ:", 1)[1].strip()
                if value:
                    _replace_text_in_paragraph(paragraph, value, data["realizer_code"])
                continue
            if text.startswith("Адреса:"):
                value = text.split("Адреса:", 1)[1].strip()
                if value:
                    _replace_text_in_paragraph(paragraph, value, data["realizer_address"])
                continue
            if text.startswith("IBAN:"):
                value = text.split("IBAN:", 1)[1].strip()
                if value:
                    _replace_text_in_paragraph(paragraph, value, data["realizer_iban"])
                continue
            if text.startswith("Телефон:"):
                value = text.split("Телефон:", 1)[1].strip()
                if value:
                    _replace_text_in_paragraph(paragraph, value, data["realizer_phone"])
                continue
            if text.startswith("E-mail:"):
                value = text.split("E-mail:", 1)[1].strip()
                if value:
                    _replace_text_in_paragraph(paragraph, value, data["realizer_email"])
                continue
        if text.startswith("1. Найменування Товару:"):
            match = re.search(r"Товару:\s+(.+?)\s+з", text)
            if match:
                _replace_text_in_paragraph(paragraph, match.group(1), type_forms["single"])
            print_match = re.search(r"«(.+?)»", text)
            if print_match:
                _replace_text_in_paragraph(paragraph, print_match.group(1), data["product_print"])
            continue
        if text.startswith("3. Загальна вартість партії:"):
            match = re.search(r"партії:\s*(\d+)", text)
            if match:
                _replace_text_in_paragraph(paragraph, match.group(1), str(data["total_sum"]))
            continue
        if text.startswith("4. Адреса доставки"):
            in_delivery_block = True
            if data["recipient_diff"]:
                _replace_text_in_paragraph(paragraph, "Реалізатора", "отримувача")
            continue
        if text.startswith("Адреса / відділення доставки:"):
            value = text.split("Адреса / відділення доставки:", 1)[1].strip()
            if value:
                _replace_text_in_paragraph(paragraph, value, data["delivery_address"])
            continue
        if text.startswith("ПІБ отримувача:"):
            value = text.split("ПІБ отримувача:", 1)[1].strip()
            if value:
                _replace_text_in_paragraph(paragraph, value, data["recipient_name"])
            continue
        if in_delivery_block and text.startswith("Телефон:"):
            value = text.split("Телефон:", 1)[1].strip()
            if value:
                _replace_text_in_paragraph(paragraph, value, data["recipient_phone"])
            continue


def _update_contract_table(doc, data):
    def _set_cell_text(cell, value):
        value_str = str(value)
        if not cell.paragraphs:
            cell.text = value_str
            return
        first = cell.paragraphs[0]
        if first.text:
            _replace_text_in_paragraph(first, first.text, value_str)
        else:
            if first.runs:
                first.runs[0].text = value_str
                for run in first.runs[1:]:
                    run.text = ""
            else:
                first.add_run(value_str)
        for extra in cell.paragraphs[1:]:
            if extra.text:
                if not _replace_text_in_paragraph(extra, extra.text, ""):
                    extra.text = ""

    target_table = None
    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = [cell.text for cell in table.rows[0].cells]
        if any("Найменування" in cell for cell in header_cells) and any("Сума" in cell for cell in header_cells):
            target_table = table
            break
    if not target_table:
        return
    sizes = data["sizes"]
    while len(target_table.rows) < len(sizes) + 1:
        target_table.add_row()
    for idx, row_data in enumerate(sizes, start=1):
        row = target_table.rows[idx]
        _set_cell_text(row.cells[0], idx)
        _set_cell_text(row.cells[1], data["product_table_name"])
        _set_cell_text(row.cells[2], row_data["size"])
        _set_cell_text(row.cells[3], row_data["qty"])
        _set_cell_text(row.cells[4], data["price_str"])
        _set_cell_text(row.cells[5], row_data["total_str"])


def _build_contract_docx(template_path, output_target, data):
    doc = Document(str(template_path))
    _update_contract_paragraphs(doc, data)
    _update_contract_table(doc, data)
    doc.save(output_target)


def _prepare_contract_data(payload, request, *, preview=False):
    missing = []
    placeholder_long = "_" * 30
    placeholder_short = "_" * 18
    placeholder_phone = "_" * 30
    placeholder_iban = "_" * 30

    def _value(field, placeholder):
        val = (payload.get(field) or "").strip()
        if val:
            return val
        missing.append(field)
        return placeholder if preview else ""

    realizer_name = _value("realizer_name", placeholder_long)
    realizer_code = _value("realizer_code", placeholder_short)
    realizer_address = _value("realizer_address", placeholder_long)
    realizer_iban = _value("realizer_iban", placeholder_iban)
    realizer_phone = _value("realizer_phone", placeholder_phone)
    realizer_email = _value("realizer_email", placeholder_long)
    delivery_address = _value("delivery_address", placeholder_long)
    recipient_name = _value("recipient_name", placeholder_long)
    recipient_phone = _value("recipient_phone", placeholder_phone)

    product_type = (payload.get("product_type") or "").strip()
    if product_type not in ("hoodie", "tshirt"):
        if preview:
            if "product_type" not in missing:
                missing.append("product_type")
            product_type = "hoodie"
        else:
            raise ContractPayloadError("invalid_product_type")

    product_print_raw = (payload.get("product_print") or "").strip()
    product_title = (payload.get("product_title") or "").strip()
    product_id = payload.get("product_id")
    if not product_print_raw:
        missing.append("product_print")

    required_fields = [
        "realizer_name",
        "realizer_code",
        "realizer_address",
        "realizer_iban",
        "realizer_phone",
        "realizer_email",
        "delivery_address",
        "recipient_name",
        "recipient_phone",
    ]

    from storefront.models import Product

    product = None
    if product_id:
        try:
            product = Product.objects.filter(id=int(product_id)).first()
        except (TypeError, ValueError):
            product = None
    detected_type = _guess_product_type(product) if product else ""
    if product and detected_type and detected_type != product_type:
        product = None

    price = None
    if product:
        price = int(product.get_drop_price(None) or 0)
    if not price:
        default_tee = DROP_FIXED_TEE_PRICE
        default_hoodie = DROP_FIXED_HOODIE_PRICE
        try:
            settings_obj = CommercialOfferEmailSettings.objects.filter(owner=request.user).first()
            if settings_obj:
                default_tee = int(settings_obj.drop_tee_price or default_tee)
                default_hoodie = int(settings_obj.drop_hoodie_price or default_hoodie)
        except Exception:
            default_tee = DROP_FIXED_TEE_PRICE
            default_hoodie = DROP_FIXED_HOODIE_PRICE
        price = int(default_hoodie if product_type == "hoodie" else default_tee)

    price_str = _format_price_uah(price)
    price_words = _ua_number_to_words(price)
    today = timezone.localdate()

    type_forms = {
        "single": "худі" if product_type == "hoodie" else "футболка",
        "plural": "худі" if product_type == "hoodie" else "футболки",
        "gen": "худі" if product_type == "hoodie" else "футболки",
        "title": "Худі" if product_type == "hoodie" else "Футболка",
    }
    if "product_type" in missing:
        type_forms = {
            "single": placeholder_short,
            "plural": placeholder_short,
            "gen": placeholder_short,
            "title": "Товар",
        }

    if not product_title:
        if product and product.title:
            product_title = product.title
        elif product_print_raw:
            product_title = f"{type_forms['title']} \"{product_print_raw}\""

    def _extract_print_from_title(title):
        match = re.search(r'[«"](.+?)[»"]', title)
        return match.group(1).strip() if match else ""

    product_print = product_print_raw
    if not product_print and product_title:
        extracted = _extract_print_from_title(product_title)
        product_print = extracted or product_title
        if "product_print" in missing:
            missing.remove("product_print")

    if not product_print:
        if preview:
            product_print = placeholder_long

    if not product_title:
        product_title = f"{type_forms['title']} \"{product_print}\""

    if not preview:
        missing_required = [field for field in required_fields if field in set(missing)]
        if not product_print:
            missing_required.append("product_print")
        if missing_required:
            raise ContractPayloadError("missing_fields", missing_required)

    quantities = payload.get('quantities') or {}
    sizes_order = ["S", "M", "L", "XL"]
    sizes = []
    total_sum = 0
    for size in sizes_order:
        qty = quantities.get(size, 1)
        try:
            qty = max(int(qty), 0)
        except (TypeError, ValueError):
            qty = 0
        line_total = qty * price
        total_sum += line_total
        sizes.append({
            "size": size,
            "qty": qty,
            "total": line_total,
            "total_str": _format_price_uah(line_total),
        })

    if preview:
        seq = ContractSequence.objects.filter(year=today.year).first()
        next_number = (seq.last_number + 1) if seq else 1
        seq_num = next_number
    else:
        with transaction.atomic():
            seq, _ = ContractSequence.objects.select_for_update().get_or_create(
                year=today.year,
                defaults={"last_number": 0},
            )
            seq.last_number = int(seq.last_number) + 1
            seq.save(update_fields=["last_number"])
            seq_num = seq.last_number

    contract_number = _format_contract_number(seq_num, today)
    contract_date_text = _format_ua_date(today)
    recipient_diff = recipient_name != realizer_name or recipient_phone != realizer_phone

    doc_data = {
        "contract_number": contract_number,
        "contract_date_text": contract_date_text,
        "realizer_name": realizer_name,
        "realizer_code": realizer_code,
        "realizer_address": realizer_address,
        "realizer_iban": realizer_iban,
        "realizer_phone": realizer_phone,
        "realizer_email": realizer_email,
        "delivery_address": delivery_address,
        "recipient_name": recipient_name,
        "recipient_phone": recipient_phone,
        "product_print": product_print,
        "product_table_name": product_title,
        "price_str": price_str,
        "price_words": price_words,
        "total_sum": total_sum,
        "sizes": sizes,
        "type_forms": type_forms,
        "recipient_diff": recipient_diff,
    }
    meta = {
        "contract_number": contract_number,
        "today": today,
        "price": price,
        "price_str": price_str,
        "total_sum": total_sum,
        "product_title": product_title,
    }
    return doc_data, meta


@login_required(login_url='management_login')
def contracts(request):
    if not user_is_management(request.user):
        return redirect('management_login')

    from storefront.models import Product

    today = timezone.localdate()
    contract_date_display = _format_ua_date(today)
    seq = ContractSequence.objects.filter(year=today.year).first()
    next_number = (seq.last_number + 1) if seq else 1
    next_contract_number = _format_contract_number(next_number, today)

    hoodie_products = []
    tshirt_products = []
    try:
        hoodie_slugs = {"hoodie", "hoodies"}
        tshirt_slugs = {"tshirts", "tshirt", "t-shirt", "tee"}
        products_qs = Product.objects.filter(
            status='published',
            is_dropship_available=True,
        ).select_related("category")
        for product in products_qs:
            prod_type = _guess_product_type(product)
            if not prod_type and product.category_id:
                slug = (product.category.slug or "").lower()
                if slug in hoodie_slugs:
                    prod_type = "hoodie"
                elif slug in tshirt_slugs:
                    prod_type = "tshirt"
            if not prod_type:
                continue
            item = {
                "id": product.id,
                "title": product.title,
                "drop_price": int(product.get_drop_price(None) or 0),
            }
            if prod_type == "hoodie":
                hoodie_products.append(item)
            elif prod_type == "tshirt":
                tshirt_products.append(item)
    except Exception:
        hoodie_products = []
        tshirt_products = []

    drop_tee_price = DROP_FIXED_TEE_PRICE
    drop_hoodie_price = DROP_FIXED_HOODIE_PRICE
    try:
        settings_obj = CommercialOfferEmailSettings.objects.filter(owner=request.user).first()
        if settings_obj:
            drop_tee_price = int(settings_obj.drop_tee_price or drop_tee_price)
            drop_hoodie_price = int(settings_obj.drop_hoodie_price or drop_hoodie_price)
    except Exception:
        drop_tee_price = DROP_FIXED_TEE_PRICE
        drop_hoodie_price = DROP_FIXED_HOODIE_PRICE

    last_contract = ManagementContract.objects.filter(created_by=request.user).first()
    prefill_payload = last_contract.payload if last_contract else {}

    contracts_history = []
    try:
        history_qs = ManagementContract.objects.filter(created_by=request.user).order_by('-created_at')[:50]
        for contract in history_qs:
            payload = contract.payload or {}
            total_sum = payload.get("total_sum")
            try:
                total_sum = int(total_sum) if total_sum not in (None, "") else None
            except Exception:
                total_sum = None
            contracts_history.append({
                "id": contract.id,
                "contract_number": contract.contract_number,
                "contract_date": contract.contract_date.strftime('%d.%m.%Y') if contract.contract_date else "",
                "created_at": timezone.localtime(contract.created_at).strftime('%d.%m.%Y %H:%M') if contract.created_at else "",
                "realizer_name": contract.realizer_name,
                "product_title": payload.get("product_title") or payload.get("product_print") or "",
                "total_sum": total_sum,
                "review_status": contract.review_status,
                "review_reject_reason": contract.review_reject_reason,
                "download_url": reverse('management_contracts_download', args=[contract.id]),
                "file_name": os.path.basename(contract.file_path) if contract.file_path else "",
            })
    except Exception:
        contracts_history = []

    stats = get_user_stats(request.user)
    report_sent_today = has_report_today(request.user)
    progress_clients_pct = min(100, int(stats['processed_today'] / TARGET_CLIENTS_DAY * 100)) if TARGET_CLIENTS_DAY else 0
    progress_points_pct = min(100, int(stats['points_today'] / TARGET_POINTS_DAY * 100)) if TARGET_POINTS_DAY else 0
    bot_username = get_manager_bot_username()
    reminders = get_reminders(request.user, stats=stats, report_sent=report_sent_today)

    return render(request, 'management/contracts.html', {
        'hoodie_products': hoodie_products,
        'tshirt_products': tshirt_products,
        'drop_hoodie_price': int(drop_hoodie_price),
        'drop_tee_price': int(drop_tee_price),
        'contract_date_display': contract_date_display,
        'next_contract_number': next_contract_number,
        'prefill_payload': prefill_payload,
        'contracts_history': contracts_history,
        'user_points_today': stats['points_today'],
        'user_points_total': stats['points_total'],
        'processed_today': stats['processed_today'],
        'target_clients': TARGET_CLIENTS_DAY,
        'target_points': TARGET_POINTS_DAY,
        'progress_clients_pct': progress_clients_pct,
        'progress_points_pct': progress_points_pct,
        'has_report_today': report_sent_today,
        'reminders': reminders,
        'manager_bot_username': bot_username,
    })


@login_required(login_url='management_login')
@require_POST
def contracts_generate_api(request):
    if not user_is_management(request.user):
        return JsonResponse({'ok': False, 'error': 'forbidden'}, status=403)

    try:
        payload = json.loads(request.body or '{}')
    except json.JSONDecodeError:
        return JsonResponse({'ok': False, 'error': 'invalid_json'}, status=400)

    try:
        doc_data, meta = _prepare_contract_data(payload, request, preview=False)
    except ContractPayloadError as exc:
        if exc.code == "missing_fields":
            return JsonResponse({'ok': False, 'error': 'missing_fields', 'fields': exc.fields}, status=400)
        if exc.code == "invalid_product_type":
            return JsonResponse({'ok': False, 'error': 'invalid_product_type'}, status=400)
        return JsonResponse({'ok': False, 'error': 'invalid_payload'}, status=400)

    template_path = _contract_template_path()
    if not template_path.exists():
        return JsonResponse({'ok': False, 'error': 'template_missing'}, status=500)

    contract_dir = Path(settings.MEDIA_ROOT) / "contracts" / str(meta["today"].year)
    contract_dir.mkdir(parents=True, exist_ok=True)
    filename = _safe_contract_filename(doc_data["realizer_name"], meta["contract_number"])
    file_path = contract_dir / filename

    try:
        _build_contract_docx(template_path, str(file_path), doc_data)
    except Exception:
        return JsonResponse({'ok': False, 'error': 'render_failed'}, status=500)

    contract = ManagementContract.objects.create(
        created_by=request.user,
        contract_number=meta["contract_number"],
        contract_date=meta["today"],
        realizer_name=doc_data["realizer_name"],
        file_path=str(file_path),
        payload={
            **payload,
            "contract_number": meta["contract_number"],
            "contract_date": meta["today"].isoformat(),
            "price": meta["price"],
            "price_str": meta["price_str"],
            "total_sum": meta["total_sum"],
            "product_title": meta["product_title"],
            "product_print": doc_data["product_print"],
        },
    )

    return JsonResponse({
        'ok': True,
        'contract': {
            'id': contract.id,
            'contract_number': contract.contract_number,
            'contract_date': meta["today"].strftime('%d.%m.%Y'),
            'created_at': timezone.localtime(contract.created_at).strftime('%d.%m.%Y %H:%M') if contract.created_at else '',
            'realizer_name': contract.realizer_name,
            'product_title': meta["product_title"],
            'total_sum': meta["total_sum"],
            'review_status': contract.review_status,
            'review_reject_reason': contract.review_reject_reason,
            'is_approved': contract.is_approved,
            'download_url': reverse('management_contracts_download', args=[contract.id]),
            'file_name': filename,
        }
    })


@login_required(login_url='management_login')
def contracts_download(request, contract_id):
    if not user_is_management(request.user):
        return redirect('management_login')

    contract = ManagementContract.objects.filter(id=contract_id).first()
    if not contract:
        return HttpResponse('Not found', status=404)
    if not (request.user.is_staff or request.user.is_superuser or contract.created_by_id == request.user.id):
        return HttpResponse('Forbidden', status=403)
    if not contract.file_path or not os.path.exists(contract.file_path):
        return HttpResponse('Not found', status=404)
    filename = os.path.basename(contract.file_path)
    return FileResponse(open(contract.file_path, 'rb'), as_attachment=True, filename=filename)


@login_required(login_url='management_login')
def contracts_list_api(request):
    if not user_is_management(request.user):
        return JsonResponse({'ok': False}, status=403)

    contracts = ManagementContract.objects.filter(created_by=request.user).order_by('-created_at')[:200]
    data = []
    for contract in contracts:
        payload = contract.payload or {}
        total_sum = payload.get("total_sum")
        try:
            total_sum = int(total_sum) if total_sum not in (None, "") else None
        except Exception:
            total_sum = None
        data.append({
            'id': contract.id,
            'contract_number': contract.contract_number,
            'contract_date': contract.contract_date.strftime('%d.%m.%Y') if contract.contract_date else '',
            'created_at': timezone.localtime(contract.created_at).strftime('%d.%m.%Y %H:%M') if contract.created_at else '',
            'realizer_name': contract.realizer_name,
            'product_title': payload.get("product_title") or payload.get("product_print") or '',
            'total_sum': total_sum,
            'review_status': contract.review_status,
            'review_reject_reason': contract.review_reject_reason,
            'is_approved': contract.is_approved,
            'download_url': reverse('management_contracts_download', args=[contract.id]),
        })
    return JsonResponse({'ok': True, 'contracts': data})


@login_required(login_url='management_login')
@require_POST
def contracts_submit_for_review_api(request, contract_id):
    if not user_is_management(request.user):
        return JsonResponse({'ok': False}, status=403)

    contract = ManagementContract.objects.filter(id=contract_id, created_by=request.user).first()
    if not contract:
        return JsonResponse({'ok': False, 'error': 'Договір не знайдено'}, status=404)

    if contract.review_status != 'draft':
        return JsonResponse({'ok': False, 'error': 'Договір вже відправлено або оброблено'}, status=400)

    contract.review_status = 'pending'
    contract.review_reject_reason = ''
    contract.reviewed_at = None
    contract.reviewed_by = None
    contract.is_approved = False
    contract.save(update_fields=[
        'review_status',
        'review_reject_reason',
        'reviewed_at',
        'reviewed_by',
        'is_approved',
    ])

    _send_contract_review_request_to_admin(contract, request=request)
    _notify_manager_contract(
        contract,
        title="📤 <b>Договір відправлено на перевірку</b>",
        body_lines=[
            f"№: <code>{escape(contract.contract_number)}</code>",
            f"Реалізатор: {escape(contract.realizer_name or '—')}",
            "Статус: ⏳ На перевірці",
        ],
    )

    return JsonResponse({'ok': True})


@login_required(login_url='management_login')
@require_POST
def contracts_preview_api(request):
    if not user_is_management(request.user):
        return JsonResponse({'ok': False, 'error': 'forbidden'}, status=403)

    try:
        payload = json.loads(request.body or '{}')
    except json.JSONDecodeError:
        return JsonResponse({'ok': False, 'error': 'invalid_json'}, status=400)

    try:
        doc_data, meta = _prepare_contract_data(payload, request, preview=True)
    except ContractPayloadError as exc:
        if exc.code == "missing_fields":
            return JsonResponse({'ok': False, 'error': 'missing_fields', 'fields': exc.fields}, status=400)
        if exc.code == "invalid_product_type":
            return JsonResponse({'ok': False, 'error': 'invalid_product_type'}, status=400)
        return JsonResponse({'ok': False, 'error': 'invalid_payload'}, status=400)

    template_path = _contract_template_path()
    if not template_path.exists():
        return JsonResponse({'ok': False, 'error': 'template_missing'}, status=500)

    buffer = BytesIO()
    try:
        _build_contract_docx(template_path, buffer, doc_data)
    except Exception:
        return JsonResponse({'ok': False, 'error': 'render_failed'}, status=500)

    buffer.seek(0)
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['X-Contract-Number'] = meta["contract_number"]
    response['Content-Disposition'] = f'inline; filename="preview_{meta["contract_number"]}.docx"'
    return response


@login_required(login_url='management_login')
def invoices(request):
    if not user_is_management(request.user):
        return redirect('management_login')

    from orders.models import WholesaleInvoice

    try:
        tshirt_products, hoodie_products = _get_wholesale_products_for_invoice()
    except Exception:
        tshirt_products, hoodie_products = [], []

    tshirt_prices, hoodie_prices = _get_wholesale_price_context()

    qs = WholesaleInvoice.objects.filter(created_by=request.user).order_by('-created_at')
    last_invoice = qs.first()
    company_data = {
        'company_name': getattr(last_invoice, 'company_name', '') if last_invoice else '',
        'company_number': getattr(last_invoice, 'company_number', '') if last_invoice else '',
        'delivery_address': getattr(last_invoice, 'delivery_address', '') if last_invoice else '',
        'phone_number': getattr(last_invoice, 'contact_phone', '') if last_invoice else '',
        'store_link': getattr(last_invoice, 'store_link', '') if last_invoice else '',
    }
    tshirt_count = tshirt_products.count() if hasattr(tshirt_products, 'count') else len(tshirt_products)
    hoodie_count = hoodie_products.count() if hasattr(hoodie_products, 'count') else len(hoodie_products)
    user_invoice_stats = {
        'total_invoices': qs.count(),
        'last_invoice_date': timezone.localtime(last_invoice.created_at).strftime('%d.%m.%Y %H:%M') if last_invoice else 'Немає накладних',
        'total_products_available': tshirt_count + hoodie_count,
        'tshirt_count': tshirt_count,
        'hoodie_count': hoodie_count,
    }

    stats = get_user_stats(request.user)
    report_sent_today = has_report_today(request.user)
    progress_clients_pct = min(100, int(stats['processed_today'] / TARGET_CLIENTS_DAY * 100)) if TARGET_CLIENTS_DAY else 0
    progress_points_pct = min(100, int(stats['points_today'] / TARGET_POINTS_DAY * 100)) if TARGET_POINTS_DAY else 0
    bot_username = get_manager_bot_username()
    reminders = get_reminders(request.user, stats=stats, report_sent=report_sent_today)

    return render(request, 'management/invoices.html', {
        'tshirt_products': tshirt_products,
        'hoodie_products': hoodie_products,
        'tshirt_prices': tshirt_prices,
        'hoodie_prices': hoodie_prices,
        'company_data': company_data,
        'user_invoice_stats': user_invoice_stats,
        'user_points_today': stats['points_today'],
        'user_points_total': stats['points_total'],
        'processed_today': stats['processed_today'],
        'target_clients': TARGET_CLIENTS_DAY,
        'target_points': TARGET_POINTS_DAY,
        'progress_clients_pct': progress_clients_pct,
        'progress_points_pct': progress_points_pct,
        'has_report_today': report_sent_today,
        'reminders': reminders,
        'manager_bot_username': bot_username,
    })


@login_required(login_url='management_login')
def invoices_list_api(request):
    if not user_is_management(request.user):
        return JsonResponse({'ok': False}, status=403)

    from orders.models import WholesaleInvoice

    invoices = WholesaleInvoice.objects.filter(created_by=request.user).order_by('-created_at')[:200]
    data = []
    for inv in invoices:
        download_filename = ''
        try:
            download_filename = os.path.basename(inv.file_path) if inv.file_path else ''
        except Exception:
            download_filename = ''
        data.append({
            'id': inv.id,
            'invoice_number': inv.invoice_number,
            'company_name': inv.company_name,
            'company_number': inv.company_number,
            'contact_phone': inv.contact_phone,
            'delivery_address': inv.delivery_address,
            'store_link': inv.store_link,
            'total_amount': float(inv.total_amount),
            'total_tshirts': inv.total_tshirts,
            'total_hoodies': inv.total_hoodies,
            'created_at': timezone.localtime(inv.created_at).strftime('%d.%m.%Y %H:%M') if inv.created_at else '',
            'status': inv.status,
            'status_display': inv.get_status_display(),
            'review_status': inv.review_status,
            'review_status_display': inv.get_review_status_display(),
            'review_reject_reason': inv.review_reject_reason,
            'is_approved': inv.is_approved,
            'payment_status': inv.payment_status,
            'payment_url': inv.payment_url,
            'download_filename': download_filename,
        })
    return JsonResponse({'ok': True, 'invoices': data})


@login_required(login_url='management_login')
@require_POST
def invoices_generate_api(request):
    if not user_is_management(request.user):
        return JsonResponse({'ok': False}, status=403)

    from orders.models import WholesaleInvoice

    try:
        payload = json.loads(request.body.decode('utf-8') or '{}')
    except Exception:
        return JsonResponse({'ok': False, 'error': 'Невірний формат даних'}, status=400)

    company_data = payload.get('companyData') or {}
    order_items = payload.get('orderItems') or []

    try:
        normalized_payload = normalize_management_invoice_payload(
            company_data=company_data,
            order_items=order_items,
        )

        normalized_company = normalized_payload["company_data"]
        normalized_items = normalized_payload["items"]
        pricing = normalized_payload["pricing"]
        totals = normalized_payload["totals"]
        serialized_payload = serialize_management_invoice_payload(normalized_payload)

        now = timezone.localtime(timezone.now())
        timestamp = now.strftime('%Y%m%d-%H%M%S')
        nonce = secrets.token_hex(2).upper()
        invoice_number = f"МЕН-{timestamp}-{nonce}"

        company_name = normalized_company["companyName"]
        safe_company = (company_name or 'Company').strip()
        safe_company = re.sub(r'[\\\\/:*?"<>|]+', '_', safe_company)
        safe_company = re.sub(r'\\s+', ' ', safe_company).strip().strip('._-')
        safe_company = (safe_company[:80] or 'Company').strip()
        beautiful_date = now.strftime('%d.%m.%Y_%H-%M')
        file_name = f"{safe_company}_накладнаОПТ_{beautiful_date}.xlsx"
        wb = build_management_invoice_workbook(
            company_data=normalized_company,
            normalized_items=normalized_items,
            pricing=pricing,
            totals=totals,
            invoice_number=invoice_number,
            created_at_label=now.strftime("%d.%m.%Y о %H:%M"),
        )

        invoice = WholesaleInvoice.objects.create(
            invoice_number=invoice_number,
            company_name=company_name,
            company_number=normalized_company["companyNumber"],
            contact_phone=normalized_company["contactPhone"],
            delivery_address=normalized_company["deliveryAddress"],
            store_link=normalized_company["storeLink"],
            total_tshirts=totals["total_tshirts"],
            total_hoodies=totals["total_hoodies"],
            total_amount=totals["total_amount"],
            status='draft',
            created_by=request.user,
            review_status='draft',
            order_details=serialized_payload,
        )

        user_folder = f"invoices/management/user_{request.user.id}"
        invoice_dir = os.path.join(settings.MEDIA_ROOT, user_folder)
        os.makedirs(invoice_dir, exist_ok=True)
        file_path = os.path.join(invoice_dir, file_name)
        wb.save(file_path)

        invoice.file_path = file_path
        invoice.save(update_fields=['file_path'])

        return JsonResponse({
            'ok': True,
            'invoice': {
                'id': invoice.id,
                'invoice_number': invoice.invoice_number,
                'company_name': invoice.company_name,
                'total_amount': float(invoice.total_amount),
                'total_tshirts': invoice.total_tshirts,
                'total_hoodies': invoice.total_hoodies,
                'created_at': timezone.localtime(invoice.created_at).strftime('%d.%m.%Y %H:%M'),
                'review_status': invoice.review_status,
                'review_status_display': invoice.get_review_status_display(),
                'payment_status': invoice.payment_status,
                'payment_url': invoice.payment_url,
                'download_filename': file_name,
            }
        })
    except InvoicePayloadError as exc:
        return JsonResponse({'ok': False, 'error': str(exc)}, status=400)
    except Exception:
        logger.exception("Failed to generate management invoice")
        return JsonResponse({'ok': False, 'error': 'Помилка при створенні накладної. Оновіть сторінку і спробуйте ще раз.'}, status=500)


@login_required(login_url='management_login')
def invoices_download(request, invoice_id):
    if not user_is_management(request.user):
        return redirect('management_login')

    from orders.models import WholesaleInvoice
    import os

    invoice = WholesaleInvoice.objects.filter(id=invoice_id).first()
    if not invoice:
        return HttpResponse('Накладна не знайдена', status=404)

    if not (request.user.is_staff or invoice.created_by_id == request.user.id):
        return HttpResponse('Доступ заборонено', status=403)

    if not invoice.file_path or not os.path.exists(invoice.file_path):
        return HttpResponse('Файл накладної не знайдено', status=404)

    filename = os.path.basename(invoice.file_path)
    return FileResponse(open(invoice.file_path, 'rb'), as_attachment=True, filename=filename)


@login_required(login_url='management_login')
@require_POST
def invoices_delete_api(request, invoice_id):
    if not user_is_management(request.user):
        return JsonResponse({'ok': False}, status=403)

    from orders.models import WholesaleInvoice
    import os

    invoice = WholesaleInvoice.objects.filter(id=invoice_id, created_by=request.user).first()
    if not invoice:
        return JsonResponse({'ok': False, 'error': 'Накладна не знайдена'}, status=404)

    if invoice.review_status == 'pending' or invoice.payment_status == 'paid':
        return JsonResponse({'ok': False, 'error': 'Накладну не можна видалити в цьому статусі'}, status=400)

    try:
        if invoice.file_path and os.path.isabs(invoice.file_path) and os.path.exists(invoice.file_path):
            os.remove(invoice.file_path)
    except Exception:
        pass

    invoice.delete()
    return JsonResponse({'ok': True})


@login_required(login_url='management_login')
@require_POST
def invoices_submit_for_review_api(request, invoice_id):
    if not user_is_management(request.user):
        return JsonResponse({'ok': False}, status=403)

    from orders.models import WholesaleInvoice

    invoice = WholesaleInvoice.objects.filter(id=invoice_id, created_by=request.user).first()
    if not invoice:
        return JsonResponse({'ok': False, 'error': 'Накладна не знайдена'}, status=404)

    if invoice.review_status != 'draft':
        return JsonResponse({'ok': False, 'error': 'Накладна вже відправлена або оброблена'}, status=400)

    invoice.review_status = 'pending'
    invoice.review_reject_reason = ''
    invoice.reviewed_at = None
    invoice.reviewed_by = None
    invoice.is_approved = False
    invoice.payment_url = None
    invoice.monobank_invoice_id = None
    invoice.status = 'pending'
    invoice.save(update_fields=[
        'review_status',
        'review_reject_reason',
        'reviewed_at',
        'reviewed_by',
        'is_approved',
        'payment_url',
        'monobank_invoice_id',
        'status',
    ])

    _send_invoice_review_request_to_admin(invoice, request=request)
    _notify_manager_invoice(
        invoice,
        title="📤 <b>Накладна відправлена на перевірку</b>",
        body_lines=[
            f"№: <code>{escape(invoice.invoice_number)}</code>",
            f"Компанія: {escape(invoice.company_name)}",
            "Статус: ⏳ На перевірці",
        ],
    )

    return JsonResponse({'ok': True})


@login_required(login_url='management_login')
@require_POST
def invoices_create_payment_api(request, invoice_id):
    if not user_is_management(request.user):
        return JsonResponse({'ok': False}, status=403)

    from decimal import Decimal
    import json
    from orders.models import WholesaleInvoice
    from storefront.views.monobank import _monobank_api_request, MonobankAPIError

    invoice = WholesaleInvoice.objects.filter(id=invoice_id, created_by=request.user).first()
    if not invoice:
        return JsonResponse({'ok': False, 'error': 'Накладна не знайдена'}, status=404)

    if invoice.payment_status == 'paid':
        return JsonResponse({'ok': False, 'error': 'Накладна вже оплачена'}, status=400)

    if invoice.review_status != 'approved' or not invoice.is_approved:
        return JsonResponse({'ok': False, 'error': 'Накладна ще не підтверджена адміністратором'}, status=400)

    if invoice.payment_url:
        return JsonResponse({'ok': True, 'payment_url': invoice.payment_url, 'monobank_invoice_id': invoice.monobank_invoice_id})

    try:
        amount_decimal = Decimal(str(invoice.total_amount))
        if amount_decimal <= 0:
            return JsonResponse({'ok': False, 'error': 'Невірна сума накладної'}, status=400)
        amount_kopecks = int(amount_decimal * 100)
    except Exception:
        return JsonResponse({'ok': False, 'error': 'Невірна сума накладної'}, status=400)

    try:
        body = json.loads(request.body.decode('utf-8') or '{}')
    except Exception:
        body = {}

    destination = (body.get('description') or '').strip() or f'Оплата накладної {invoice.invoice_number}'
    payload = {
        'amount': amount_kopecks,
        'ccy': 980,
        'merchantPaymInfo': {
            'reference': f'MGMT-INV-{invoice.id}',
            'destination': destination,
            'basketOrder': [
                {
                    'name': f'Накладна {invoice.invoice_number}',
                    'qty': 1,
                    'sum': amount_kopecks,
                    'icon': '',
                    'unit': 'шт',
                }
            ]
        },
    }

    try:
        creation_data = _monobank_api_request('POST', '/api/merchant/invoice/create', json_payload=payload)
    except MonobankAPIError as exc:
        return JsonResponse({'ok': False, 'error': f'Помилка створення платежу: {str(exc)}'}, status=500)
    except Exception:
        return JsonResponse({'ok': False, 'error': 'Помилка створення платежу'}, status=500)

    payment_url = creation_data.get('invoiceUrl') or creation_data.get('pageUrl')
    monobank_invoice_id = creation_data.get('invoiceId') or creation_data.get('invoice_id')
    if not payment_url:
        return JsonResponse({'ok': False, 'error': 'Не отримано URL для оплати'}, status=500)

    invoice.payment_status = 'pending'
    invoice.payment_url = payment_url
    invoice.monobank_invoice_id = monobank_invoice_id
    invoice.save(update_fields=['payment_status', 'payment_url', 'monobank_invoice_id'])

    return JsonResponse({'ok': True, 'payment_url': payment_url, 'monobank_invoice_id': monobank_invoice_id})


def _default_manager_name(user):
    try:
        prof = user.userprofile
        name = (getattr(prof, "full_name", "") or "").strip()
        if name:
            return name
    except Exception:
        pass
    name = (user.get_full_name() or "").strip()
    if name:
        return name
    return (getattr(user, "username", "") or "").strip()


def _get_profile_phone(user):
    try:
        prof = user.userprofile
        return (getattr(prof, "phone", "") or "").strip()
    except Exception:
        return ""


def _manager_photo_url(user, request) -> str:
    try:
        prof = user.userprofile
        avatar = getattr(prof, "avatar", None)
        if not avatar:
            return ""
        try:
            url = avatar.url
        except Exception:
            return ""
        try:
            return request.build_absolute_uri(url)
        except Exception:
            return url
    except Exception:
        return ""


def _offer_payload_from_form(form, default_name, initial, request):
    def str_val(key: str) -> str:
        if form.is_bound:
            return (form.data.get(key) or "").strip()
        val = initial.get(key)
        if val is None:
            return ""
        return str(val).strip()

    def bool_val(key: str) -> bool:
        if form.is_bound:
            return key in form.data
        return bool(initial.get(key))

    def json_gallery_val(key: str) -> list[dict]:
        raw = form.data.get(key) if form.is_bound else initial.get(key)
        if raw is None:
            return []
        if isinstance(raw, list):
            data = raw
        else:
            s = str(raw).strip()
            if not s:
                return []
            try:
                data = json.loads(s)
            except Exception:
                return []
        if not isinstance(data, list):
            return []

        slots: list[dict] = []
        for item in data:
            if isinstance(item, str):
                url = item.strip()
                caption = ""
            elif isinstance(item, dict):
                url = str(item.get("url") or item.get("link") or item.get("href") or "").strip()
                caption = str(item.get("caption") or item.get("title") or "").strip()
            else:
                continue
            if url:
                slots.append({"url": url, "caption": caption})
        return slots[:6]

    show_manager = bool_val("show_manager")
    manager_name = str_val("manager_name") or default_name

    phone_enabled = bool_val("phone_enabled")

    viber_enabled = bool_val("viber_enabled")
    whatsapp_enabled = bool_val("whatsapp_enabled")
    telegram_enabled = bool_val("telegram_enabled")

    phone = str_val("phone") if (show_manager and phone_enabled) else ""
    viber = str_val("viber") if (show_manager and viber_enabled) else ""
    whatsapp = str_val("whatsapp") if (show_manager and whatsapp_enabled) else ""
    telegram = str_val("telegram") if (show_manager and telegram_enabled) else ""

    segment_mode = str_val("segment_mode") or "NEUTRAL"
    gallery_neutral = json_gallery_val("gallery_neutral")
    gallery_edgy = json_gallery_val("gallery_edgy")
    gallery_urls = gallery_edgy if segment_mode.upper() == "EDGY" else gallery_neutral

    return {
        "shop_name": str_val("recipient_name"),
        "mode": str_val("mode") or "VISUAL",
        "segment_mode": segment_mode,
        "subject_preset": str_val("subject_preset") or "PRESET_1",
        "subject_custom": str_val("subject_custom"),
        "cta_type": str_val("cta_type"),
        "cta_button_text": str_val("cta_button_text"),
        "cta_custom_url": str_val("cta_custom_url"),
        "cta_microtext": str_val("cta_microtext"),
        "pricing_mode": str_val("pricing_mode") or "OPT",
        "opt_tier": str_val("opt_tier") or "8_15",
        "drop_tee_price": str_val("drop_tee_price"),
        "drop_hoodie_price": str_val("drop_hoodie_price"),
        "dropship_loyalty_bonus": bool_val("dropship_loyalty_bonus"),
        "tee_entry": str_val("tee_entry"),
        "tee_retail_example": str_val("tee_retail_example"),
        "hoodie_entry": str_val("hoodie_entry"),
        "hoodie_retail_example": str_val("hoodie_retail_example"),
        "show_manager": show_manager,
        "manager_name": manager_name,
        "phone_enabled": phone_enabled,
        "phone": phone,
        "viber_enabled": viber_enabled,
        "viber": viber,
        "whatsapp_enabled": whatsapp_enabled,
        "whatsapp": whatsapp,
        "telegram_enabled": telegram_enabled,
        "telegram": telegram,
        "general_tg": str_val("general_tg"),
        "include_catalog_link": bool_val("include_catalog_link"),
        "include_wholesale_link": bool_val("include_wholesale_link"),
        "include_dropship_link": bool_val("include_dropship_link"),
        "include_instagram_link": bool_val("include_instagram_link"),
        "include_site_link": bool_val("include_site_link"),
        "gallery_urls": gallery_urls,
        "manager_photo_url": _manager_photo_url(request.user, request) if show_manager else "",
    }


def _build_cp_messenger_context(*, user, settings_obj, default_name, default_phone):
    # Контекст для генератора Telegram-повідомлень у месенджерах (plain text)
    try:
        prof = user.userprofile
    except Exception:
        prof = None

    manager_name = (
        (getattr(prof, "full_name", "") or "").strip()
        or (default_name or "").strip()
        or (user.get_full_name() or "").strip()
        or (getattr(user, "username", "") or "").strip()
    )

    phone = ((getattr(prof, "phone", "") or "").strip() or (default_phone or "").strip()).strip()

    telegram_raw = ""
    try:
        if getattr(settings_obj, "show_manager", True) and getattr(settings_obj, "telegram_enabled", False):
            telegram_raw = (getattr(settings_obj, "telegram", "") or "").strip()
    except Exception:
        telegram_raw = ""

    if not telegram_raw:
        telegram_raw = (getattr(prof, "telegram", "") or "").strip() if prof else ""

    telegram = telegram_raw
    if telegram and not telegram.startswith("@"):
        if "/" not in telegram and ":" not in telegram and not telegram.startswith("+"):
            telegram = f"@{telegram}"

    whatsapp = ""
    viber = ""
    try:
        if getattr(settings_obj, "show_manager", True) and getattr(settings_obj, "whatsapp_enabled", False):
            whatsapp = (getattr(settings_obj, "whatsapp", "") or "").strip()
        if getattr(settings_obj, "show_manager", True) and getattr(settings_obj, "viber_enabled", False):
            viber = (getattr(settings_obj, "viber", "") or "").strip()
    except Exception:
        pass

    base_url = (getattr(settings, "SITE_BASE_URL", "") or "").strip() or "https://twocomms.shop"
    if not base_url.endswith("/"):
        base_url += "/"

    def abs_url(path: str) -> str:
        if not path:
            return base_url
        if path.startswith(("http://", "https://")):
            return path
        return f"{base_url}{path.lstrip('/')}"

    def normalize_tg_link(value: str) -> str:
        v = (value or "").strip()
        if not v:
            return ""
        if v.startswith(("http://", "https://", "tg:")):
            return v
        if v.startswith("@"):
            return f"https://t.me/{v[1:]}"
        if v.startswith("t.me/"):
            return f"https://{v}"
        if "t.me/" in v:
            return v if v.startswith(("http://", "https://")) else f"https://{v}"
        return f"https://t.me/{v}"

    general_tg_raw = (getattr(settings_obj, "general_tg", "") or "").strip()
    general_tg = normalize_tg_link(general_tg_raw) or "https://t.me/twocomms"

    opt_tiers = {}
    for tier_key, label in OPT_TIER_LABELS.items():
        opt_tiers[tier_key] = {
            "label": label,
            "tee": int(OPT_TIER_WHOLESALE_TEE[tier_key]),
            "hoodie": int(OPT_TIER_WHOLESALE_HOODIE[tier_key]),
        }

    retail_defaults = get_twocomms_cp_unit_defaults(pricing_mode="OPT", opt_tier="100_PLUS")

    pricing_config = {
        "optTiers": opt_tiers,
        "dropFixed": {
            "tee": int(DROP_FIXED_TEE_PRICE),
            "hoodie": int(DROP_FIXED_HOODIE_PRICE),
        },
        "retailExamples": {
            "tee": int(retail_defaults.get("tee_retail_default") or 0),
            "hoodie": int(retail_defaults.get("hoodie_retail_default") or 0),
        },
        "dropshipLoyaltyStep": 10,
        "maxDropDiscount": 120,
    }

    return {
        "manager": {
            "name": manager_name or "менеджер TwoComms",
            "phone": phone,
            "telegram": telegram,
        },
        "links": {
            "catalog": abs_url("/catalog/"),
            "wholesale": abs_url("/wholesale/"),
            "dropship": abs_url("/dropshipper/"),
            "general_tg": general_tg,
        },
        "pricing": pricing_config,
    }


@login_required(login_url='management_login')
def commercial_offer_email(request):

    if not user_is_management(request.user):
        return redirect('management_login')

    settings_obj, _ = CommercialOfferEmailSettings.objects.get_or_create(owner=request.user)
    default_name = (settings_obj.manager_name or "").strip() or _default_manager_name(request.user)
    default_phone = (settings_obj.phone or "").strip() or _get_profile_phone(request.user)

    initial = {
        "show_manager": settings_obj.show_manager,
        "manager_name": default_name,
        "phone_enabled": getattr(settings_obj, "phone_enabled", False) or bool(default_phone),
        "phone": default_phone,
        "viber_enabled": settings_obj.viber_enabled,
        "viber": (settings_obj.viber or "").strip(),
        "whatsapp_enabled": settings_obj.whatsapp_enabled,
        "whatsapp": (settings_obj.whatsapp or "").strip(),
        "telegram_enabled": settings_obj.telegram_enabled,
        "telegram": (settings_obj.telegram or "").strip(),
        "general_tg": (getattr(settings_obj, "general_tg", "") or "").strip(),
        "pricing_mode": getattr(settings_obj, "pricing_mode", "OPT") or "OPT",
        "opt_tier": getattr(settings_obj, "opt_tier", "8_15") or "8_15",
        "drop_tee_price": getattr(settings_obj, "drop_tee_price", None),
        "drop_hoodie_price": getattr(settings_obj, "drop_hoodie_price", None),
        "dropship_loyalty_bonus": bool(getattr(settings_obj, "dropship_loyalty_bonus", False)),
        "include_catalog_link": getattr(settings_obj, "include_catalog_link", True),
        "include_wholesale_link": getattr(settings_obj, "include_wholesale_link", True),
        "include_dropship_link": getattr(settings_obj, "include_dropship_link", True),
        "include_instagram_link": getattr(settings_obj, "include_instagram_link", True),
        "include_site_link": getattr(settings_obj, "include_site_link", True),
        "cta_type": (getattr(settings_obj, "cta_type", "") or "").strip(),
        "cta_custom_url": (getattr(settings_obj, "cta_custom_url", "") or "").strip(),
        "cta_button_text": (getattr(settings_obj, "cta_button_text", "") or "").strip(),
        "cta_microtext": (getattr(settings_obj, "cta_microtext", "") or "").strip(),
        "gallery_neutral": list(getattr(settings_obj, "gallery_neutral", []) or []),
        "gallery_edgy": list(getattr(settings_obj, "gallery_edgy", []) or []),
        "mode": getattr(settings_obj, "mode", "VISUAL") or "VISUAL",
        "segment_mode": getattr(settings_obj, "segment_mode", "NEUTRAL") or "NEUTRAL",
        "subject_preset": getattr(settings_obj, "subject_preset", "PRESET_1") or "PRESET_1",
        "subject_custom": (getattr(settings_obj, "subject_custom", "") or "").strip(),
        "tee_entry": getattr(settings_obj, "tee_entry", None),
        "tee_retail_example": getattr(settings_obj, "tee_retail_example", None),
        "hoodie_entry": getattr(settings_obj, "hoodie_entry", None),
        "hoodie_retail_example": getattr(settings_obj, "hoodie_retail_example", None),
    }

    if not getattr(settings_obj, "gallery_initialized", False):
        try:
            if not initial.get("gallery_neutral"):
                initial["gallery_neutral"] = build_twocomms_cp_email({"segment_mode": "NEUTRAL"}).get("gallery_urls") or []
            if not initial.get("gallery_edgy"):
                initial["gallery_edgy"] = build_twocomms_cp_email({"segment_mode": "EDGY"}).get("gallery_urls") or []
        except Exception:
            pass

    sent_success = request.GET.get("sent") == "1"
    send_error = ""

    cp_tab = (request.GET.get("tab") or "email").strip().lower()
    if cp_tab not in ("email", "messengers"):
        cp_tab = "email"

    messenger_context = _build_cp_messenger_context(
        user=request.user,
        settings_obj=settings_obj,
        default_name=default_name,
        default_phone=default_phone,
    )

    if request.method == "POST":
        form = CommercialOfferEmailForm(request.POST, user=request.user)
        if form.is_valid():
            cd = form.cleaned_data
            recipient_email = (cd.get("recipient_email") or "").strip()

            duplicate_qs = CommercialOfferEmailLog.objects.filter(
                recipient_email__iexact=recipient_email,
                status=CommercialOfferEmailLog.Status.SENT,
            )
            if duplicate_qs.exists() and cd.get("confirm_resend") != 1:
                form.add_error(
                    "recipient_email",
                    "На цей email вже надсилали КП. Підтвердіть повторну відправку.",
                )
            else:
                settings_obj.show_manager = bool(cd.get("show_manager"))
                settings_obj.manager_name = (cd.get("manager_name") or "").strip()
                settings_obj.phone_enabled = bool(cd.get("phone_enabled"))
                settings_obj.phone = (cd.get("phone") or "").strip()
                settings_obj.viber_enabled = bool(cd.get("viber_enabled"))
                settings_obj.viber = (cd.get("viber") or "").strip()
                settings_obj.whatsapp_enabled = bool(cd.get("whatsapp_enabled"))
                settings_obj.whatsapp = (cd.get("whatsapp") or "").strip()
                settings_obj.telegram_enabled = bool(cd.get("telegram_enabled"))
                settings_obj.telegram = (cd.get("telegram") or "").strip()

                settings_obj.general_tg = (cd.get("general_tg") or "").strip()
                settings_obj.pricing_mode = (cd.get("pricing_mode") or "OPT").strip().upper()
                settings_obj.opt_tier = (cd.get("opt_tier") or "8_15").strip().upper()
                settings_obj.drop_tee_price = cd.get("drop_tee_price") or None
                settings_obj.drop_hoodie_price = cd.get("drop_hoodie_price") or None
                settings_obj.dropship_loyalty_bonus = bool(cd.get("dropship_loyalty_bonus"))
                settings_obj.include_catalog_link = bool(cd.get("include_catalog_link"))
                settings_obj.include_wholesale_link = bool(cd.get("include_wholesale_link"))
                settings_obj.include_dropship_link = bool(cd.get("include_dropship_link"))
                settings_obj.include_instagram_link = bool(cd.get("include_instagram_link"))
                settings_obj.include_site_link = bool(cd.get("include_site_link"))

                settings_obj.cta_type = (cd.get("cta_type") or "").strip().upper()
                settings_obj.cta_custom_url = (cd.get("cta_custom_url") or "").strip()
                settings_obj.cta_button_text = (cd.get("cta_button_text") or "").strip()
                settings_obj.cta_microtext = (cd.get("cta_microtext") or "").strip()

                settings_obj.gallery_neutral = cd.get("gallery_neutral") or []
                settings_obj.gallery_edgy = cd.get("gallery_edgy") or []
                settings_obj.gallery_initialized = True

                settings_obj.mode = (cd.get("mode") or "VISUAL").strip().upper()
                settings_obj.segment_mode = (cd.get("segment_mode") or "NEUTRAL").strip().upper()
                settings_obj.subject_preset = (cd.get("subject_preset") or "PRESET_1").strip().upper()
                settings_obj.subject_custom = (cd.get("subject_custom") or "").strip()
                settings_obj.tee_entry = cd.get("tee_entry") or None
                settings_obj.tee_retail_example = cd.get("tee_retail_example") or None
                settings_obj.hoodie_entry = cd.get("hoodie_entry") or None
                settings_obj.hoodie_retail_example = cd.get("hoodie_retail_example") or None
                settings_obj.save()

                payload = {
                    "shop_name": (cd.get("recipient_name") or "").strip(),
                    "mode": settings_obj.mode,
                    "segment_mode": settings_obj.segment_mode,
                    "subject_preset": settings_obj.subject_preset,
                    "subject_custom": settings_obj.subject_custom,
                    "cta_type": settings_obj.cta_type,
                    "cta_custom_url": settings_obj.cta_custom_url,
                    "cta_button_text": settings_obj.cta_button_text,
                    "cta_microtext": settings_obj.cta_microtext,
                    "pricing_mode": getattr(settings_obj, "pricing_mode", "OPT") or "OPT",
                    "opt_tier": getattr(settings_obj, "opt_tier", "8_15") or "8_15",
                    "drop_tee_price": getattr(settings_obj, "drop_tee_price", None),
                    "drop_hoodie_price": getattr(settings_obj, "drop_hoodie_price", None),
                    "dropship_loyalty_bonus": bool(getattr(settings_obj, "dropship_loyalty_bonus", False)),
                    "tee_entry": settings_obj.tee_entry,
                    "tee_retail_example": settings_obj.tee_retail_example,
                    "hoodie_entry": settings_obj.hoodie_entry,
                    "hoodie_retail_example": settings_obj.hoodie_retail_example,
                    "show_manager": settings_obj.show_manager,
                    "manager_name": settings_obj.manager_name or default_name,
                    "phone_enabled": settings_obj.phone_enabled,
                    "phone": settings_obj.phone if (settings_obj.show_manager and settings_obj.phone_enabled) else "",
                    "viber": settings_obj.viber if (settings_obj.show_manager and settings_obj.viber_enabled) else "",
                    "whatsapp": settings_obj.whatsapp if (settings_obj.show_manager and settings_obj.whatsapp_enabled) else "",
                    "telegram": settings_obj.telegram if (settings_obj.show_manager and settings_obj.telegram_enabled) else "",
                    "general_tg": settings_obj.general_tg,
                    "include_catalog_link": settings_obj.include_catalog_link,
                    "include_wholesale_link": settings_obj.include_wholesale_link,
                    "include_dropship_link": settings_obj.include_dropship_link,
                    "include_instagram_link": settings_obj.include_instagram_link,
                    "include_site_link": settings_obj.include_site_link,
                    "gallery_urls": settings_obj.gallery_edgy if settings_obj.segment_mode == "EDGY" else settings_obj.gallery_neutral,
                    "manager_photo_url": _manager_photo_url(request.user, request) if settings_obj.show_manager else "",
                }
                email_build = build_twocomms_cp_email(payload)
                subject = email_build["subject"]
                preheader = email_build["preheader"]
                text_body = email_build["text"]
                html_body_to_send = email_build["html"] if settings_obj.mode == "VISUAL" else email_build["html_light"]

                from_email = getattr(settings, "DEFAULT_FROM_EMAIL", None) or "TwoComms <cooperation@twocomms.shop>"
                reply_to = [settings.EMAIL_HOST_USER] if getattr(settings, "EMAIL_HOST_USER", "") else None

                status = CommercialOfferEmailLog.Status.SENT
                error_text = ""
                try:
                    msg = EmailMultiAlternatives(
                        subject=subject,
                        body=text_body,
                        from_email=from_email,
                        to=[recipient_email],
                        reply_to=reply_to,
                    )
                    msg.attach_alternative(html_body_to_send, "text/html")
                    msg.send(fail_silently=False)
                except Exception as exc:
                    status = CommercialOfferEmailLog.Status.FAILED
                    error_text = str(exc)
                    send_error = "Не вдалося відправити лист. Перевірте SMTP налаштування та спробуйте ще раз."

                CommercialOfferEmailLog.objects.create(
                    owner=request.user,
                    recipient_email=recipient_email,
                    recipient_name=(cd.get("recipient_name") or "").strip(),
                    subject=subject,
                    preheader=preheader,
                    body_html=html_body_to_send,
                    body_text=text_body,
                    mode=settings_obj.mode,
                    segment_mode=settings_obj.segment_mode,
                    subject_preset=settings_obj.subject_preset,
                    subject_custom=settings_obj.subject_custom,
                    cta_type=email_build.get("cta_type") or settings_obj.cta_type,
                    cta_url=email_build.get("cta_url") or "",
                    cta_custom_url=settings_obj.cta_custom_url,
                    cta_button_text=email_build.get("cta_button_text") or settings_obj.cta_button_text,
                    cta_microtext=email_build.get("cta_microtext") or settings_obj.cta_microtext,
                    general_tg=settings_obj.general_tg,
                    pricing_mode=email_build.get("pricing_mode") or getattr(settings_obj, "pricing_mode", "OPT") or "OPT",
                    opt_tier=email_build.get("opt_tier") or getattr(settings_obj, "opt_tier", "8_15") or "8_15",
                    drop_tee_price=email_build.get("drop_tee_price") or getattr(settings_obj, "drop_tee_price", None),
                    drop_hoodie_price=email_build.get("drop_hoodie_price") or getattr(settings_obj, "drop_hoodie_price", None),
                    dropship_loyalty_bonus=bool(
                        email_build.get("dropship_loyalty_bonus")
                        if ("dropship_loyalty_bonus" in email_build)
                        else getattr(settings_obj, "dropship_loyalty_bonus", False)
                    ),
                    include_catalog_link=settings_obj.include_catalog_link,
                    include_wholesale_link=settings_obj.include_wholesale_link,
                    include_dropship_link=settings_obj.include_dropship_link,
                    include_instagram_link=settings_obj.include_instagram_link,
                    include_site_link=settings_obj.include_site_link,
                    gallery_urls=email_build.get("gallery_urls") or [],
                    gallery_items=email_build.get("gallery_items") or [],
                    tee_entry=email_build.get("tee_entry"),
                    tee_retail_example=email_build.get("tee_retail_example"),
                    tee_profit=email_build.get("tee_profit"),
                    hoodie_entry=email_build.get("hoodie_entry"),
                    hoodie_retail_example=email_build.get("hoodie_retail_example"),
                    hoodie_profit=email_build.get("hoodie_profit"),
                    show_manager=settings_obj.show_manager,
                    manager_name=(settings_obj.manager_name or default_name) if settings_obj.show_manager else "",
                    phone_enabled=settings_obj.phone_enabled,
                    phone=settings_obj.phone if (settings_obj.show_manager and settings_obj.phone_enabled) else "",
                    viber=settings_obj.viber if (settings_obj.show_manager and settings_obj.viber_enabled) else "",
                    whatsapp=settings_obj.whatsapp if (settings_obj.show_manager and settings_obj.whatsapp_enabled) else "",
                    telegram=settings_obj.telegram if (settings_obj.show_manager and settings_obj.telegram_enabled) else "",
                    status=status,
                    error=error_text,
                )

                if status == CommercialOfferEmailLog.Status.SENT:
                    return redirect(f"{reverse('management_commercial_offer_email')}?sent=1")
    else:
        form = CommercialOfferEmailForm(initial=initial, user=request.user)

    preview_payload = _offer_payload_from_form(form, default_name, initial, request)
    preview_build = build_twocomms_cp_email(preview_payload)
    preview_visual_html = preview_build["html"]
    preview_light_text = preview_build["text"]
    preview_subject = preview_build["subject"]
    preview_preheader = preview_build["preheader"]

    logs = CommercialOfferEmailLog.objects.filter(owner=request.user).order_by("-created_at")[:30]

    gallery_neutral_json = "[]"
    gallery_edgy_json = "[]"
    try:
        if form.is_bound:
            gallery_neutral_json = (form.data.get("gallery_neutral") or "[]").strip() or "[]"
            gallery_edgy_json = (form.data.get("gallery_edgy") or "[]").strip() or "[]"
        else:
            gallery_neutral_json = json.dumps(initial.get("gallery_neutral") or [])
            gallery_edgy_json = json.dumps(initial.get("gallery_edgy") or [])
    except Exception:
        gallery_neutral_json = "[]"
        gallery_edgy_json = "[]"

    return render(
        request,
        "management/commercial_offer_email.html",
        {
            "form": form,
            "preview_visual_html": preview_visual_html,
            "preview_light_text": preview_light_text,
            "preview_subject": preview_subject,
            "preview_preheader": preview_preheader,
            "preview_mode": preview_payload.get("mode", "VISUAL"),
            "logs": logs,
            "sent_success": sent_success,
            "send_error": send_error,
            "gallery_neutral_json": gallery_neutral_json,
            "gallery_edgy_json": gallery_edgy_json,
            "cp_tab": cp_tab,
            "messenger_context": messenger_context,
        },
    )


@login_required(login_url='management_login')
def commercial_offer_email_check_api(request):
    if not user_is_management(request.user):
        return JsonResponse({"ok": False, "error": "forbidden"}, status=403)

    email = (request.GET.get("email") or "").strip()
    if not email:
        return JsonResponse({"ok": True, "exists": False, "count": 0, "lastSentAtDisplay": ""})

    try:
        validate_email(email)
    except ValidationError:
        return JsonResponse({"ok": True, "exists": False, "count": 0, "lastSentAtDisplay": ""})

    qs = CommercialOfferEmailLog.objects.filter(
        recipient_email__iexact=email,
        status=CommercialOfferEmailLog.Status.SENT,
    )
    last = qs.order_by("-created_at").first()
    last_display = timezone.localtime(last.created_at).strftime("%d.%m.%Y %H:%M") if last else ""

    return JsonResponse(
        {
            "ok": True,
            "exists": qs.exists(),
            "count": qs.count(),
            "lastSentAtDisplay": last_display,
        }
    )


@login_required(login_url="management_login")
def commercial_offer_email_unit_defaults_api(request):
    if not user_is_management(request.user):
        return JsonResponse({"ok": False, "error": "forbidden"}, status=403)

    pricing_mode = (request.GET.get("pricing_mode") or "OPT").strip().upper()
    opt_tier = (request.GET.get("opt_tier") or "8_15").strip().upper()
    drop_tee_price = request.GET.get("drop_tee_price")
    drop_hoodie_price = request.GET.get("drop_hoodie_price")
    try:
        drop_tee_val = int(drop_tee_price) if (drop_tee_price is not None and str(drop_tee_price).strip()) else None
    except Exception:
        drop_tee_val = None
    try:
        drop_hoodie_val = int(drop_hoodie_price) if (drop_hoodie_price is not None and str(drop_hoodie_price).strip()) else None
    except Exception:
        drop_hoodie_val = None

    defaults = get_twocomms_cp_unit_defaults(
        pricing_mode="DROP" if pricing_mode == "DROP" else "OPT",
        opt_tier=opt_tier if opt_tier in {"8_15", "16_31", "32_63", "64_99", "100_PLUS"} else "8_15",
        drop_tee_price=drop_tee_val,
        drop_hoodie_price=drop_hoodie_val,
    )
    return JsonResponse({"ok": True, "defaults": defaults})


@login_required(login_url="management_login")
def commercial_offer_email_gallery_defaults_api(request):
    if not user_is_management(request.user):
        return JsonResponse({"ok": False, "error": "forbidden"}, status=403)

    seg = (request.GET.get("segment") or "NEUTRAL").strip().upper()
    if seg not in {"NEUTRAL", "EDGY"}:
        seg = "NEUTRAL"

    email_build = build_twocomms_cp_email({"segment_mode": seg})
    urls = email_build.get("gallery_urls") or []
    return JsonResponse({"ok": True, "segment_mode": seg, "urls": urls, "source": "site" if urls else "fallback"})


@login_required(login_url="management_login")
def commercial_offer_email_resolve_product_api(request):
    if not user_is_management(request.user):
        return JsonResponse({"ok": False, "error": "forbidden"}, status=403)

    raw_url = (request.GET.get("url") or "").strip()
    if not raw_url:
        return JsonResponse({"ok": False, "error": "no_url"}, status=400)

    def site_base_url() -> str:
        base = (getattr(settings, "SITE_BASE_URL", "") or "").strip() or "https://twocomms.shop"
        if not base.endswith("/"):
            base += "/"
        return base

    def abs_url(path_or_url: str) -> str:
        if not path_or_url:
            return site_base_url()
        if path_or_url.startswith(("http://", "https://")):
            return path_or_url
        from urllib.parse import urljoin

        return urljoin(site_base_url(), path_or_url.lstrip("/"))

    try:
        import re
        from urllib.parse import urlparse

        path = urlparse(raw_url).path if raw_url.startswith(("http://", "https://")) else raw_url
        match = re.search(r"/product/(?P<slug>[-a-zA-Z0-9_]+)/?", path)
        slug = match.group("slug") if match else ""
    except Exception:
        slug = ""

    if not slug:
        try:
            import re as _re
            from urllib.parse import urlparse as _urlparse

            path = _urlparse(raw_url).path if raw_url.startswith(("http://", "https://")) else raw_url
            is_image = bool(_re.search(r"\.(png|jpe?g|webp|gif)(?:\\?|$)", path or "", _re.IGNORECASE))
        except Exception:
            is_image = False
        if not is_image:
            return JsonResponse({"ok": False, "error": "bad_url"}, status=400)

        img_url = abs_url(raw_url)
        item = {
            "title": "",
            "img_url": img_url,
            "link_url": img_url,
            "retail": None,
        }
        return JsonResponse({"ok": True, "item": item})

    try:
        from storefront.models import Product, ProductStatus

        product = Product.objects.select_related("category").filter(status=ProductStatus.PUBLISHED, slug=slug).first()
    except Exception:
        product = None

    if not product:
        return JsonResponse({"ok": False, "error": "not_found"}, status=404)

    img = getattr(product, "display_image", None)
    img_url = ""
    try:
        img_url = abs_url(img.url) if img and getattr(img, "url", None) else ""
    except Exception:
        img_url = ""

    if not img_url:
        return JsonResponse({"ok": False, "error": "no_image"}, status=404)

    item = {
        "title": (product.title or "").strip(),
        "img_url": img_url,
        "link_url": abs_url(f"/product/{product.slug}/"),
        "retail": getattr(product, "final_price", None),
    }
    return JsonResponse({"ok": True, "item": item})


@require_POST
@login_required(login_url="management_login")
def commercial_offer_email_resend_api(request, log_id: int):
    if not user_is_management(request.user):
        return JsonResponse({"ok": False, "error": "forbidden"}, status=403)

    try:
        original = CommercialOfferEmailLog.objects.get(id=log_id, owner=request.user)
    except CommercialOfferEmailLog.DoesNotExist:
        return JsonResponse({"ok": False, "error": "not_found"}, status=404)

    from_email = getattr(settings, "DEFAULT_FROM_EMAIL", None) or "TwoComms <cooperation@twocomms.shop>"
    reply_to = [settings.EMAIL_HOST_USER] if getattr(settings, "EMAIL_HOST_USER", "") else None

    status = CommercialOfferEmailLog.Status.SENT
    error_text = ""
    try:
        msg = EmailMultiAlternatives(
            subject=original.subject,
            body=original.body_text or "",
            from_email=from_email,
            to=[original.recipient_email],
            reply_to=reply_to,
        )
        if original.body_html:
            msg.attach_alternative(original.body_html, "text/html")
        msg.send(fail_silently=False)
    except Exception as exc:
        status = CommercialOfferEmailLog.Status.FAILED
        error_text = str(exc)

    new_log = CommercialOfferEmailLog.objects.create(
        owner=request.user,
        recipient_email=original.recipient_email,
        recipient_name=original.recipient_name,
        subject=original.subject,
        preheader=getattr(original, "preheader", "") or "",
        body_html=original.body_html,
        body_text=original.body_text,
        mode=original.mode,
        segment_mode=original.segment_mode,
        subject_preset=original.subject_preset,
        subject_custom=original.subject_custom,
        cta_type=getattr(original, "cta_type", "") or "",
        cta_url=getattr(original, "cta_url", "") or "",
        cta_custom_url=getattr(original, "cta_custom_url", "") or "",
        cta_button_text=getattr(original, "cta_button_text", "") or "",
        cta_microtext=getattr(original, "cta_microtext", "") or "",
        general_tg=getattr(original, "general_tg", "") or "",
        pricing_mode=getattr(original, "pricing_mode", "OPT") or "OPT",
        opt_tier=getattr(original, "opt_tier", "8_15") or "8_15",
        drop_tee_price=getattr(original, "drop_tee_price", None),
        drop_hoodie_price=getattr(original, "drop_hoodie_price", None),
        dropship_loyalty_bonus=bool(getattr(original, "dropship_loyalty_bonus", False)),
        include_catalog_link=bool(getattr(original, "include_catalog_link", True)),
        include_wholesale_link=bool(getattr(original, "include_wholesale_link", True)),
        include_dropship_link=bool(getattr(original, "include_dropship_link", True)),
        include_instagram_link=bool(getattr(original, "include_instagram_link", True)),
        include_site_link=bool(getattr(original, "include_site_link", True)),
        gallery_urls=getattr(original, "gallery_urls", []) or [],
        gallery_items=getattr(original, "gallery_items", []) or [],
        tee_entry=original.tee_entry,
        tee_retail_example=original.tee_retail_example,
        tee_profit=original.tee_profit,
        hoodie_entry=original.hoodie_entry,
        hoodie_retail_example=original.hoodie_retail_example,
        hoodie_profit=original.hoodie_profit,
        show_manager=original.show_manager,
        manager_name=original.manager_name,
        phone_enabled=bool(getattr(original, "phone_enabled", False)),
        phone=original.phone,
        viber=original.viber,
        whatsapp=original.whatsapp,
        telegram=original.telegram,
        status=status,
        error=error_text,
    )

    row_html = render_to_string("management/partials/commercial_offer_log_row.html", {"log": new_log})
    return JsonResponse(
        {
            "ok": True,
            "sent": status == CommercialOfferEmailLog.Status.SENT,
            "status": status,
            "row_html": row_html,
            "message": "КП успішно надіслано." if status == CommercialOfferEmailLog.Status.SENT else "Не вдалося відправити лист.",
            "error": error_text,
        }
    )


@require_POST
@login_required(login_url="management_login")
def commercial_offer_email_preview_api(request):
    if not user_is_management(request.user):
        return JsonResponse({"ok": False, "error": "forbidden"}, status=403)

    settings_obj, _ = CommercialOfferEmailSettings.objects.get_or_create(owner=request.user)
    default_name = (settings_obj.manager_name or "").strip() or _default_manager_name(request.user)
    default_phone = (settings_obj.phone or "").strip() or _get_profile_phone(request.user)

    data = request.POST

    def str_val(key: str) -> str:
        return (data.get(key) or "").strip()

    def bool_val(key: str) -> bool:
        return key in data

    show_manager = bool_val("show_manager")
    manager_name = str_val("manager_name") or default_name

    phone_enabled = bool_val("phone_enabled")
    viber_enabled = bool_val("viber_enabled")
    whatsapp_enabled = bool_val("whatsapp_enabled")
    telegram_enabled = bool_val("telegram_enabled")

    def json_gallery_val(key: str) -> list[dict]:
        raw = (data.get(key) or "").strip()
        if not raw:
            return []
        try:
            parsed = json.loads(raw)
        except Exception:
            return []
        if not isinstance(parsed, list):
            return []
        slots: list[dict] = []
        for item in parsed:
            if isinstance(item, str):
                url = item.strip()
                caption = ""
            elif isinstance(item, dict):
                url = str(item.get("url") or item.get("link") or item.get("href") or "").strip()
                caption = str(item.get("caption") or item.get("title") or "").strip()
            else:
                continue
            if url:
                slots.append({"url": url, "caption": caption})
        return slots[:6]

    segment_mode_val = (str_val("segment_mode") or getattr(settings_obj, "segment_mode", "NEUTRAL") or "NEUTRAL").upper()
    gallery_neutral = json_gallery_val("gallery_neutral")
    gallery_edgy = json_gallery_val("gallery_edgy")
    gallery_urls = gallery_edgy if segment_mode_val == "EDGY" else gallery_neutral

    pricing_mode_val = (str_val("pricing_mode") or getattr(settings_obj, "pricing_mode", "OPT") or "OPT").upper()
    opt_tier_val = (str_val("opt_tier") or getattr(settings_obj, "opt_tier", "8_15") or "8_15").upper()

    payload = {
        "shop_name": str_val("recipient_name"),
        "mode": (str_val("mode") or getattr(settings_obj, "mode", "VISUAL") or "VISUAL").upper(),
        "segment_mode": segment_mode_val,
        "subject_preset": (str_val("subject_preset") or getattr(settings_obj, "subject_preset", "PRESET_1") or "PRESET_1").upper(),
        "subject_custom": str_val("subject_custom"),
        "cta_type": str_val("cta_type"),
        "cta_custom_url": str_val("cta_custom_url"),
        "cta_button_text": str_val("cta_button_text"),
        "cta_microtext": str_val("cta_microtext"),
        "pricing_mode": pricing_mode_val,
        "opt_tier": opt_tier_val,
        "drop_tee_price": str_val("drop_tee_price"),
        "drop_hoodie_price": str_val("drop_hoodie_price"),
        "dropship_loyalty_bonus": bool_val("dropship_loyalty_bonus"),
        "tee_entry": str_val("tee_entry"),
        "tee_retail_example": str_val("tee_retail_example"),
        "hoodie_entry": str_val("hoodie_entry"),
        "hoodie_retail_example": str_val("hoodie_retail_example"),
        "show_manager": show_manager,
        "manager_name": manager_name if show_manager else "",
        "phone_enabled": phone_enabled,
        "phone": (str_val("phone") or default_phone) if (show_manager and phone_enabled) else "",
        "viber": str_val("viber") if (show_manager and viber_enabled) else "",
        "whatsapp": str_val("whatsapp") if (show_manager and whatsapp_enabled) else "",
        "telegram": str_val("telegram") if (show_manager and telegram_enabled) else "",
        "general_tg": str_val("general_tg"),
        "include_catalog_link": bool_val("include_catalog_link"),
        "include_wholesale_link": bool_val("include_wholesale_link"),
        "include_dropship_link": bool_val("include_dropship_link"),
        "include_instagram_link": bool_val("include_instagram_link"),
        "include_site_link": bool_val("include_site_link"),
        "gallery_urls": gallery_urls,
        "manager_photo_url": _manager_photo_url(request.user, request) if show_manager else "",
    }

    email_build = build_twocomms_cp_email(payload)
    return JsonResponse(
        {
            "ok": True,
            "subject": email_build["subject"],
            "preheader": email_build["preheader"],
            "html": email_build["html"],
            "text": email_build["text"],
            "mode": email_build.get("mode", payload.get("mode")),
            "segment_mode": email_build.get("segment_mode", payload.get("segment_mode")),
            "tee_profit": email_build.get("tee_profit"),
            "hoodie_profit": email_build.get("hoodie_profit"),
        }
    )


@login_required(login_url="management_login")
def commercial_offer_email_log_detail_api(request, log_id: int):
    if not user_is_management(request.user):
        return JsonResponse({"ok": False, "error": "forbidden"}, status=403)

    try:
        log = CommercialOfferEmailLog.objects.get(id=log_id, owner=request.user)
    except CommercialOfferEmailLog.DoesNotExist:
        return JsonResponse({"ok": False, "error": "not_found"}, status=404)

    return JsonResponse(
        {
            "ok": True,
            "log": {
                "id": log.id,
                "recipient_email": log.recipient_email,
                "recipient_name": log.recipient_name,
                "subject": log.subject,
                "preheader": log.preheader,
                "mode": log.mode,
                "segment_mode": log.segment_mode,
                "subject_preset": log.subject_preset,
                "subject_custom": log.subject_custom,
                "cta_type": getattr(log, "cta_type", ""),
                "cta_url": getattr(log, "cta_url", ""),
                "cta_custom_url": getattr(log, "cta_custom_url", ""),
                "cta_button_text": getattr(log, "cta_button_text", ""),
                "cta_microtext": getattr(log, "cta_microtext", ""),
                "general_tg": getattr(log, "general_tg", ""),
                "pricing_mode": getattr(log, "pricing_mode", "OPT") or "OPT",
                "opt_tier": getattr(log, "opt_tier", "8_15") or "8_15",
                "drop_tee_price": getattr(log, "drop_tee_price", None),
                "drop_hoodie_price": getattr(log, "drop_hoodie_price", None),
                "dropship_loyalty_bonus": bool(getattr(log, "dropship_loyalty_bonus", False)),
                "include_catalog_link": bool(getattr(log, "include_catalog_link", True)),
                "include_wholesale_link": bool(getattr(log, "include_wholesale_link", True)),
                "include_dropship_link": bool(getattr(log, "include_dropship_link", True)),
                "include_instagram_link": bool(getattr(log, "include_instagram_link", True)),
                "include_site_link": bool(getattr(log, "include_site_link", True)),
                "gallery_urls": getattr(log, "gallery_urls", []) or [],
                "gallery_items": getattr(log, "gallery_items", []) or [],
                "tee_entry": log.tee_entry,
                "tee_retail_example": log.tee_retail_example,
                "tee_profit": log.tee_profit,
                "hoodie_entry": log.hoodie_entry,
                "hoodie_retail_example": log.hoodie_retail_example,
                "hoodie_profit": log.hoodie_profit,
                "show_manager": log.show_manager,
                "manager_name": log.manager_name,
                "phone_enabled": bool(getattr(log, "phone_enabled", False)),
                "phone": log.phone,
                "viber": log.viber,
                "whatsapp": log.whatsapp,
                "telegram": log.telegram,
                "status": log.status,
                "error": log.error,
                "created_at_display": timezone.localtime(log.created_at).strftime("%d.%m.%Y %H:%M"),
                "body_html": log.body_html,
                "body_text": log.body_text,
            },
        }
    )


@require_POST
@login_required(login_url='management_login')
def commercial_offer_email_send_api(request):
    if not user_is_management(request.user):
        return JsonResponse({"ok": False, "error": "forbidden"}, status=403)

    form = CommercialOfferEmailForm(request.POST, user=request.user)
    if not form.is_valid():
        return JsonResponse({"ok": False, "errors": form.errors}, status=400)

    cd = form.cleaned_data
    recipient_email = (cd.get("recipient_email") or "").strip()

    duplicate_qs = CommercialOfferEmailLog.objects.filter(
        recipient_email__iexact=recipient_email,
        status=CommercialOfferEmailLog.Status.SENT,
    )
    if duplicate_qs.exists() and cd.get("confirm_resend") != 1:
        last = duplicate_qs.order_by("-created_at").first()
        last_display = timezone.localtime(last.created_at).strftime("%d.%m.%Y %H:%M") if last else ""
        return JsonResponse(
            {
                "ok": False,
                "needs_confirmation": True,
                "message": "На цей email вже відправляли комерційну пропозицію. Відправити ще раз?",
                "duplicate": {"count": duplicate_qs.count(), "lastSentAtDisplay": last_display},
            },
            status=409,
        )

    settings_obj, _ = CommercialOfferEmailSettings.objects.get_or_create(owner=request.user)

    default_name = (settings_obj.manager_name or "").strip() or _default_manager_name(request.user)
    settings_obj.show_manager = bool(cd.get("show_manager"))
    settings_obj.manager_name = (cd.get("manager_name") or "").strip() or default_name
    settings_obj.phone_enabled = bool(cd.get("phone_enabled"))
    settings_obj.phone = (cd.get("phone") or "").strip()
    settings_obj.viber_enabled = bool(cd.get("viber_enabled"))
    settings_obj.viber = (cd.get("viber") or "").strip()
    settings_obj.whatsapp_enabled = bool(cd.get("whatsapp_enabled"))
    settings_obj.whatsapp = (cd.get("whatsapp") or "").strip()
    settings_obj.telegram_enabled = bool(cd.get("telegram_enabled"))
    settings_obj.telegram = (cd.get("telegram") or "").strip()

    settings_obj.general_tg = (cd.get("general_tg") or "").strip()
    settings_obj.pricing_mode = (cd.get("pricing_mode") or "OPT").strip().upper()
    settings_obj.opt_tier = (cd.get("opt_tier") or "8_15").strip().upper()
    settings_obj.drop_tee_price = cd.get("drop_tee_price") or None
    settings_obj.drop_hoodie_price = cd.get("drop_hoodie_price") or None
    settings_obj.dropship_loyalty_bonus = bool(cd.get("dropship_loyalty_bonus"))
    settings_obj.include_catalog_link = bool(cd.get("include_catalog_link"))
    settings_obj.include_wholesale_link = bool(cd.get("include_wholesale_link"))
    settings_obj.include_dropship_link = bool(cd.get("include_dropship_link"))
    settings_obj.include_instagram_link = bool(cd.get("include_instagram_link"))
    settings_obj.include_site_link = bool(cd.get("include_site_link"))

    settings_obj.cta_type = (cd.get("cta_type") or "").strip().upper()
    settings_obj.cta_custom_url = (cd.get("cta_custom_url") or "").strip()
    settings_obj.cta_button_text = (cd.get("cta_button_text") or "").strip()
    settings_obj.cta_microtext = (cd.get("cta_microtext") or "").strip()

    settings_obj.gallery_neutral = cd.get("gallery_neutral") or []
    settings_obj.gallery_edgy = cd.get("gallery_edgy") or []
    settings_obj.gallery_initialized = True

    settings_obj.mode = (cd.get("mode") or "VISUAL").strip().upper()
    settings_obj.segment_mode = (cd.get("segment_mode") or "NEUTRAL").strip().upper()
    settings_obj.subject_preset = (cd.get("subject_preset") or "PRESET_1").strip().upper()
    settings_obj.subject_custom = (cd.get("subject_custom") or "").strip()
    settings_obj.tee_entry = cd.get("tee_entry") or None
    settings_obj.tee_retail_example = cd.get("tee_retail_example") or None
    settings_obj.hoodie_entry = cd.get("hoodie_entry") or None
    settings_obj.hoodie_retail_example = cd.get("hoodie_retail_example") or None
    settings_obj.save()

    payload = {
        "shop_name": (cd.get("recipient_name") or "").strip(),
        "mode": settings_obj.mode,
        "segment_mode": settings_obj.segment_mode,
        "subject_preset": settings_obj.subject_preset,
        "subject_custom": settings_obj.subject_custom,
        "cta_type": settings_obj.cta_type,
        "cta_custom_url": settings_obj.cta_custom_url,
        "cta_button_text": settings_obj.cta_button_text,
        "cta_microtext": settings_obj.cta_microtext,
        "tee_entry": settings_obj.tee_entry,
        "tee_retail_example": settings_obj.tee_retail_example,
        "hoodie_entry": settings_obj.hoodie_entry,
        "hoodie_retail_example": settings_obj.hoodie_retail_example,
        "show_manager": settings_obj.show_manager,
        "manager_name": settings_obj.manager_name,
        "phone_enabled": settings_obj.phone_enabled,
        "phone": settings_obj.phone if (settings_obj.show_manager and settings_obj.phone_enabled) else "",
        "viber": settings_obj.viber if (settings_obj.show_manager and settings_obj.viber_enabled) else "",
        "whatsapp": settings_obj.whatsapp if (settings_obj.show_manager and settings_obj.whatsapp_enabled) else "",
        "telegram": settings_obj.telegram if (settings_obj.show_manager and settings_obj.telegram_enabled) else "",
        "general_tg": settings_obj.general_tg,
        "pricing_mode": getattr(settings_obj, "pricing_mode", "OPT") or "OPT",
        "opt_tier": getattr(settings_obj, "opt_tier", "8_15") or "8_15",
        "drop_tee_price": getattr(settings_obj, "drop_tee_price", None),
        "drop_hoodie_price": getattr(settings_obj, "drop_hoodie_price", None),
        "dropship_loyalty_bonus": bool(getattr(settings_obj, "dropship_loyalty_bonus", False)),
        "include_catalog_link": settings_obj.include_catalog_link,
        "include_wholesale_link": settings_obj.include_wholesale_link,
        "include_dropship_link": settings_obj.include_dropship_link,
        "include_instagram_link": settings_obj.include_instagram_link,
        "include_site_link": settings_obj.include_site_link,
        "gallery_urls": settings_obj.gallery_edgy if settings_obj.segment_mode == "EDGY" else settings_obj.gallery_neutral,
        "manager_photo_url": _manager_photo_url(request.user, request) if settings_obj.show_manager else "",
    }
    email_build = build_twocomms_cp_email(payload)
    subject = email_build["subject"]
    preheader = email_build["preheader"]
    text_body = email_build["text"]
    html_body_to_send = email_build["html"] if settings_obj.mode == "VISUAL" else email_build["html_light"]

    from_email = getattr(settings, "DEFAULT_FROM_EMAIL", None) or "TwoComms <cooperation@twocomms.shop>"
    reply_to = [settings.EMAIL_HOST_USER] if getattr(settings, "EMAIL_HOST_USER", "") else None

    status = CommercialOfferEmailLog.Status.SENT
    error_text = ""
    try:
        msg = EmailMultiAlternatives(
            subject=subject,
            body=text_body,
            from_email=from_email,
            to=[recipient_email],
            reply_to=reply_to,
        )
        msg.attach_alternative(html_body_to_send, "text/html")
        msg.send(fail_silently=False)
    except Exception as exc:
        status = CommercialOfferEmailLog.Status.FAILED
        error_text = str(exc)

    log = CommercialOfferEmailLog.objects.create(
        owner=request.user,
        recipient_email=recipient_email,
        recipient_name=(cd.get("recipient_name") or "").strip(),
        subject=subject,
        preheader=preheader,
        body_html=html_body_to_send,
        body_text=text_body,
        mode=settings_obj.mode,
        segment_mode=settings_obj.segment_mode,
        subject_preset=settings_obj.subject_preset,
        subject_custom=settings_obj.subject_custom,
        cta_type=email_build.get("cta_type") or settings_obj.cta_type,
        cta_url=email_build.get("cta_url") or "",
        cta_custom_url=settings_obj.cta_custom_url,
        cta_button_text=email_build.get("cta_button_text") or settings_obj.cta_button_text,
        cta_microtext=email_build.get("cta_microtext") or settings_obj.cta_microtext,
        general_tg=settings_obj.general_tg,
        pricing_mode=email_build.get("pricing_mode") or getattr(settings_obj, "pricing_mode", "OPT") or "OPT",
        opt_tier=email_build.get("opt_tier") or getattr(settings_obj, "opt_tier", "8_15") or "8_15",
        drop_tee_price=email_build.get("drop_tee_price") or getattr(settings_obj, "drop_tee_price", None),
        drop_hoodie_price=email_build.get("drop_hoodie_price") or getattr(settings_obj, "drop_hoodie_price", None),
        dropship_loyalty_bonus=bool(
            email_build.get("dropship_loyalty_bonus")
            if ("dropship_loyalty_bonus" in email_build)
            else getattr(settings_obj, "dropship_loyalty_bonus", False)
        ),
        include_catalog_link=settings_obj.include_catalog_link,
        include_wholesale_link=settings_obj.include_wholesale_link,
        include_dropship_link=settings_obj.include_dropship_link,
        include_instagram_link=settings_obj.include_instagram_link,
        include_site_link=settings_obj.include_site_link,
        gallery_urls=email_build.get("gallery_urls") or [],
        gallery_items=email_build.get("gallery_items") or [],
        tee_entry=email_build.get("tee_entry"),
        tee_retail_example=email_build.get("tee_retail_example"),
        tee_profit=email_build.get("tee_profit"),
        hoodie_entry=email_build.get("hoodie_entry"),
        hoodie_retail_example=email_build.get("hoodie_retail_example"),
        hoodie_profit=email_build.get("hoodie_profit"),
        show_manager=settings_obj.show_manager,
        manager_name=settings_obj.manager_name if settings_obj.show_manager else "",
        phone_enabled=settings_obj.phone_enabled,
        phone=settings_obj.phone if (settings_obj.show_manager and settings_obj.phone_enabled) else "",
        viber=settings_obj.viber if (settings_obj.show_manager and settings_obj.viber_enabled) else "",
        whatsapp=settings_obj.whatsapp if (settings_obj.show_manager and settings_obj.whatsapp_enabled) else "",
        telegram=settings_obj.telegram if (settings_obj.show_manager and settings_obj.telegram_enabled) else "",
        status=status,
        error=error_text,
    )

    row_html = render_to_string("management/partials/commercial_offer_log_row.html", {"log": log})

    return JsonResponse(
        {
            "ok": True,
            "status": status,
            "sent": status == CommercialOfferEmailLog.Status.SENT,
            "row_html": row_html,
            "created_at_display": timezone.localtime(log.created_at).strftime("%d.%m.%Y %H:%M"),
            "message": "КП успішно надіслано." if status == CommercialOfferEmailLog.Status.SENT else "Не вдалося відправити лист.",
            "error": error_text,
        }
    )


@login_required(login_url='management_login')
def info(request):
    if not user_is_management(request.user):
        return redirect('management_login')

    stats = get_user_stats(request.user)
    report_sent_today = has_report_today(request.user)
    progress_clients_pct = min(100, int(stats['processed_today'] / TARGET_CLIENTS_DAY * 100)) if TARGET_CLIENTS_DAY else 0
    progress_points_pct = min(100, int(stats['points_today'] / TARGET_POINTS_DAY * 100)) if TARGET_POINTS_DAY else 0
    bot_username = get_manager_bot_username()
    reminders = get_reminders(request.user, stats=stats, report_sent=report_sent_today)

    faq_items = [
        {
            'q': 'Коли можна вивести нарахування?',
            'a': 'Незабаром з’явиться.',
        },
        {
            'q': 'Чому частина балансу заморожена на 14 днів?',
            'a': 'Незабаром з’явиться.',
        },
        {
            'q': 'Коли виплачується ставка і як формується комісія?',
            'a': 'Незабаром з’явиться.',
        },
        {
            'q': 'Як оновити картку для виплат?',
            'a': 'Незабаром з’явиться.',
        },
        {
            'q': 'Як прив’язати Telegram-бота менеджменту?',
            'a': 'Незабаром з’явиться.',
        },
        {
            'q': 'Куди писати, якщо є питання щодо виплат?',
            'a': 'Незабаром з’явиться.',
        },
    ]

    return render(
        request,
        'management/info.html',
        {
            'faq_items': faq_items,
            'user_points_today': stats['points_today'],
            'user_points_total': stats['points_total'],
            'processed_today': stats['processed_today'],
            'target_clients': TARGET_CLIENTS_DAY,
            'target_points': TARGET_POINTS_DAY,
            'progress_clients_pct': progress_clients_pct,
            'progress_points_pct': progress_points_pct,
            'has_report_today': report_sent_today,
            'reminders': reminders,
            'manager_bot_username': bot_username,
        }
    )


@login_required(login_url='management_login')
def payouts(request):
    if not user_is_management(request.user):
        return redirect('management_login')

    import re
    from decimal import Decimal

    from django.db import models
    from django.db.models import Sum
    from django.db.models.functions import Coalesce

    from orders.models import WholesaleInvoice
    from management.models import ManagerCommissionAccrual, ManagerPayoutRequest, NightlyScoreSnapshot

    def mask_card(raw):
        s = (raw or '').strip()
        digits = re.sub(r'\D+', '', s)
        if len(digits) < 4:
            return '—'
        return '**** ' + digits[-4:]

    now = timezone.now()
    zero = Decimal('0')
    money_field = models.DecimalField(max_digits=12, decimal_places=2)

    accruals_qs = ManagerCommissionAccrual.objects.filter(owner=request.user)
    total_accrued = accruals_qs.aggregate(
        total=Coalesce(Sum('amount'), zero, output_field=money_field),
    )['total'] or zero

    frozen_amount = accruals_qs.filter(frozen_until__gt=now).aggregate(
        total=Coalesce(Sum('amount'), zero, output_field=money_field),
    )['total'] or zero

    paid_total = ManagerPayoutRequest.objects.filter(
        owner=request.user,
        status=ManagerPayoutRequest.Status.PAID,
    ).aggregate(
        total=Coalesce(Sum('amount'), zero, output_field=money_field),
    )['total'] or zero

    reserved_amount = ManagerPayoutRequest.objects.filter(
        owner=request.user,
        status__in=[ManagerPayoutRequest.Status.PROCESSING, ManagerPayoutRequest.Status.APPROVED],
    ).aggregate(
        total=Coalesce(Sum('amount'), zero, output_field=money_field),
    )['total'] or zero

    balance = (total_accrued - paid_total)
    available = balance - frozen_amount - reserved_amount
    if available < 0:
        available = zero

    active_request = ManagerPayoutRequest.objects.filter(
        owner=request.user,
        status__in=[ManagerPayoutRequest.Status.PROCESSING, ManagerPayoutRequest.Status.APPROVED],
    ).order_by('-created_at').first()

    latest_snapshot = NightlyScoreSnapshot.objects.filter(owner=request.user).order_by('-snapshot_date').first()

    history = list(
        ManagerPayoutRequest.objects.filter(owner=request.user).order_by('-created_at')[:50]
    )
    frozen_items = list(
        accruals_qs.filter(frozen_until__gt=now)
        .order_by('frozen_until')
        .values('amount', 'freeze_reason_text', 'note', 'frozen_until', 'freeze_reason_code')[:10]
    )

    deals_count = WholesaleInvoice.objects.filter(created_by=request.user, payment_status='paid').count()

    first_deal_at = accruals_qs.order_by('created_at').values_list('created_at', flat=True).first()
    last_payout_at = ManagerPayoutRequest.objects.filter(
        owner=request.user,
        status=ManagerPayoutRequest.Status.PAID,
        paid_at__isnull=False,
    ).order_by('-paid_at').values_list('paid_at', flat=True).first()

    total_received = paid_total

    try:
        prof = request.user.userprofile
    except Exception:
        prof = None

    position = (getattr(prof, 'manager_position', '') or '').strip() if prof else ''
    base_salary = getattr(prof, 'manager_base_salary_uah', 0) if prof else 0
    commission_percent = getattr(prof, 'manager_commission_percent', 0) if prof else 0

    started_at = getattr(prof, 'manager_started_at', None) if prof else None
    if not started_at:
        try:
            started_at = timezone.localtime(request.user.date_joined).date() if request.user.date_joined else None
        except Exception:
            started_at = None

    days_worked = None
    try:
        if started_at:
            days_worked = max(0, (timezone.localdate() - started_at).days)
    except Exception:
        days_worked = None

    card_mask = mask_card(getattr(prof, 'payment_details', '') if prof else '')

    stats = get_user_stats(request.user)
    report_sent_today = has_report_today(request.user)
    progress_clients_pct = min(100, int(stats['processed_today'] / TARGET_CLIENTS_DAY * 100)) if TARGET_CLIENTS_DAY else 0
    progress_points_pct = min(100, int(stats['points_today'] / TARGET_POINTS_DAY * 100)) if TARGET_POINTS_DAY else 0
    bot_username = get_manager_bot_username()
    reminders = get_reminders(request.user, stats=stats, report_sent=report_sent_today)
    salary_simulator = build_salary_simulator(
        user=request.user,
        summary={
            'invoices': {
                'amount': str(
                    WholesaleInvoice.objects.filter(created_by=request.user, payment_status='paid')
                    .aggregate(total=models.Sum('total_amount'))
                    .get('total')
                    or 0
                ),
                'paid': deals_count,
            }
        },
        shadow_score={
            'mosaic_score': float(latest_snapshot.mosaic_score or 0) if latest_snapshot else 0,
            'confidence_band': ((latest_snapshot.payload or {}).get('trust') or {}).get('confidence_band', 'LOW') if latest_snapshot else 'LOW',
            'freshness_seconds': int(latest_snapshot.freshness_seconds or 0) if latest_snapshot else 0,
        },
    )

    return render(
        request,
        'management/payouts.html',
        {
            'balance': balance,
            'available': available,
            'frozen_amount': frozen_amount,
            'reserved_amount': reserved_amount,
            'active_request': active_request,
            'history': history,
            'frozen_items': frozen_items,
            'position': position or '—',
            'days_worked': days_worked,
            'deals_count': deals_count,
            'base_salary': base_salary,
            'commission_percent': commission_percent,
            'first_deal_at': first_deal_at,
            'last_payout_at': last_payout_at,
            'total_received': total_received,
            'card_mask': card_mask,
            'user_points_today': stats['points_today'],
            'user_points_total': stats['points_total'],
            'processed_today': stats['processed_today'],
            'target_clients': TARGET_CLIENTS_DAY,
            'target_points': TARGET_POINTS_DAY,
            'progress_clients_pct': progress_clients_pct,
            'progress_points_pct': progress_points_pct,
            'has_report_today': report_sent_today,
            'reminders': reminders,
            'manager_bot_username': bot_username,
            'salary_simulator': salary_simulator,
        }
    )


@login_required(login_url='management_login')
@require_POST
def payouts_request_api(request):
    if not user_is_management(request.user):
        return JsonResponse({'ok': False}, status=403)

    import json
    from decimal import Decimal

    from django.db import models, transaction
    from django.db.models import Sum
    from django.db.models.functions import Coalesce

    from management.models import ManagerCommissionAccrual, ManagerPayoutRequest

    try:
        prof = request.user.userprofile
    except Exception:
        prof = None

    card_details = (getattr(prof, 'payment_details', '') or '').strip() if prof else ''
    if not card_details:
        return JsonResponse({'ok': False, 'error': 'Вкажіть картку в профілі перед запитом на виплату.'}, status=400)

    try:
        payload = json.loads(request.body.decode('utf-8') or '{}')
    except Exception:
        payload = {}

    money_field = models.DecimalField(max_digits=12, decimal_places=2)
    zero = Decimal('0')
    now = timezone.now()

    def calc_available():
        accruals_qs = ManagerCommissionAccrual.objects.filter(owner=request.user)
        total_accrued = accruals_qs.aggregate(
            total=Coalesce(Sum('amount'), zero, output_field=money_field),
        )['total'] or zero
        frozen_amount = accruals_qs.filter(frozen_until__gt=now).aggregate(
            total=Coalesce(Sum('amount'), zero, output_field=money_field),
        )['total'] or zero
        paid_total = ManagerPayoutRequest.objects.filter(
            owner=request.user,
            status=ManagerPayoutRequest.Status.PAID,
        ).aggregate(
            total=Coalesce(Sum('amount'), zero, output_field=money_field),
        )['total'] or zero
        reserved_amount = ManagerPayoutRequest.objects.filter(
            owner=request.user,
            status__in=[ManagerPayoutRequest.Status.PROCESSING, ManagerPayoutRequest.Status.APPROVED],
        ).aggregate(
            total=Coalesce(Sum('amount'), zero, output_field=money_field),
        )['total'] or zero
        balance = total_accrued - paid_total
        available_val = balance - frozen_amount - reserved_amount
        if available_val < 0:
            available_val = zero
        return available_val

    amount_raw = payload.get('amount')
    try:
        amount = Decimal(str(amount_raw)) if amount_raw is not None else None
    except Exception:
        amount = None

    with transaction.atomic():
        existing = ManagerPayoutRequest.objects.select_for_update().filter(
            owner=request.user,
            status__in=[ManagerPayoutRequest.Status.PROCESSING, ManagerPayoutRequest.Status.APPROVED],
        ).first()
        if existing:
            return JsonResponse({'ok': False, 'error': 'У вас вже є активний запит на виплату.'}, status=400)

        available = calc_available()
        if amount is None:
            amount = available

        if amount <= 0:
            return JsonResponse({'ok': False, 'error': 'Недостатньо коштів для виплати.'}, status=400)
        if amount > available:
            return JsonResponse({'ok': False, 'error': 'Сума перевищує доступний баланс.'}, status=400)

        req = ManagerPayoutRequest.objects.create(
            owner=request.user,
            kind=ManagerPayoutRequest.Kind.REQUEST,
            amount=amount,
            status=ManagerPayoutRequest.Status.PROCESSING,
        )

    # Telegram notifications (best-effort)
    try:
        import re
        digits = re.sub(r'\D+', '', card_details)
        card_mask = ('**** ' + digits[-4:]) if len(digits) >= 4 else '—'

        _notify_manager_payout(
            req,
            title='💸 <b>Запит на виплату створено</b>',
            body_lines=[
                f"Ви запросили виплату на суму <b>{escape(str(req.amount))} грн</b>.",
                f"Картка: <code>{escape(card_mask)}</code>",
                'Статус: ⏳ <b>в обробці</b>.',
            ],
        )
        _send_payout_request_to_admin(req)
    except Exception:
        pass

    return JsonResponse({
        'ok': True,
        'request': {
            'id': req.id,
            'status': req.status,
            'amount': str(req.amount),
            'created_at': timezone.localtime(req.created_at).strftime('%d.%m.%Y %H:%M') if req.created_at else '',
        }
    })


@login_required(login_url='management_login')
@require_POST
def admin_payout_settings_save_api(request):
    if not request.user.is_staff:
        return JsonResponse({'ok': False}, status=403)

    import json
    from decimal import Decimal

    from django.contrib.auth import get_user_model

    try:
        payload = json.loads(request.body.decode('utf-8') or '{}')
    except Exception:
        payload = {}

    try:
        user_id = int(payload.get('user_id'))
    except Exception:
        return JsonResponse({'ok': False, 'error': 'Невірний user_id'}, status=400)

    position = (payload.get('position') or '').strip()

    try:
        base_salary = int(payload.get('base_salary') or 0)
    except Exception:
        base_salary = 0
    base_salary = max(0, min(10_000_000, base_salary))

    try:
        percent = Decimal(str(payload.get('percent') or 0))
    except Exception:
        percent = Decimal('0')
    if percent < 0:
        percent = Decimal('0')
    if percent > 100:
        percent = Decimal('100')

    started_at = None
    started_at_raw = (payload.get('started_at') or '').strip()
    if started_at_raw:
        try:
            from datetime import datetime
            started_at = datetime.strptime(started_at_raw, '%Y-%m-%d').date()
        except Exception:
            started_at = None

    User = get_user_model()
    user = User.objects.filter(id=user_id).select_related('userprofile').first()
    if not user:
        return JsonResponse({'ok': False, 'error': 'Користувача не знайдено'}, status=404)

    profile, _ = UserProfile.objects.get_or_create(user=user)
    profile.manager_position = position
    profile.manager_base_salary_uah = base_salary
    profile.manager_commission_percent = percent
    profile.manager_started_at = started_at
    profile.save()

    return JsonResponse({'ok': True})


@login_required(login_url='management_login')
@require_POST
def admin_payout_approve_api(request, request_id):
    if not request.user.is_staff:
        return JsonResponse({'ok': False}, status=403)

    from django.db import transaction
    from management.models import ManagerPayoutRequest, PayoutRejectionReasonRequest

    with transaction.atomic():
        req = ManagerPayoutRequest.objects.select_for_update().filter(id=request_id).first()
        if not req:
            return JsonResponse({'ok': False, 'error': 'Запит не знайдено'}, status=404)

        if req.status in (ManagerPayoutRequest.Status.REJECTED, ManagerPayoutRequest.Status.PAID):
            return JsonResponse({'ok': False, 'error': 'Запит вже оброблено'}, status=400)

        req.status = ManagerPayoutRequest.Status.APPROVED
        req.approved_at = timezone.now()
        req.rejected_at = None
        req.rejection_reason = ''
        req.processed_by = request.user
        req.save(update_fields=['status', 'approved_at', 'rejected_at', 'rejection_reason', 'processed_by'])

        PayoutRejectionReasonRequest.objects.filter(payout_request=req, is_active=True).update(is_active=False)

    # Telegram sync (best-effort)
    try:
        req_full = ManagerPayoutRequest.objects.select_related('owner', 'owner__userprofile').filter(id=req.id).first()
        if req_full:
            _try_update_admin_payout_message(req_full, final=False)

            import re
            card_raw = ''
            try:
                card_raw = getattr(req_full.owner.userprofile, 'payment_details', '')
            except Exception:
                card_raw = ''
            digits = re.sub(r'\D+', '', (card_raw or ''))
            card_mask = ('**** ' + digits[-4:]) if len(digits) >= 4 else '—'

            _notify_manager_payout(
                req_full,
                title='✅ <b>Виплату схвалено</b>',
                body_lines=[
                    f"Сума: <b>{escape(str(req_full.amount))} грн</b>.",
                    f"Протягом 3 годин сума буде зарахована на картку <code>{escape(card_mask)}</code>.",
                ],
            )
    except Exception:
        pass

    return JsonResponse({'ok': True})


@login_required(login_url='management_login')
@require_POST
def admin_payout_reject_api(request, request_id):
    if not request.user.is_staff:
        return JsonResponse({'ok': False}, status=403)

    import json
    from django.db import transaction
    from management.models import ManagerPayoutRequest, PayoutRejectionReasonRequest

    try:
        payload = json.loads(request.body.decode('utf-8') or '{}')
    except Exception:
        payload = {}

    reason = (payload.get('reason') or '').strip()
    if not reason:
        return JsonResponse({'ok': False, 'error': 'Вкажіть причину'}, status=400)

    with transaction.atomic():
        req = ManagerPayoutRequest.objects.select_for_update().filter(id=request_id).first()
        if not req:
            return JsonResponse({'ok': False, 'error': 'Запит не знайдено'}, status=404)

        if req.status in (ManagerPayoutRequest.Status.REJECTED, ManagerPayoutRequest.Status.PAID):
            return JsonResponse({'ok': False, 'error': 'Запит вже оброблено'}, status=400)

        req.status = ManagerPayoutRequest.Status.REJECTED
        req.rejection_reason = reason
        req.rejected_at = timezone.now()
        req.approved_at = None
        req.processed_by = request.user
        req.save(update_fields=['status', 'rejection_reason', 'rejected_at', 'approved_at', 'processed_by'])

        PayoutRejectionReasonRequest.objects.filter(payout_request=req, is_active=True).update(is_active=False)

    # Telegram sync (best-effort)
    try:
        req_full = ManagerPayoutRequest.objects.select_related('owner', 'owner__userprofile').filter(id=req.id).first()
        if req_full:
            _try_update_admin_payout_message(req_full, final=True)
            _notify_manager_payout(
                req_full,
                title='❌ <b>Виплату відхилено</b>',
                body_lines=[
                    f"Сума: <b>{escape(str(req_full.amount))} грн</b>.",
                    f"Причина: {escape(reason)}",
                ],
            )
    except Exception:
        pass

    return JsonResponse({'ok': True})


@login_required(login_url='management_login')
@require_POST
def admin_payout_paid_api(request, request_id):
    if not request.user.is_staff:
        return JsonResponse({'ok': False}, status=403)

    from django.db import transaction
    from management.models import ManagerPayoutRequest, PayoutRejectionReasonRequest

    with transaction.atomic():
        req = ManagerPayoutRequest.objects.select_for_update().filter(id=request_id).first()
        if not req:
            return JsonResponse({'ok': False, 'error': 'Запит не знайдено'}, status=404)

        if req.status == ManagerPayoutRequest.Status.PAID:
            return JsonResponse({'ok': False, 'error': 'Вже виплачено'}, status=400)
        if req.status == ManagerPayoutRequest.Status.REJECTED:
            return JsonResponse({'ok': False, 'error': 'Запит відхилено'}, status=400)

        req.status = ManagerPayoutRequest.Status.PAID
        req.paid_at = timezone.now()
        req.processed_by = request.user
        req.save(update_fields=['status', 'paid_at', 'processed_by'])

        PayoutRejectionReasonRequest.objects.filter(payout_request=req, is_active=True).update(is_active=False)

    # Telegram sync (best-effort)
    try:
        req_full = ManagerPayoutRequest.objects.select_related('owner', 'owner__userprofile').filter(id=req.id).first()
        if req_full:
            _try_update_admin_payout_message(req_full, final=True)

            import re
            card_raw = ''
            try:
                card_raw = getattr(req_full.owner.userprofile, 'payment_details', '')
            except Exception:
                card_raw = ''
            digits = re.sub(r'\D+', '', (card_raw or ''))
            card_mask = ('**** ' + digits[-4:]) if len(digits) >= 4 else '—'

            _notify_manager_payout(
                req_full,
                title='💳 <b>Виплату здійснено</b>',
                body_lines=[
                    f"Сума: <b>{escape(str(req_full.amount))} грн</b>.",
                    f"Зачислено на карту <code>{escape(card_mask)}</code>.",
                ],
            )
    except Exception:
        pass

    return JsonResponse({'ok': True})


@login_required(login_url='management_login')
@require_POST
def admin_payout_manual_create_api(request):
    if not request.user.is_staff:
        return JsonResponse({'ok': False}, status=403)

    import json
    from decimal import Decimal

    from management.models import ManagerPayoutRequest

    try:
        payload = json.loads(request.body.decode('utf-8') or '{}')
    except Exception:
        payload = {}

    try:
        user_id = int(payload.get('user_id'))
    except Exception:
        return JsonResponse({'ok': False, 'error': 'Невірний user_id'}, status=400)

    try:
        amount = Decimal(str(payload.get('amount') or '0'))
    except Exception:
        amount = Decimal('0')

    if amount <= 0:
        return JsonResponse({'ok': False, 'error': 'Невірна сума'}, status=400)

    paid_at = None
    paid_at_raw = (payload.get('paid_at') or '').strip()
    if paid_at_raw:
        try:
            from datetime import datetime
            paid_at = datetime.strptime(paid_at_raw, '%Y-%m-%d').date()
        except Exception:
            paid_at = None

    User = get_user_model()
    user = User.objects.filter(id=user_id).first()
    if not user:
        return JsonResponse({'ok': False, 'error': 'Користувача не знайдено'}, status=404)

    paid_dt = timezone.now()
    if paid_at:
        try:
            from datetime import datetime, time
            tz = timezone.get_current_timezone()
            paid_dt = timezone.make_aware(datetime.combine(paid_at, time(12, 0)), tz)
        except Exception:
            paid_dt = timezone.now()

    req = ManagerPayoutRequest.objects.create(
        owner=user,
        kind=ManagerPayoutRequest.Kind.MANUAL,
        amount=amount,
        status=ManagerPayoutRequest.Status.PAID,
        paid_at=paid_dt,
        processed_by=request.user,
    )

    return JsonResponse({'ok': True, 'id': req.id})


@login_required(login_url='management_login')
@require_POST
def admin_payout_adjust_api(request):
    if not request.user.is_staff:
        return JsonResponse({'ok': False}, status=403)

    import json
    from decimal import Decimal

    from management.models import ManagerCommissionAccrual

    try:
        payload = json.loads(request.body.decode('utf-8') or '{}')
    except Exception:
        payload = {}

    try:
        user_id = int(payload.get('user_id'))
    except Exception:
        return JsonResponse({'ok': False, 'error': 'Невірний user_id'}, status=400)

    direction = (payload.get('direction') or 'credit').strip().lower()
    if direction not in ('credit', 'debit'):
        direction = 'credit'

    try:
        amount = Decimal(str(payload.get('amount') or '0'))
    except Exception:
        amount = Decimal('0')

    if amount <= 0:
        return JsonResponse({'ok': False, 'error': 'Невірна сума'}, status=400)

    note = (payload.get('note') or '').strip()[:255]

    frozen_until_raw = (payload.get('frozen_until') or '').strip()
    frozen_dt = timezone.now()
    if frozen_until_raw:
        try:
            from datetime import datetime, time
            tz = timezone.get_current_timezone()
            frozen_date = datetime.strptime(frozen_until_raw, '%Y-%m-%d').date()
            frozen_dt = timezone.make_aware(datetime.combine(frozen_date, time(12, 0)), tz)
        except Exception:
            frozen_dt = timezone.now()

    User = get_user_model()
    user = User.objects.filter(id=user_id).first()
    if not user:
        return JsonResponse({'ok': False, 'error': 'Користувача не знайдено'}, status=404)

    signed_amount = amount if direction == 'credit' else (-amount)

    ManagerCommissionAccrual.objects.create(
        owner=user,
        base_amount=Decimal('0'),
        percent=Decimal('0'),
        amount=signed_amount,
        frozen_until=frozen_dt,
        note=note,
    )

    return JsonResponse({'ok': True})
